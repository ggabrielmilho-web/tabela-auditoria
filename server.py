"""
Auditoria Receita — Backend
Rode: python server.py
Acesse: http://localhost:5000
"""

import os
import io
import json
import time
import tempfile
import functools
import psycopg2
import requests
import pgr
import placas
from flask import Flask, Response, jsonify, send_from_directory, request, session, redirect, url_for, send_file, stream_with_context
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_folder='.')
app.secret_key = os.getenv('SECRET_KEY', 'dev-secret-change-me')
CORS(app, supports_credentials=True)

# ── Config Power BI ──
CONFIG = {
    'tenant_id':       os.getenv('POWERBI_TENANT_ID', ''),
    'client_id':       os.getenv('POWERBI_CLIENT_ID', ''),
    'client_secret':   os.getenv('POWERBI_CLIENT_SECRET', ''),
    'dataset_id':      os.getenv('POWERBI_DATASET_ID', ''),       # Auditoria + Tarifas
    'group_id':        os.getenv('POWERBI_GROUP_ID', ''),
    'dre_dataset_id':  os.getenv('POWERBI_DRE_DATASET_ID', ''),   # DRE
    # tabelas.contabil — 456 extrato, 441 faturamento, 571 ACNI, 479 eventos.
    'contabil_dataset_id': os.getenv('POWERBI_CONTABIL_DATASET_ID', ''),
}

# ── Config Mercado Livre (OAuth) ──
ML_CONFIG = {
    'client_id':     os.getenv('ML_CLIENT_ID', ''),      # App ID do Mercado Livre
    'client_secret': os.getenv('ML_CLIENT_SECRET', ''),  # Secret Key do app
    'redirect_uri':  os.getenv('ML_REDIRECT_URI', 'https://rizza.carvalhoia.com/mercadolivre/callback'),
}
ML_TOKEN_URL = 'https://api.mercadolibre.com/oauth/token'


# ── Banco de dados ──
def get_db():
    return psycopg2.connect(
        host=os.getenv('DB_HOST', 'localhost'),
        port=os.getenv('DB_PORT', '5432'),
        dbname=os.getenv('DB_NAME', 'postgres'),
        user=os.getenv('DB_USER', 'postgres'),
        password=os.getenv('DB_PASSWORD', ''),
    )

# ── Decorators de autenticação ──
def login_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            if request.path.startswith('/api/'):
                return jsonify({'ok': False, 'error': 'Não autenticado'}), 401
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            if request.path.startswith('/api/'):
                return jsonify({'ok': False, 'error': 'Não autenticado'}), 401
            return redirect('/login')
        if session.get('role') != 'admin':
            if request.path.startswith('/api/'):
                return jsonify({'ok': False, 'error': 'Acesso negado'}), 403
            return redirect('/')
        return f(*args, **kwargs)
    return decorated

# Abas concedíveis por usuário (a aba Admin NÃO entra — é exclusiva de role=admin).
PAGINAS_VALIDAS = {'auditoria', 'tarifas', 'embarques', 'reuniao', 'dre',
                   'despesas', 'conhecimentos', 'faturamento', 'contratos', 'veiculos',
                   'pgr', 'contabil'}
# Chave da aba → rota inicial (para redirect sem loop).
_PAGINA_ROTA = {
    'auditoria': '/', 'tarifas': '/tarifas', 'embarques': '/embarques',
    'reuniao': '/reuniao', 'dre': '/dre', 'despesas': '/dre/despesas',
    'conhecimentos': '/dre/conhecimentos', 'faturamento': '/faturamento',
    'contratos': '/contratos', 'veiculos': '/veiculos', 'pgr': '/pgr',
    'contabil': '/contabil',
}
# Ordem de preferência ao escolher a primeira aba permitida.
# 'embarques' vem antes de 'tarifas': quem não tem Auditoria e tem Embarques
# cai direto no Embarques (só depois em Tarifas).
# 'pgr' fica junto da família de rastreamento (perto de Embarques): é segurança
# operacional, não financeiro.
_PAGINA_ORDEM = ['auditoria', 'embarques', 'pgr', 'tarifas', 'reuniao', 'dre',
                 'despesas', 'conhecimentos', 'faturamento', 'contratos', 'veiculos',
                 'contabil']


def _primeira_pagina_permitida():
    """Rota da 1ª aba concedida ao usuário (evita loop de redirect). Admin → '/'.
    Sem nenhuma aba → página neutra de 'sem acesso'."""
    if session.get('role') == 'admin':
        return '/'
    permitidas = session.get('paginas_permitidas') or []
    for chave in _PAGINA_ORDEM:
        if chave in permitidas:
            return _PAGINA_ROTA[chave]
    return '/sem-acesso'


def page_required(page_key):
    """Libera a rota se o usuário for admin (bypass) ou tiver a aba concedida."""
    def deco(f):
        @functools.wraps(f)
        def inner(*args, **kwargs):
            if 'user_id' not in session:
                if request.path.startswith('/api/'):
                    return jsonify({'ok': False, 'error': 'Não autenticado'}), 401
                return redirect('/login')
            if session.get('role') == 'admin':
                return f(*args, **kwargs)
            if page_key not in (session.get('paginas_permitidas') or []):
                if request.path.startswith('/api/'):
                    return jsonify({'ok': False, 'error': 'Acesso negado'}), 403
                return redirect(_primeira_pagina_permitida())
            return f(*args, **kwargs)
        return inner
    return deco


# ── Power BI helpers ──
def get_token():
    url = f"https://login.microsoftonline.com/{CONFIG['tenant_id']}/oauth2/v2.0/token"
    data = {
        'grant_type': 'client_credentials',
        'client_id': CONFIG['client_id'],
        'client_secret': CONFIG['client_secret'],
        'scope': 'https://analysis.windows.net/powerbi/api/.default'
    }
    resp = requests.post(url, data=data, timeout=30)
    resp.raise_for_status()
    return resp.json()['access_token']


def execute_dax(token, query, dataset_id=None):
    ds = dataset_id or CONFIG['dataset_id']
    url = (
        f"https://api.powerbi.com/v1.0/myorg/groups/"
        f"{CONFIG['group_id']}/datasets/{ds}/executeQueries"
    )
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/json'
    }
    body = {
        'queries': [{'query': query}],
        'serializerSettings': {'includeNulls': True}
    }
    resp = requests.post(url, json=body, headers=headers, timeout=120)
    resp.raise_for_status()
    return resp.json()


def clean_rows(rows):
    result = []
    for row in rows:
        clean = {}
        for k, v in row.items():
            short_key = k.split('[')[-1].rstrip(']') if '[' in k else k
            clean[short_key] = v
        result.append(clean)
    return result


# ════════════════════════════════════════
# ROTAS DE AUTENTICAÇÃO
# ════════════════════════════════════════

@app.route('/login', methods=['GET'])
def login_page():
    if 'user_id' in session:
        return redirect('/')
    return send_from_directory('.', 'login.html')


@app.route('/nav-perms.js')
def nav_perms_js():
    """Script de gating do menu por permissão de aba (servido a qualquer um)."""
    return send_from_directory('.', 'nav-perms.js', mimetype='application/javascript')


@app.route('/report-filter.js')
def report_filter_js():
    """Componente de AutoFilter (estilo Excel) dos relatórios densos."""
    return send_from_directory('.', 'report-filter.js', mimetype='application/javascript')


@app.route('/sem-acesso')
@login_required
def sem_acesso_page():
    """Página neutra para usuário sem nenhuma aba liberada (evita loop de redirect)."""
    return (
        "<!DOCTYPE html><html lang='pt-BR'><head><meta charset='UTF-8'>"
        "<title>Sem acesso</title><style>body{font-family:sans-serif;background:#0a0e17;"
        "color:#e2e8f0;display:flex;min-height:100vh;align-items:center;justify-content:center;"
        "margin:0;text-align:center}a{color:#38bdf8}</style></head><body><div>"
        "<h1>Sem acesso</h1><p>Seu usuário ainda não tem nenhuma aba liberada.<br>"
        "Contate o administrador.</p><p><a href='/logout'>Sair</a></p></div></body></html>"
    )


@app.route('/login', methods=['POST'])
def login_post():
    data = request.get_json() or {}
    email = data.get('email', '').strip().lower()
    senha = data.get('senha', '')

    if not email or not senha:
        return jsonify({'ok': False, 'error': 'Preencha e-mail e senha'}), 400

    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            "SELECT id, nome, password_hash, role, ativo, tipos_permitidos, paginas_permitidas FROM auditoria_users WHERE email = %s",
            (email,)
        )
        user = cur.fetchone()
        cur.close()
        conn.close()
    except Exception as e:
        return jsonify({'ok': False, 'error': f'Erro de banco: {str(e)}'}), 500

    if not user:
        return jsonify({'ok': False, 'error': 'E-mail ou senha inválidos'}), 401

    uid, nome, pw_hash, role, ativo, tipos_permitidos, paginas_permitidas = user

    if not ativo:
        return jsonify({'ok': False, 'error': 'Conta desativada. Contate o administrador.'}), 403

    if not check_password_hash(pw_hash, senha):
        return jsonify({'ok': False, 'error': 'E-mail ou senha inválidos'}), 401

    session['user_id']         = uid
    session['nome']            = nome
    session['role']            = role
    session['tipos_permitidos'] = tipos_permitidos or []
    session['paginas_permitidas'] = paginas_permitidas or []
    return jsonify({'ok': True, 'redirect': _primeira_pagina_permitida()})


@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')


def _ml_page(titulo, corpo):
    """Página neutra de retorno do Mercado Livre (mesmo visual do /sem-acesso)."""
    return (
        "<!DOCTYPE html><html lang='pt-BR'><head><meta charset='UTF-8'>"
        "<meta name='viewport' content='width=device-width, initial-scale=1'>"
        f"<title>{titulo}</title><style>body{{font-family:sans-serif;background:#0a0e17;"
        "color:#e2e8f0;display:flex;min-height:100vh;align-items:center;justify-content:center;"
        "margin:0;text-align:center;padding:1.5rem}code{background:#1e293b;padding:.2rem .4rem;"
        "border-radius:4px;word-break:break-all}a{color:#38bdf8}</style></head><body><div>"
        f"{corpo}</div></body></html>"
    )


@app.route('/mercadolivre/callback')
def mercadolivre_callback():
    """
    Retorno do OAuth do Mercado Livre.

    Fase 1: confirma que a autorização chegou. Se ML_CLIENT_ID/ML_CLIENT_SECRET
    estiverem configurados, já troca o `code` por um access_token e mostra o
    resultado. A persistência do token e o uso da API vêm na fase 2.
    """
    erro = request.args.get('error')
    code = request.args.get('code')

    if erro:
        desc = request.args.get('error_description', '')
        return _ml_page(
            'Autorização negada',
            f"<h1>Autorização não concluída</h1><p>Mercado Livre retornou: "
            f"<code>{erro}</code></p><p>{desc}</p>"
        ), 400

    if not code:
        return _ml_page(
            'Callback Mercado Livre',
            "<h1>Callback ativo ✅</h1><p>Esta é a URI de redirecionamento do app "
            "do Mercado Livre. Ela só recebe dados durante o fluxo de login.</p>"
        )

    # Sem credenciais configuradas ainda: só confirma que o retorno chegou.
    if not ML_CONFIG['client_id'] or not ML_CONFIG['client_secret']:
        return _ml_page(
            'Autorização recebida',
            "<h1>Autorização recebida ✅</h1>"
            f"<p>Código de autorização:</p><p><code>{code}</code></p>"
            "<p>Falta configurar <code>ML_CLIENT_ID</code> e "
            "<code>ML_CLIENT_SECRET</code> para trocar por um token.</p>"
        )

    # Troca o code por access_token.
    try:
        resp = requests.post(ML_TOKEN_URL, data={
            'grant_type':    'authorization_code',
            'client_id':     ML_CONFIG['client_id'],
            'client_secret': ML_CONFIG['client_secret'],
            'code':          code,
            'redirect_uri':  ML_CONFIG['redirect_uri'],
        }, headers={'Accept': 'application/json'}, timeout=20)
    except Exception as e:
        return _ml_page('Erro', f"<h1>Falha ao contatar o Mercado Livre</h1><p>{e}</p>"), 502

    if resp.status_code != 200:
        return _ml_page(
            'Erro na troca de token',
            f"<h1>Não foi possível gerar o token</h1><p><code>{resp.status_code}</code></p>"
            f"<p><code>{resp.text}</code></p>"
        ), resp.status_code

    dados = resp.json()
    # Fase 2: persistir dados['access_token'] / dados['refresh_token'] / dados['user_id'].
    return _ml_page(
        'Conectado',
        "<h1>Conta conectada ✅</h1><p>Token gerado com sucesso para o usuário "
        f"<code>{dados.get('user_id', '?')}</code>.</p>"
        "<p>Já dá para usar a API do Mercado Livre.</p>"
    )


# ════════════════════════════════════════
# ROTAS PRINCIPAIS
# ════════════════════════════════════════

@app.route('/')
@page_required('auditoria')
def index():
    return send_from_directory('.', 'index.html')


@app.route('/admin')
@admin_required
def admin_page():
    return send_from_directory('.', 'admin.html')


@app.route('/tarifas')
@page_required('tarifas')
def tarifas_page():
    return send_from_directory('.', 'tarifas.html')


@app.route('/reuniao')
@page_required('reuniao')
def reuniao_page():
    return send_from_directory('.', 'reuniao.html')


@app.route('/contratos')
@page_required('contratos')
def contratos_page():
    return send_from_directory('.', 'contratos.html')


@app.route('/dre')
@page_required('dre')
def dre_page():
    return send_from_directory('.', 'dre.html')


@app.route('/dre/despesas')
@page_required('despesas')
def dre_despesas_page():
    return send_from_directory('.', 'dre-despesas.html')


@app.route('/dre/conhecimentos')
@page_required('conhecimentos')
def dre_conhecimentos_page():
    return send_from_directory('.', 'dre-conhecimentos.html')


@app.route('/contabil')
@page_required('contabil')
def contabil_page():
    return send_from_directory('.', 'contabil.html')


@app.route('/contabil/extrato')
@page_required('contabil')
def contabil_extrato_page():
    return send_from_directory('.', 'contabil-extrato.html')


@app.route('/contabil/faturas')
@page_required('contabil')
def contabil_faturas_page():
    return send_from_directory('.', 'contabil-faturas.html')


@app.route('/contabil/acni')
@page_required('contabil')
def contabil_acni_page():
    return send_from_directory('.', 'contabil-acni.html')


@app.route('/contabil/eventos')
@page_required('contabil')
def contabil_eventos_page():
    return send_from_directory('.', 'contabil-eventos.html')


@app.route('/contabil/contas-fixas')
@page_required('contabil')
def contabil_contas_fixas_page():
    return send_from_directory('.', 'contabil-contas-fixas.html')


@app.route('/contabil/de-para')
@page_required('contabil')
def contabil_depara_page():
    # A tela de status virou a de configuração, que faz as duas coisas.
    return redirect('/contabil/eventos')


@app.route('/faturamento')
@page_required('faturamento')
def faturamento_page():
    return send_from_directory('.', 'faturamento.html')


@app.route('/veiculos')
@page_required('veiculos')
def veiculos_page():
    return send_from_directory('.', 'veiculos.html')


# ── PGR — excesso de velocidade ──────────────────────────────────────
#
# Acesso duplo, de propósito: o time interno abre pelo app com o login normal
# (aba 'pgr' concedida no Admin), e o diretor abre pelo link do WhatsApp sem
# logar — se ele tiver que autenticar às 7h da manhã no celular, não abre.
# O token é POR RELATÓRIO e tem validade: vazou um link, expôs um dia. A página
# é beco sem saída, sem navegação para o resto do app.

def _pgr_dia_autorizado(dia_str):
    """(dia, erro_http). Libera por sessão (aba/admin) ou por token de leitura."""
    from datetime import date as _date
    try:
        dia = _date.fromisoformat(dia_str) if dia_str else None
    except ValueError:
        return None, ('Data inválida', 400)

    # 1) sessão com a aba concedida (ou admin)
    if 'user_id' in session and (
            session.get('role') == 'admin'
            or 'pgr' in (session.get('paginas_permitidas') or [])):
        return dia or _pgr_ultimo_dia(), None

    # 2) token de leitura
    token = request.args.get('t')
    if token:
        conn = get_db()
        try:
            cur = conn.cursor()
            dia_tok = pgr.validar_token(cur, token, dia)
            conn.commit()
            cur.close()
        finally:
            conn.close()
        if dia_tok:
            return dia_tok, None
        return None, ('Link inválido ou expirado', 403)

    return None, ('Não autorizado', 401)


def _pgr_ultimo_dia():
    """Dia mais recente já apurado (a página abre nele quando não pedem data)."""
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT MAX(dia) FROM pgr_eventos")
        r = cur.fetchone()
        cur.close()
    finally:
        conn.close()
    from datetime import datetime as _dt, timedelta as _td
    return (r and r[0]) or (_dt.utcnow() - _td(hours=3) - _td(days=1)).date()


_PGR_EXPIRADO_HTML = """<!DOCTYPE html><html lang="pt-BR"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>PGR · link expirado</title><style>
body{background:#0a0e17;color:#e2e8f0;font-family:'DM Sans',-apple-system,'Segoe UI',sans-serif;
display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0;padding:24px}
div{max-width:340px;text-align:center}
h1{font-size:1.05rem;font-weight:700;margin-bottom:10px}
p{font-size:.82rem;line-height:1.6;color:#94a3b8}
</style></head><body><div>
<h1>Este relatório expirou</h1>
<p>O link do relatório de PGR vale por alguns dias. Peça um link novo para ver
este dia, ou acesse pelo sistema.</p>
</div></body></html>"""


@app.route('/pgr')
def pgr_page():
    dia, erro = _pgr_dia_autorizado(request.args.get('data'))
    if erro:
        # Sem sessão e sem token: é usuário interno, manda logar.
        if 'user_id' not in session and not request.args.get('t'):
            return redirect('/login')
        # Com token inválido/vencido: o destinatário é o diretor às 7h da
        # manhã — 403 cru é o pior desfecho possível para esse usuário.
        if request.args.get('t'):
            return Response(_PGR_EXPIRADO_HTML, status=410, mimetype='text/html')
        return erro[0], erro[1]
    return send_from_directory('.', 'pgr.html')


def sincronizar_cadastro_pgr():
    """veiculos_045 (Power BI) → pgr_cadastro_veiculos.

    Roda aqui, e não no worker, de propósito: o módulo de rastreamento é
    Postgres puro e dar DAX a ele acoplaria dois mundos limpos, além de criar
    dependência de credencial num job de madrugada. O job de apuração só lê a
    tabela.
    """
    token = get_token()
    cad = _cadastro_veiculos(token)
    conn = get_db()
    try:
        cur = conn.cursor()
        n = 0
        for placa, v in cad.items():
            cur.execute("""
                INSERT INTO pgr_cadastro_veiculos
                    (placa_norm, proprietario, tipo, modelo, eh_rizza, atualizado_em)
                VALUES (%s,%s,%s,%s,%s,NOW())
                ON CONFLICT (placa_norm) DO UPDATE SET
                    proprietario = EXCLUDED.proprietario,
                    tipo = EXCLUDED.tipo,
                    modelo = EXCLUDED.modelo,
                    eh_rizza = EXCLUDED.eh_rizza,
                    atualizado_em = NOW()
            """, (placa, v.get('proprietario'), _norm_tipo_veiculo(v.get('tipo')),
                  v.get('modelo'), _eh_rizza(v.get('proprietario'))))
            n += 1
        conn.commit()
        cur.close()
    finally:
        conn.close()
    return n


PGR_MANIFESTOS_DIAS = int(os.getenv('PGR_MANIFESTOS_DIAS', '30'))


def _txt(v, n):
    """Corta no tamanho da coluna. O cadastro do Winthor tem razão social longa
    (passa de 160 em tomador/motorista) e derrubaria a sincronização inteira."""
    s = (str(v).strip() if v is not None else '')
    return s[:n] or None


def _centroides_municipios(cur):
    """(cidade_normalizada, uf) → (lat, lng), de municipios_ibge."""
    cur.execute("SELECT cidade_normalizada, uf, latitude, longitude FROM municipios_ibge")
    return {(r[0], r[1]): (float(r[2]), float(r[3])) for r in cur.fetchall()}


def _geo_cidade_uf(centroides, cidade_uf):
    """'Uberlândia/MG' → (lat, lng). None se não achar."""
    import geocoding      # importado localmente, como nas demais funções do arquivo
    if not cidade_uf or '/' not in str(cidade_uf):
        return None
    cid, uf = str(cidade_uf).rsplit('/', 1)
    return centroides.get((geocoding.normalizar_cidade(cid), uf.strip().upper()))


def sincronizar_manifestos_pgr(dias=None):
    """Auditoria Receita (Power BI) → pgr_manifestos, já geocodificado.

    Janela móvel: o casamento precisa de manifesto EMITIDO ANTES do excesso, e
    uma viagem longa pode ter o CTRC emitido dias antes (§8 do handoff: o
    TZC0I41 carregou em 06/08 e o excesso foi em 10/08). Por isso a janela é
    generosa, não só o dia anterior.

    Geocodifica na sincronização para que o job de apuração seja aritmética
    local pura, sem depender do Power BI de madrugada.
    """
    dias = dias or PGR_MANIFESTOS_DIAS
    from datetime import date as _date, timedelta as _td
    ini = _date.today() - _td(days=dias)

    token = get_token()
    AR = "'Auditoria Receita'"
    linhas = _dax_rows(token, (
        f"EVALUATE SELECTCOLUMNS(FILTER({AR}, "
        f"{AR}[data_ref_ctrc] >= DATE({ini.year},{ini.month},{ini.day})), "
        f"\"car\",{AR}[placa_carreta],\"cav\",{AR}[placa_cavalo],"
        f"\"dt\",{AR}[data_ref_ctrc],\"manif\",{AR}[Manifesto],"
        f"\"orig\",{AR}[cidade_uf_origem],\"dest\",{AR}[cidade_uf_destino],"
        f"\"tom\",{AR}[cliente_pagador],\"mot\",{AR}[motorista],"
        f"\"tipo\",{AR}[Tipo Operacao])"))

    conn = get_db()
    try:
        cur = conn.cursor()
        centroides = _centroides_municipios(cur)
        vistos, n = set(), 0
        for r in linhas:
            dt = str(r.get('dt') or '')[:10]
            if not dt:
                continue
            cav = _placa_mercosul(r.get('cav')) or None
            car = _placa_mercosul(r.get('car')) or None
            if not cav and not car:
                continue
            chave = (r.get('manif'), dt, cav, car)
            if chave in vistos:      # o grão da Auditoria é CTRB: várias linhas por manifesto
                continue
            vistos.add(chave)
            o = _geo_cidade_uf(centroides, r.get('orig'))
            d = _geo_cidade_uf(centroides, r.get('dest'))
            cur.execute("""
                INSERT INTO pgr_manifestos (manifesto, data_ref, placa_cavalo, placa_carreta,
                    origem, destino, origem_lat, origem_lng, destino_lat, destino_lng,
                    tomador, motorista, tipo_operacao, atualizado_em)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,NOW())
                ON CONFLICT (manifesto, data_ref, placa_cavalo, placa_carreta) DO UPDATE SET
                    origem = EXCLUDED.origem, destino = EXCLUDED.destino,
                    origem_lat = EXCLUDED.origem_lat, origem_lng = EXCLUDED.origem_lng,
                    destino_lat = EXCLUDED.destino_lat, destino_lng = EXCLUDED.destino_lng,
                    tomador = EXCLUDED.tomador, motorista = EXCLUDED.motorista,
                    tipo_operacao = EXCLUDED.tipo_operacao, atualizado_em = NOW()
            """, (_txt(r.get('manif'), 30), dt, cav, car,
                  _txt(r.get('orig'), 80), _txt(r.get('dest'), 80),
                  o[0] if o else None, o[1] if o else None,
                  d[0] if d else None, d[1] if d else None,
                  _txt(r.get('tom'), 160), _txt(r.get('mot'), 160),
                  _txt((r.get('tipo') or '').upper(), 12)))
            n += 1
        # Fora da janela não serve mais e a tabela cresceria sem teto.
        cur.execute("DELETE FROM pgr_manifestos WHERE data_ref < %s",
                    (_date.today() - _td(days=dias * 2),))
        conn.commit()
        cur.close()
    finally:
        conn.close()
    return n


@app.route('/api/pgr/sync-cadastro', methods=['POST'])
@admin_required
def api_pgr_sync_cadastro():
    try:
        n = sincronizar_cadastro_pgr()
        m = sincronizar_manifestos_pgr()
        return jsonify({'ok': True, 'veiculos': n, 'manifestos': m})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _pgr_sessao_ok():
    """Visão de período e filtros são SÓ para quem está logado.

    O token do link do WhatsApp é preso a UM dia, de propósito: se ele abrisse o
    modo período, um link vazado exporia meses de operação em vez de um dia.
    """
    return 'user_id' in session and (
        session.get('role') == 'admin'
        or 'pgr' in (session.get('paginas_permitidas') or []))


def _meses_pedidos():
    return [m.strip() for m in (request.args.get('meses') or '').split(',')
            if m.strip()]


@app.route('/api/pgr')
def api_pgr():
    meses = _meses_pedidos()
    if meses:
        if not _pgr_sessao_ok():
            return jsonify({'ok': False, 'error': 'Não autorizado'}), 401
        conn = get_db()
        try:
            cur = conn.cursor()
            dados = pgr.listar_periodo(
                cur, meses,
                cavalo=request.args.get('cavalo') or None,
                carreta=request.args.get('carreta') or None,
                motorista=request.args.get('motorista') or None)
            cur.close()
        finally:
            conn.close()
        return jsonify({'ok': True, 'modo': 'periodo', **dados})

    dia, erro = _pgr_dia_autorizado(request.args.get('data'))
    if erro:
        return jsonify({'ok': False, 'error': erro[0]}), erro[1]
    conn = get_db()
    try:
        cur = conn.cursor()
        dados = pgr.listar_dia(cur, dia)
        cur.close()
    finally:
        conn.close()
    # Mesmo no dia único devolve `dias`, para a página ter um formato só.
    return jsonify({'ok': True, 'modo': 'dia',
                    'dias': [dados], 'totais': dados['totais'],
                    'dias_apurados': 1, 'limiar': dados['limiar']})


@app.route('/api/pgr/opcoes')
def api_pgr_opcoes():
    """Meses com dia apurado + valores que existem no período selecionado."""
    if not _pgr_sessao_ok():
        return jsonify({'ok': False, 'error': 'Não autorizado'}), 401
    conn = get_db()
    try:
        cur = conn.cursor()
        meses = pgr.meses_apurados(cur)
        sel = _meses_pedidos() or [m['mes'] for m in meses[:1]]
        opcoes = pgr.opcoes_filtro(cur, sel)
        cur.close()
    finally:
        conn.close()
    return jsonify({'ok': True, 'meses': meses, **opcoes})


@app.route('/embarques')
@page_required('embarques')
def embarques_page():
    return send_from_directory('.', 'embarques.html')


@app.route('/embarques/novo')
@page_required('embarques')
def embarques_novo_page():
    return send_from_directory('.', 'embarques-novo.html')


@app.route('/embarques/relatorio')
@page_required('embarques')
def embarques_relatorio_page():
    return send_from_directory('.', 'embarques-relatorio.html')


@app.route('/embarques/<int:carga_id>/editar')
@page_required('embarques')
def embarques_editar_page(carga_id):
    # A permissão é verificada na API ao buscar a carga; aqui só serve o HTML
    return send_from_directory('.', 'embarques-novo.html')


@app.route('/embarques/mapa')
@page_required('embarques')
def embarques_mapa_page():
    return send_from_directory('.', 'mapa.html')


@app.route('/embarques/cargas/<int:carga_id>/mapa')
@page_required('embarques')
def embarques_mapa_carga_page(carga_id):
    return send_from_directory('.', 'mapa-carga.html')


@app.route('/api/tarifas')
@login_required
def tarifas():
    try:
        token = get_token()
        result = execute_dax(token, "EVALUATE 'public tarifas_frete'")
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)

        # Mantém apenas a última versão (maior versao_id) por cliente
        max_versao = {}
        for r in data:
            cliente = r.get('cliente_nome')
            v = r.get('versao_id')
            if cliente and v is not None:
                if cliente not in max_versao or v > max_versao[cliente]:
                    max_versao[cliente] = v

        data = [r for r in data if r.get('versao_id') == max_versao.get(r.get('cliente_nome'))]

        return jsonify({'ok': True, 'data': data, 'count': len(data)})

    except requests.exceptions.HTTPError as e:
        detail = ''
        try:
            detail = e.response.json()
        except Exception:
            detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/icms')
@login_required
def api_icms():
    """Consulta a matriz de ICMS de transporte por UF (icms_aliquota).
    GET /api/icms?origem=XX&destino=YY -> {aliquota, tipo, isento, observacao}."""
    origem  = (request.args.get('origem')  or '').upper().strip()
    destino = (request.args.get('destino') or '').upper().strip()
    if len(origem) != 2 or len(destino) != 2:
        return jsonify({'ok': False, 'error': 'Informe origem e destino (UF de 2 letras)'}), 400
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute(
            "SELECT aliquota, tipo, isento, observacao FROM icms_aliquota WHERE uf_origem=%s AND uf_destino=%s",
            (origem, destino)
        )
        r = cur.fetchone()
        cur.close(); conn.close()
        if not r:
            return jsonify({'ok': False, 'error': f'Par {origem}->{destino} não encontrado'}), 404
        return jsonify({
            'ok': True,
            'origem': origem, 'destino': destino,
            'aliquota': float(r[0]) if r[0] is not None else None,
            'tipo': r[1],
            'isento': bool(r[2]),
            'observacao': r[3],
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/me')
@login_required
def me():
    return jsonify({
        'ok':              True,
        'nome':            session.get('nome'),
        'role':            session.get('role'),
        'tipos_permitidos': session.get('tipos_permitidos', []),
        'paginas_permitidas': session.get('paginas_permitidas', []),
    })


@app.route('/api/status')
@login_required
def status():
    missing = [k for k, v in CONFIG.items() if not v]
    if missing:
        return jsonify({'ok': False, 'missing': missing}), 400
    return jsonify({'ok': True})


def _ancoragem_vazio(v):
    return v is None or (isinstance(v, str) and v.strip() == '')


def _emissao_por_ctrb(token):
    """Mapa {ctrb: emissao} de 'public ctrbs_oss' (data da OS), cacheado 5min
    (_cache_get/_cache_set). Usado só p/ o fallback de data dos CTRBs órfãos."""
    cached = _cache_get('ctrbs_emissao')
    if cached is not None:
        return cached
    res = execute_dax(token, "EVALUATE SELECTCOLUMNS('public ctrbs_oss', "
                      "\"ctrb\",'public ctrbs_oss'[ctrb], \"emissao\",'public ctrbs_oss'[emissao])")
    linhas = clean_rows(res.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))
    mapa = {str(r.get('ctrb')): r.get('emissao') for r in linhas if r.get('ctrb')}
    _cache_set('ctrbs_emissao', mapa)
    return mapa


def _anexar_ancoragem(token, data):
    """Acrescenta (ADITIVO — não altera nenhum campo existente) a cada linha:
      - 'data_efetiva' = data_ref_ctrc, com fallback p/ emissao da OS (só difere nos sem-data);
      - 'ancoragem'    = 'orfao' (sem Manifesto) | 'sem_nota' (com Manifesto, sem data fiscal) | 'ok'.
    Degrada com segurança: se a busca de emissao falhar, data_efetiva = data_ref_ctrc (= hoje)."""
    try:
        emap = _emissao_por_ctrb(token)
    except Exception:
        emap = {}
    for r in data:
        data_ref = r.get('data_ref_ctrc')
        emissao = emap.get(str(r.get('CTRB')))
        r['data_efetiva'] = data_ref if not _ancoragem_vazio(data_ref) else emissao
        if _ancoragem_vazio(r.get('Manifesto')):
            r['ancoragem'] = 'orfao'
        elif _ancoragem_vazio(data_ref):
            r['ancoragem'] = 'sem_nota'
        else:
            r['ancoragem'] = 'ok'
    return data


@app.route('/api/auditoria')
@login_required
def auditoria():
    try:
        token = get_token()
        result = execute_dax(token, "EVALUATE 'Auditoria Receita'")
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)

        # Filtrar por tipos permitidos
        tipos = session.get('tipos_permitidos', [])
        if tipos:
            tipos_lower = [t.lower() for t in tipos]
            data = [
                r for r in data
                if any(t in (r.get('Tipo Operacao') or '').lower() for t in tipos_lower)
            ]

        # Aditivo: data_efetiva (fallback p/ órfãos) + rótulo de ancoragem.
        data = _anexar_ancoragem(token, data)

        return jsonify({'ok': True, 'data': data, 'count': len(data)})

    except requests.exceptions.HTTPError as e:
        detail = ''
        try:
            detail = e.response.json()
        except Exception:
            detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/dax', methods=['POST'])
@login_required
def dax_query():
    try:
        body = request.get_json()
        query = body.get('query', '')
        if not query:
            return jsonify({'ok': False, 'error': 'Query vazia'}), 400

        token = get_token()
        result = execute_dax(token, query)
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)
        return jsonify({'ok': True, 'data': data, 'count': len(data)})

    except requests.exceptions.HTTPError as e:
        detail = ''
        try:
            detail = e.response.json()
        except Exception:
            detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ════════════════════════════════════════
# ROTAS ADMIN — GERENCIAMENTO DE USUÁRIOS
# ════════════════════════════════════════

@app.route('/api/admin/users', methods=['GET'])
@admin_required
def admin_list_users():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT id, nome, email, role, ativo, tipos_permitidos, criado_em, paginas_permitidas FROM auditoria_users ORDER BY criado_em")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        users = [
            {'id': r[0], 'nome': r[1], 'email': r[2], 'role': r[3], 'ativo': r[4],
             'tipos_permitidos': r[5] or [],
             'criado_em': r[6].strftime('%d/%m/%Y %H:%M') if r[6] else '',
             'paginas_permitidas': r[7] or []}
            for r in rows
        ]
        return jsonify({'ok': True, 'users': users})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


TIPOS_VALIDOS = {'Carreteiro', 'Agregado', 'Frota'}
PAGINAS_PADRAO = ['auditoria', 'tarifas', 'embarques']  # abas liberadas por padrão

@app.route('/api/admin/users', methods=['POST'])
@admin_required
def admin_create_user():
    data = request.get_json() or {}
    nome             = data.get('nome', '').strip()
    email            = data.get('email', '').strip().lower()
    senha            = data.get('senha', '')
    role             = data.get('role', 'viewer')
    tipos_permitidos = data.get('tipos_permitidos', list(TIPOS_VALIDOS))
    paginas_permitidas = data.get('paginas_permitidas', PAGINAS_PADRAO)

    if not nome or not email or not senha:
        return jsonify({'ok': False, 'error': 'Nome, e-mail e senha são obrigatórios'}), 400
    if role not in ('admin', 'viewer'):
        return jsonify({'ok': False, 'error': 'Role inválido'}), 400
    tipos_permitidos = [t for t in tipos_permitidos if t in TIPOS_VALIDOS]
    if not tipos_permitidos:
        return jsonify({'ok': False, 'error': 'Selecione ao menos um tipo de operação'}), 400
    paginas_permitidas = [p for p in paginas_permitidas if p in PAGINAS_VALIDAS]

    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            """INSERT INTO auditoria_users (nome, email, password_hash, role, tipos_permitidos, paginas_permitidas)
               VALUES (%s, %s, %s, %s, %s, %s) RETURNING id""",
            (nome, email, generate_password_hash(senha), role, tipos_permitidos, paginas_permitidas)
        )
        new_id = cur.fetchone()[0]
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'ok': True, 'id': new_id})
    except psycopg2.errors.UniqueViolation:
        return jsonify({'ok': False, 'error': 'E-mail já cadastrado'}), 409
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/admin/users/<int:uid>', methods=['PATCH'])
@admin_required
def admin_toggle_user(uid):
    data = request.get_json() or {}

    # Atualizar perfil (role)
    if 'role' in data:
        novo_role = data['role']
        if novo_role not in ('admin', 'viewer'):
            return jsonify({'ok': False, 'error': 'Perfil inválido'}), 400
        if uid == session.get('user_id'):
            return jsonify({'ok': False, 'error': 'Você não pode mudar o próprio perfil'}), 400
        try:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("UPDATE auditoria_users SET role = %s WHERE id = %s", (novo_role, uid))
            conn.commit()
            cur.close()
            conn.close()
            return jsonify({'ok': True})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    # Atualizar paginas_permitidas (abas visíveis)
    if 'paginas_permitidas' in data:
        paginas = [p for p in data['paginas_permitidas'] if p in PAGINAS_VALIDAS]
        try:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("UPDATE auditoria_users SET paginas_permitidas = %s WHERE id = %s", (paginas, uid))
            conn.commit()
            cur.close()
            conn.close()
            return jsonify({'ok': True})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    # Atualizar tipos_permitidos
    if 'tipos_permitidos' in data:
        tipos = [t for t in data['tipos_permitidos'] if t in TIPOS_VALIDOS]
        if not tipos:
            return jsonify({'ok': False, 'error': 'Selecione ao menos um tipo'}), 400
        try:
            conn = get_db()
            cur = conn.cursor()
            cur.execute("UPDATE auditoria_users SET tipos_permitidos = %s WHERE id = %s", (tipos, uid))
            conn.commit()
            cur.close()
            conn.close()
            return jsonify({'ok': True})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    # Ativar/Desativar
    ativo = data.get('ativo')
    if ativo is None:
        return jsonify({'ok': False, 'error': 'Campo "ativo" ou "tipos_permitidos" obrigatório'}), 400
    if uid == session.get('user_id') and not ativo:
        return jsonify({'ok': False, 'error': 'Você não pode desativar sua própria conta'}), 400
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE auditoria_users SET ativo = %s WHERE id = %s", (ativo, uid))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/admin/users/<int:uid>', methods=['DELETE'])
@admin_required
def admin_delete_user(uid):
    if uid == session.get('user_id'):
        return jsonify({'ok': False, 'error': 'Você não pode excluir sua própria conta'}), 400
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM auditoria_users WHERE id = %s", (uid,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ════════════════════════════════════════
# ROTAS REUNIÃO — TRANSCRIÇÃO + ATA
# ════════════════════════════════════════

def _transcrever_com_assemblyai(caminho_arquivo, speakers_expected=None):
    import assemblyai as aai
    aai.settings.api_key = os.getenv('ASSEMBLYAI_API_KEY')
    config_args = dict(
        speech_models=['universal-3-pro', 'universal-2'],
        language_detection=True,
        speaker_labels=True,
    )
    if speakers_expected:
        config_args['speakers_expected'] = int(speakers_expected)
    config = aai.TranscriptionConfig(**config_args)
    transcriber = aai.Transcriber(config=config)
    transcript = transcriber.transcribe(caminho_arquivo)
    if transcript.status == aai.TranscriptStatus.error:
        raise Exception(f'AssemblyAI erro: {transcript.error}')
    if transcript.utterances:
        linhas = [f'Speaker {u.speaker}: {u.text}' for u in transcript.utterances]
        return '\n'.join(linhas)
    return transcript.text


def _gerar_ata(tema, transcricao):
    from openai import OpenAI
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

    prompt_geracao = f"""Você é um assistente especializado em redigir atas de reunião para empresas de transporte rodoviário de cargas.

Contexto da empresa: transportadora com operações em múltiplos estados (SP, RJ, GO, ES, BA e outros), frota própria de carretas, motoristas agregados e carreteiros, clientes industriais (L'Oreal, Nestlé, Heinz, etc.), atuação com cargas spot e contratos fixos.

Tema da reunião: {tema}

Transcrição:
{transcricao}

Instruções obrigatórias:
- Speakers identificados (Speaker A, Speaker B, etc.) = participantes reais da reunião. NENHUM outro nome deve aparecer como participante
- Tente inferir o nome real de cada Speaker pelo contexto (ex: se alguém chama "João" e Speaker A responde, provavelmente é João). Se não for possível inferir, use "Speaker A"
- Nomes citados na conversa (motoristas, coordenadores, clientes) aparecem APENAS no corpo do texto, nunca como participantes
- Nomes como "Martins" podem ser terminais/clientes — use o contexto para distinguir pessoas de empresas/localidades
- Prazos: use apenas datas mencionadas explicitamente na transcrição. Se não houver data, escreva "A definir". NUNCA invente datas
- Termos do setor são válidos: carreta, frota, agregado, carreteiro, spot, frete líquido, recuperação judicial (RJ), diária, escala, etc.
- Linguagem: profissional e objetiva, sem excesso de formalidade

Estrutura obrigatória da ata (nesta ordem):
1. **Cabeçalho** — data/hora se mencionada, senão deixar em branco; tema; local se mencionado
2. **Participantes** — apenas os Speakers com nome inferido ou label (ex: "João (Speaker A)")
3. **Resumo Executivo** — 3 a 5 decisões/pontos principais em bullets, para leitura rápida
4. **Encaminhamentos** — tabela com: Encaminhamento | Responsável | Prazo
5. **Pauta abordada** — tópicos discutidos
6. **Discussões e Deliberações** — detalhamento por tópico
7. **Próximos Passos** — se houver"""

    primeira_ata = client.chat.completions.create(
        model='gpt-4.1-mini',
        messages=[{'role': 'user', 'content': prompt_geracao}],
        max_tokens=4000
    ).choices[0].message.content

    prompt_revisao = f"""Revise a ata de reunião de uma transportadora abaixo. Corrija especificamente:

1. **Participantes incorretos** — remova qualquer nome que não seja um Speaker identificado na transcrição original
2. **Datas inventadas** — substitua por "A definir" qualquer prazo que não foi mencionado explicitamente na reunião
3. **Pessoas vs empresas/terminais** — confirme que "Martins", "Raiz", "Start" e similares estão como clientes/terminais, não como pessoas
4. **Encaminhamentos sem responsável real** — se o responsável for desconhecido, use o Speaker mais provável pelo contexto
5. **Repetições e redundâncias** entre seções
6. **Ordem da estrutura** — garanta que Resumo Executivo e Encaminhamentos vêm ANTES das discussões detalhadas

Retorne apenas a ata final revisada, sem comentários.

Ata a revisar:
{primeira_ata}"""

    ata_final = client.chat.completions.create(
        model='gpt-4.1-mini',
        messages=[{'role': 'user', 'content': prompt_revisao}],
        max_tokens=4000
    ).choices[0].message.content

    return ata_final


@app.route('/api/reuniao/processar', methods=['POST'])
@page_required('reuniao')
def processar_reuniao():
    if 'audio' not in request.files:
        return jsonify({'ok': False, 'error': 'Arquivo de áudio não enviado'}), 400

    audio_file = request.files['audio']
    tema = request.form.get('tema', 'Reunião').strip()
    participantes = request.form.get('participantes')

    ext = os.path.splitext(audio_file.filename)[1].lower() or '.mp3'
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
        audio_file.save(tmp.name)
        tmp_path = tmp.name

    try:
        transcricao = _transcrever_com_assemblyai(tmp_path, speakers_expected=participantes)
        ata = _gerar_ata(tema, transcricao)
        return jsonify({'ok': True, 'ata': ata, 'transcricao': transcricao})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.route('/api/reuniao/exportar', methods=['POST'])
@page_required('reuniao')
def exportar_reuniao():
    data = request.get_json() or {}
    ata = data.get('ata', '')
    tema = data.get('tema', 'Ata de Reunião')
    formato = data.get('formato', 'docx').lower()

    if not ata:
        return jsonify({'ok': False, 'error': 'Conteúdo da ata não informado'}), 400

    def _parse_md_linha(texto):
        import re
        texto = re.sub(r'\*\*(.+?)\*\*', r'\1', texto)
        texto = re.sub(r'\*(.+?)\*', r'\1', texto)
        return texto.strip()

    def _e_separador_tabela(linha):
        return all(c in '|- :' for c in linha) and '|' in linha

    def _extrair_tabela(linhas, idx):
        rows = []
        while idx < len(linhas) and '|' in linhas[idx]:
            celulas = [c.strip() for c in linhas[idx].strip().strip('|').split('|')]
            if not _e_separador_tabela(linhas[idx]):
                rows.append(celulas)
            idx += 1
        return rows, idx

    if formato == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_heading(tema, level=1)
        linhas = ata.split('\n')
        i = 0
        while i < len(linhas):
            linha = linhas[i].strip()
            if not linha or linha == '---':
                i += 1
                continue
            if linha.startswith('### '):
                doc.add_heading(_parse_md_linha(linha[4:]), level=3)
            elif linha.startswith('## '):
                doc.add_heading(_parse_md_linha(linha[3:]), level=2)
            elif linha.startswith('# '):
                doc.add_heading(_parse_md_linha(linha[2:]), level=1)
            elif '|' in linha and not _e_separador_tabela(linha):
                rows, i = _extrair_tabela(linhas, i)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            cell = table.cell(r_idx, c_idx)
                            cell.text = _parse_md_linha(cell_text)
                            if r_idx == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                continue
            elif linha.startswith('- ') or linha.startswith('* '):
                doc.add_paragraph(_parse_md_linha(linha[2:]), style='List Bullet')
            else:
                p = doc.add_paragraph()
                partes = linha.split('**')
                for j, parte in enumerate(partes):
                    run = p.add_run(parte)
                    run.bold = (j % 2 == 1)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        nome_arquivo = f"{tema.replace(' ', '_')}.docx"
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=nome_arquivo)

    elif formato == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        s_normal  = ParagraphStyle('n', parent=styles['Normal'], fontSize=10, leading=15, spaceAfter=3)
        s_h1      = ParagraphStyle('h1', parent=styles['Heading1'], fontSize=15, spaceAfter=8)
        s_h2      = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=12, spaceAfter=6, spaceBefore=10)
        s_h3      = ParagraphStyle('h3', parent=styles['Heading3'], fontSize=11, spaceAfter=4, spaceBefore=8)
        s_bullet  = ParagraphStyle('b', parent=s_normal, leftIndent=14, bulletIndent=4)
        s_title   = ParagraphStyle('t', parent=styles['Title'], fontSize=16, spaceAfter=12)
        s_cell    = ParagraphStyle('c', parent=s_normal, fontSize=9, leading=13)
        s_cell_h  = ParagraphStyle('ch', parent=s_cell, fontName='Helvetica-Bold')

        story = [Paragraph(tema, s_title), Spacer(1, 8)]
        linhas = ata.split('\n')
        i = 0
        while i < len(linhas):
            linha = linhas[i].strip()
            if not linha or linha == '---':
                story.append(Spacer(1, 6))
                i += 1
                continue
            if linha.startswith('### '):
                story.append(Paragraph(_parse_md_linha(linha[4:]), s_h3))
            elif linha.startswith('## '):
                story.append(Paragraph(_parse_md_linha(linha[3:]), s_h2))
            elif linha.startswith('# '):
                story.append(Paragraph(_parse_md_linha(linha[2:]), s_h1))
            elif '|' in linha and not _e_separador_tabela(linha):
                rows, i = _extrair_tabela(linhas, i)
                if rows:
                    col_w = (A4[0] - 5*cm) / max(len(rows[0]), 1)
                    data = []
                    for r_idx, row in enumerate(rows):
                        estilo = s_cell_h if r_idx == 0 else s_cell
                        data.append([Paragraph(_parse_md_linha(c), estilo) for c in row])
                    t = Table(data, colWidths=[col_w]*len(rows[0]))
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a2235')),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#94a3b8')),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#1e293b')),
                        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('PADDING', (0,0), (-1,-1), 6),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 8))
                continue
            elif linha.startswith('- ') or linha.startswith('* '):
                story.append(Paragraph(f'• {_parse_md_linha(linha[2:])}', s_bullet))
            else:
                import re
                texto_html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', linha)
                story.append(Paragraph(texto_html, s_normal))
            i += 1

        doc.build(story)
        buf.seek(0)
        nome_arquivo = f"{tema.replace(' ', '_')}.pdf"
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name=nome_arquivo)

    return jsonify({'ok': False, 'error': 'Formato inválido. Use docx ou pdf'}), 400


# ════════════════════════════════════════
# CONTRATOS — Emissão de contrato TAC Agregado
# ════════════════════════════════════════

@app.route('/api/contratos/extrair', methods=['POST'])
@page_required('contratos')
def contratos_extrair():
    import contratos_service as cs
    arquivos = []
    for f in request.files.getlist('documentos'):
        if f and f.filename:
            arquivos.append((f.filename, f.read()))
    if not arquivos:
        return jsonify({'ok': False, 'error': 'Envie ao menos um documento.'}), 400

    # Dados já editados pelo operador (para mesclar numa reextração com mais documentos).
    dados_existentes = None
    if request.form.get('dados'):
        try:
            dados_existentes = json.loads(request.form['dados'])
        except (ValueError, TypeError):
            dados_existentes = None

    try:
        dados = cs.extrair_documentos(arquivos)
        if dados_existentes:
            dados = cs.merge_dados(dados_existentes, dados)
        pendencias = cs.checar_pendencias(dados)
        return jsonify({'ok': True, 'dados': dados, 'pendencias': pendencias})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _contrato_contexto_do_request(data):
    """Monta o contexto do template a partir do payload JSON do frontend."""
    import contratos_service as cs
    dados = data.get('dados') or {}
    contexto = cs.montar_contexto(
        dados,
        usa_rastreador_proprio=bool(data.get('usa_rastreador_proprio')),
        vigencia_inicio=data.get('vigencia_inicio', ''),
        vigencia_termino=data.get('vigencia_termino', ''),
        comodato_numero_serie=data.get('comodato_numero_serie', ''),
        comodato_estado=data.get('comodato_estado', ''),
        comodato_marca_modelo=data.get('comodato_marca_modelo', ''),
    )
    return dados, contexto


@app.route('/api/contratos/gerar', methods=['POST'])
@page_required('contratos')
def contratos_gerar():
    import contratos_service as cs
    data = request.get_json() or {}
    dados = data.get('dados') or {}
    pendencias = cs.checar_pendencias(dados)
    if not data.get('usa_rastreador_proprio'):
        if not (data.get('comodato_marca_modelo') or '').strip():
            pendencias.append('Informe a marca/modelo do rastreador (Comodato).')
        if not (data.get('comodato_numero_serie') or '').strip():
            pendencias.append('Informe o nº de série/ID do rastreador (Comodato).')
    if pendencias:
        return jsonify({'ok': False, 'error': 'Existem pendências impeditivas.',
                        'pendencias': pendencias}), 400
    try:
        _, contexto = _contrato_contexto_do_request(data)
        docx_bytes = cs.gerar_docx(contexto)
        buf = io.BytesIO(docx_bytes)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True, download_name=cs.nome_arquivo(contexto) + '.docx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contratos/preview', methods=['POST'])
@page_required('contratos')
def contratos_preview():
    """Devolve o contrato preenchido em HTML para visualização e 'Salvar como PDF'
    pelo navegador (derivado do mesmo .docx gerado — fonte única)."""
    import contratos_service as cs
    data = request.get_json() or {}
    dados = data.get('dados') or {}
    pendencias = cs.checar_pendencias(dados)
    if not data.get('usa_rastreador_proprio'):
        if not (data.get('comodato_marca_modelo') or '').strip():
            pendencias.append('Informe a marca/modelo do rastreador (Comodato).')
        if not (data.get('comodato_numero_serie') or '').strip():
            pendencias.append('Informe o nº de série/ID do rastreador (Comodato).')
    if pendencias:
        return jsonify({'ok': False, 'error': 'Existem pendências impeditivas.',
                        'pendencias': pendencias}), 400
    try:
        _, contexto = _contrato_contexto_do_request(data)
        docx_bytes = cs.gerar_docx(contexto)
        return jsonify({'ok': True, 'html': cs.gerar_html(docx_bytes),
                        'titulo': cs.nome_arquivo(contexto)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ════════════════════════════════════════
# DRE — Demonstrativo do Resultado do Exercício
# ════════════════════════════════════════

# Mapeamento de cada descr_evento para (Grupo, Subgrupo) — traduzido do DAX Mapa_DRE
MAPA_DRE = {
    # OPERACIONAL — FRETES
    'FRETE COLETA/ENTREGA DE VEICULOS TERCEIROS': ('Operacional', 'Fretes'),
    'FRETE TRANSFERENCIA C/ AGREGADOS': ('Operacional', 'Fretes'),
    'FRETE TRANSFERENCIA C/ TERCEIROS': ('Operacional', 'Fretes'),
    'FRETE FLUVIAL': ('Operacional', 'Fretes'),
    # COMBUSTÍVEL
    'COMBUSTIVEIS E LUBRIFICANTES': ('Operacional', 'Combustível'),
    # MANUTENÇÃO
    'MANUTENCAO E CONSERVACAO DE VEICULOS': ('Operacional', 'Manutenção'),
    'SERVICO MANUTENCAO CAVALOS': ('Operacional', 'Manutenção'),
    'SERVICOS MANUTENCAO CARRETA': ('Operacional', 'Manutenção'),
    'PECAS MANUTENCAO CAVALO': ('Operacional', 'Manutenção'),
    'PECAS MANUTENCAO CARRETA': ('Operacional', 'Manutenção'),
    'PECAS E ACESSORIOS': ('Operacional', 'Manutenção'),
    # PNEUS
    'PNEUS E CAMARAS': ('Operacional', 'Pneus'),
    'RECAPAGEM DE PNEUS': ('Operacional', 'Pneus'),
    # DESLOCAMENTO
    'PEDAGIOS': ('Operacional', 'Deslocamento'),
    'ESTACIONAMENTOS': ('Operacional', 'Deslocamento'),
    'DIARIAS': ('Operacional', 'Deslocamento'),
    # SEGUROS
    'SEGURO DE CARGAS': ('Operacional', 'Seguros'),
    'SEGURO DE VEICULOS': ('Operacional', 'Seguros'),
    'GERENCIAMENTO DE RISCO': ('Operacional', 'Seguros'),
    # MÃO DE OBRA OPERACIONAL
    'SALARIO MENSAL - OPERACIONAL': ('Operacional', 'Mão de Obra'),
    'SALARIO MENSAL - APOIO OPERACIONAL': ('Operacional', 'Mão de Obra'),
    'ADIANTAMENTO SALARIAL OPERACIONAL': ('Operacional', 'Mão de Obra'),
    'ADIANTAMENTO SALARIAL - APOIO OPERACIONAL': ('Operacional', 'Mão de Obra'),
    'SALARIOS': ('Operacional', 'Mão de Obra'),
    'ADIANTAMENTO SALARIOS': ('Operacional', 'Mão de Obra'),
    'RPA- RECIBO DE PAGAMENTO AUTONOMO': ('Operacional', 'Mão de Obra'),
    'ADIANTAMENTO SALARIAL - ADMINISTRATIVO/ APOIO': ('Operacional', 'Mão de Obra'),
    # OUTROS OPERACIONAIS
    'CARGA E DESCARGA C/ TERCEIROS': ('Operacional', 'Outros'),
    'INDENIZACAO DE MERCADORIAS - ONUS DE CONTRATO': ('Operacional', 'Outros'),
    'LOCACAO DE CARRETA': ('Operacional', 'Outros'),
    'DESPESAS PJ - OPERACIONAL': ('Operacional', 'Outros'),
    'OUTRAS DESPESAS OPERACIONAIS': ('Operacional', 'Outros'),
    'ACERTO CONTA FORNECEDOR': ('Operacional', 'Outros'),
    # ADMINISTRATIVO — MÃO DE OBRA
    'SALARIOS ADMINISTRATIVOS - APOIO': ('Administrativo', 'Mão de Obra'),
    'PRO-LABORE': ('Administrativo', 'Mão de Obra'),
    'FERIAS': ('Administrativo', 'Mão de Obra'),
    'RESCISOES': ('Administrativo', 'Mão de Obra'),
    '13O SALARIOS': ('Administrativo', 'Mão de Obra'),
    'PENSAO ALIMENTICIA': ('Administrativo', 'Mão de Obra'),
    # ENCARGOS
    'INSS': ('Administrativo', 'Encargos'),
    'FGTS': ('Administrativo', 'Encargos'),
    'IRRF- AUTONOMOS - CLT': ('Administrativo', 'Encargos'),
    # ESTRUTURA
    'ALUGUEL DO IMOVEL': ('Administrativo', 'Estrutura'),
    'ENERGIA ELETRICA': ('Administrativo', 'Estrutura'),
    'AGUA E ESGOTO': ('Administrativo', 'Estrutura'),
    'MANUTENCAO E CONSERVACAO DO IMOVEL': ('Administrativo', 'Estrutura'),
    'DESPESAS COM HIGIENE E LIMPEZA': ('Administrativo', 'Estrutura'),
    # SISTEMAS
    'SOFTWARE E LICENCAS': ('Administrativo', 'Sistemas'),
    'TELEFONIA E COMUNICACAO DE DADOS': ('Administrativo', 'Sistemas'),
    'DESPESAS CONTABEIS': ('Administrativo', 'Sistemas'),
    'DESPESAS PJ - APOIO ADMINISTRATIVO E COMERCIAL': ('Administrativo', 'Sistemas'),
    'OUTROS SERVICOS PJ': ('Administrativo', 'Sistemas'),
    # JURÍDICO
    'DESPESAS JURIDICAS - INDENIZAACAOES TRABALHISTAS': ('Administrativo', 'Jurídico'),
    'DESPESAS SINDICAIS': ('Administrativo', 'Jurídico'),
    # SAÚDE E SEGURANÇA
    'PLANO DE SAUDE': ('Administrativo', 'Saúde'),
    'SAUDE OCUPACIONAL - EXAMES LTCAT PPRA PCMSO': ('Administrativo', 'Saúde'),
    'SEGURO DE VIDA': ('Administrativo', 'Saúde'),
    'SEGURANCA DO TRABALHO - BRIGADA - EPIS - UNIFORME': ('Administrativo', 'Segurança'),
    'SEGURANCA E VIGILANCIA PATRIMONIAL': ('Administrativo', 'Segurança'),
    'SEGURO DE IMOVEIS': ('Administrativo', 'Segurança'),
    # TAXAS
    'LICENCAS- SUATRANS - ANVISA- BOMBEIRO': ('Administrativo', 'Taxas'),
    'TAXAS PREFEITURA - TAXA INCENDIO RENOVACAO': ('Administrativo', 'Taxas'),
    'IPVA': ('Administrativo', 'Taxas'),
    'IPTU': ('Administrativo', 'Taxas'),
    'MULTAS E INFRACOES JUNTO AOS ORGAOS FEDERAIS MUNIC': ('Administrativo', 'Taxas'),
    'INFRACOES DE TRANSITO': ('Administrativo', 'Taxas'),
    'LICENCIAMENTO DE VEICULOS': ('Administrativo', 'Taxas'),
    # COMERCIAL
    'COMISSAO AGENTE': ('Administrativo', 'Comercial'),
    'DESPESA COMERCIAL - VIAGENS E RELATORIO DESPESA': ('Administrativo', 'Comercial'),
    'CORREIOS E TELEGRAFOS': ('Administrativo', 'Comercial'),
    'BRINDES DOACOES CONFRATERNIZACOES': ('Administrativo', 'Comercial'),
    # BENEFÍCIOS
    'VALE REFEICAO': ('Administrativo', 'Benefícios'),
    'VALE TRANSPORTE': ('Administrativo', 'Benefícios'),
    'ALIMENTACAO': ('Administrativo', 'Benefícios'),
    # OUTROS ADM
    'DESPACHANTE': ('Administrativo', 'Outros'),
    'MATERIAL DE ESCRITORIO': ('Administrativo', 'Outros'),
    'TARIFA - PEF CTRB': ('Administrativo', 'Outros'),
    'TREINAMENTO DESENVOLVIMENTO BENEFICIOS': ('Administrativo', 'Outros'),
    # FINANCEIRO
    'EMPRESTIMOS': ('Financeiro', 'Dívida'),
    'CAPITAL DE GIRO': ('Financeiro', 'Dívida'),
    'JUROS E ENCARGOS': ('Financeiro', 'Custos Financeiros'),
    'IOF': ('Financeiro', 'Custos Financeiros'),
    'TARIFA BANCARIA': ('Financeiro', 'Custos Financeiros'),
    # IMPOSTOS
    'IR': ('Impostos', 'Impostos'),
    'CSLL': ('Impostos', 'Impostos'),
    # DEDUÇÕES
    'ICMS': ('Deduções', 'Deduções'),
    'PIS': ('Deduções', 'Deduções'),
    'COFINS': ('Deduções', 'Deduções'),
    'ISS': ('Deduções', 'Deduções'),
    # INVESTIMENTOS
    'ATIVO IMOBILIZADO - IMOVEIS': ('Investimento', 'Investimentos'),
    'ATIVO IMOBILIZADO- VEICULOS': ('Investimento', 'Investimentos'),
    'INVESTIMENTO - CONSORCIO': ('Investimento', 'Investimentos'),
    'INVESTIMENTO- CDC': ('Investimento', 'Investimentos'),
    'INVESTIMENTO- FINAME': ('Investimento', 'Investimentos'),
    # RETIRADAS
    'RETIRADA CLEIVON': ('Retirada', 'Retirada'),
    'RETIRADA SOCIOS': ('Retirada', 'Retirada'),
    'RETIRADA PATRICIA': ('Retirada', 'Retirada'),
    'MANUTENCAO DE MAQUINAS E EQUIPAMENTOS': ('Retirada', 'Retirada'),
}

# Filtragem das tabelas (via Power BI DAX, igual /api/tarifas e /api/auditoria):
# - despesas: filtra pela coluna REF (formato 'YYYY/MM') — competência
# - conhecimentos: filtra por data_autorizacao


def _dre_zerado(receita_bruta=0.0):
    return {
        'receita_bruta': receita_bruta, 'deducoes': 0.0, 'receita_liquida': receita_bruta,
        'custo_operacional': 0.0, 'despesas_administrativas': 0.0, 'ebitda': receita_bruta,
        'despesas_financeiras': 0.0, 'lair': receita_bruta, 'impostos': 0.0,
        'lucro_liquido': receita_bruta, 'investimentos': 0.0, 'pos_investimento': receita_bruta,
        'retiradas': 0.0, 'resultado_final': receita_bruta,
    }


def _gerar_refs_periodo(start_date, end_date):
    """Gera lista de strings YYYY/MM cobrindo o período."""
    refs = []
    cur_y, cur_m = start_date.year, start_date.month
    end_y, end_m = end_date.year, end_date.month
    while (cur_y, cur_m) <= (end_y, end_m):
        refs.append(f"{cur_y:04d}/{cur_m:02d}")
        cur_m += 1
        if cur_m > 12:
            cur_m = 1
            cur_y += 1
    return refs

DRE_LINHAS = [
    ('Receita Bruta',                'Subtotal', 'receita_bruta'),
    ('(-) Deduções',                 'Grupo',    'deducoes'),
    ('= Receita Líquida',            'Subtotal', 'receita_liquida'),
    ('(-) Custo Operacional',        'Grupo',    'custo_operacional'),
    ('(-) Despesas Administrativas', 'Grupo',    'despesas_administrativas'),
    ('= EBITDA',                     'Subtotal', 'ebitda'),
    ('(-) Despesas Financeiras',     'Grupo',    'despesas_financeiras'),
    ('= LAIR',                       'Subtotal', 'lair'),
    ('(-) Impostos',                 'Grupo',    'impostos'),
    ('= Lucro Líquido',              'Subtotal', 'lucro_liquido'),
    ('(-) Investimentos',            'Grupo',    'investimentos'),
    ('= Pós Investimento',           'Subtotal', 'pos_investimento'),
    ('(-) Retiradas',                'Grupo',    'retiradas'),
    ('= Resultado Final',            'Subtotal', 'resultado_final'),
]

# Mapeamento: linha da DRE → Grupo do MAPA_DRE (para drilldown)
DRE_LINHA_GRUPO = {
    '(-) Deduções':                 'Deduções',
    '(-) Custo Operacional':        'Operacional',
    '(-) Despesas Administrativas': 'Administrativo',
    '(-) Despesas Financeiras':     'Financeiro',
    '(-) Impostos':                 'Impostos',
    '(-) Investimentos':            'Investimento',
    '(-) Retiradas':                'Retirada',
}


def _dax_lista_refs(refs):
    """Formata lista Python ['2026/01','2026/02'] como literal DAX: { "2026/01", "2026/02" }"""
    return '{ ' + ', '.join(f'"{r}"' for r in refs) + ' }'


def _dax_data(d):
    """Formata date como literal DAX: DATE(2026,3,1)"""
    return f'DATE({d.year},{d.month},{d.day})'


def _calcular_dre_periodo(start_date, end_date):
    """Calcula uma DRE para um período fechado via Power BI DAX."""
    token = get_token()

    # Receita Bruta = SUM(valor_frete) filtrado por data_autorizacao
    dax_receita = (
        f'EVALUATE ROW("total", '
        f'CALCULATE(SUM(\'public conhecimentos_emitidos\'[valor_frete]), '
        f'FILTER(\'public conhecimentos_emitidos\', '
        f'\'public conhecimentos_emitidos\'[data_autorizacao] >= {_dax_data(start_date)} && '
        f'\'public conhecimentos_emitidos\'[data_autorizacao] <= {_dax_data(end_date)})))'
    )
    result = execute_dax(token, dax_receita, dataset_id=CONFIG['dre_dataset_id'])
    rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
    receita_bruta = float((rows[0].get('[total]') if rows else 0) or 0)

    # Despesas filtradas por REF (YYYY/MM)
    refs = _gerar_refs_periodo(start_date, end_date)
    grupos = {'Deduções': 0.0, 'Operacional': 0.0, 'Administrativo': 0.0,
              'Financeiro': 0.0, 'Impostos': 0.0, 'Investimento': 0.0, 'Retirada': 0.0}
    if refs:
        dax_despesas = (
            f'EVALUATE CALCULATETABLE('
            f'SUMMARIZE(\'public consulta_despesas_477\', '
            f'\'public consulta_despesas_477\'[descr_evento], '
            f'"Total", SUM(\'public consulta_despesas_477\'[vlr_final])), '
            f'\'public consulta_despesas_477\'[REF] IN {_dax_lista_refs(refs)})'
        )
        result = execute_dax(token, dax_despesas, dataset_id=CONFIG['dre_dataset_id'])
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)
        for r in data:
            descr = r.get('descr_evento')
            valor = float(r.get('Total') or 0)
            if descr in MAPA_DRE:
                grupo = MAPA_DRE[descr][0]
                grupos[grupo] += valor

    receita_liquida   = receita_bruta - grupos['Deduções']
    ebitda            = receita_liquida - grupos['Operacional'] - grupos['Administrativo']
    lair              = ebitda - grupos['Financeiro']
    lucro_liquido     = lair - grupos['Impostos']
    pos_investimento  = lucro_liquido - grupos['Investimento']
    resultado_final   = pos_investimento - grupos['Retirada']

    return {
        'receita_bruta':            receita_bruta,
        'deducoes':                 grupos['Deduções'],
        'receita_liquida':          receita_liquida,
        'custo_operacional':        grupos['Operacional'],
        'despesas_administrativas': grupos['Administrativo'],
        'ebitda':                   ebitda,
        'despesas_financeiras':     grupos['Financeiro'],
        'lair':                     lair,
        'impostos':                 grupos['Impostos'],
        'lucro_liquido':            lucro_liquido,
        'investimentos':            grupos['Investimento'],
        'pos_investimento':         pos_investimento,
        'retiradas':                grupos['Retirada'],
        'resultado_final':          resultado_final,
    }


def _parse_meses_param(meses_str):
    """'2025-01,2025-03,2026-01' → lista ordenada [(2025,1),(2025,3),(2026,1)]"""
    pares = []
    for m in meses_str.split(','):
        m = m.strip()
        if not m: continue
        y, mo = m.split('-')
        pares.append((int(y), int(mo)))
    return sorted(set(pares))


def _meses_para_periodos(pares, dia_ini=None, dia_fim=None):
    """[(2025,1),(2025,3)] → [(y, m, nome_curto, primeiro_dia, ultimo_dia), ...]

    `dia_ini`/`dia_fim` recortam a janela DENTRO de cada mês (modo fracionado).
    Os dias são grampeados no tamanho real do mês — pedir "1 a 31" em fevereiro
    devolve 1 a 28/29, senão date() estouraria em mês curto.

    ATENÇÃO: o recorte por dia só afeta a RECEITA. A despesa é filtrada por
    competência (`REF`, YYYY/MM) em `_calcular_dre_periodo` e continua vindo do
    mês inteiro — não existe fracionamento de despesa por dia (decisão do
    negócio: a competência do SSW não é uma data). Ver `fracionado` na resposta
    de /api/dre, que obriga o front a rotular as duas bases.
    """
    from datetime import date
    from calendar import monthrange
    nomes_curtos = ['jan', 'fev', 'mar', 'abr', 'mai', 'jun',
                    'jul', 'ago', 'set', 'out', 'nov', 'dez']
    result = []
    for y, m in pares:
        ult = monthrange(y, m)[1]
        d1 = min(max(dia_ini or 1, 1), ult)
        d2 = min(dia_fim or ult, ult)
        if d2 < d1:
            d2 = d1
        nome = f"{nomes_curtos[m-1]}/{str(y)[2:]}"
        result.append((y, m, nome, date(y, m, d1), date(y, m, d2)))
    return result


def _iterar_meses(start_date, end_date):
    """Gera tuplas (ano, mes, nome_mes, primeiro_dia, ultimo_dia) entre as datas."""
    from datetime import date
    from calendar import monthrange
    nomes = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    cur_y, cur_m = start_date.year, start_date.month
    end_y, end_m = end_date.year, end_date.month
    while (cur_y, cur_m) <= (end_y, end_m):
        ult_dia = monthrange(cur_y, cur_m)[1]
        primeiro = date(cur_y, cur_m, 1)
        ultimo = date(cur_y, cur_m, ult_dia)
        if primeiro < start_date: primeiro = start_date
        if ultimo > end_date: ultimo = end_date
        yield (cur_y, cur_m, nomes[cur_m - 1], primeiro, ultimo)
        cur_m += 1
        if cur_m > 12:
            cur_m = 1
            cur_y += 1


@app.route('/api/dre')
@page_required('dre')
def api_dre():
    from datetime import datetime

    # ── Recorte por dia (modo fracionado) — só vale junto com `meses=` ──
    # Fraciona APENAS a receita; a despesa continua por competência (mês cheio).
    dia_ini = dia_fim = None
    if request.args.get('dia_ini') or request.args.get('dia_fim'):
        try:
            dia_ini = int(request.args.get('dia_ini') or 1)
            dia_fim = int(request.args.get('dia_fim') or 31)
        except (TypeError, ValueError):
            return jsonify({'ok': False, 'error': 'dia_ini/dia_fim devem ser números'}), 400
        if not (1 <= dia_ini <= 31 and 1 <= dia_fim <= 31):
            return jsonify({'ok': False, 'error': 'dia_ini/dia_fim devem estar entre 1 e 31'}), 400
        if dia_ini > dia_fim:
            return jsonify({'ok': False, 'error': 'dia_ini não pode ser maior que dia_fim'}), 400

    fracionado = (dia_ini, dia_fim) != (None, None) and (dia_ini, dia_fim) != (1, 31)

    meses_param = request.args.get('meses')
    if meses_param:
        try:
            pares = _parse_meses_param(meses_param)
            if not pares:
                raise ValueError('lista vazia')
            meses = _meses_para_periodos(pares, dia_ini, dia_fim)
        except Exception:
            return jsonify({'ok': False, 'error': 'Parâmetro meses inválido. Use YYYY-MM,YYYY-MM,...'}), 400
    else:
        try:
            start = datetime.strptime(request.args.get('start'), '%Y-%m-%d').date()
            end   = datetime.strptime(request.args.get('end'), '%Y-%m-%d').date()
        except (TypeError, ValueError):
            return jsonify({'ok': False, 'error': 'Informe meses=YYYY-MM,... ou start/end'}), 400
        meses = list(_iterar_meses(start, end))

    estrutura = [{'linha': l, 'tipo': t, 'key': k} for l, t, k in DRE_LINHAS]

    # Metadados do recorte por dia — o front usa p/ rotular as bases e suprimir
    # os visuais que subtraem despesa cheia de receita parcial.
    meta_frac = {'fracionado': fracionado, 'dia_ini': dia_ini, 'dia_fim': dia_fim}

    if len(meses) == 1:
        (_, _, _, prim, ult) = meses[0]
        totais = _calcular_dre_periodo(prim, ult)
        rb = totais['receita_bruta']
        for item in estrutura:
            v = totais[item['key']]
            item['valor'] = v
            item['pct'] = (v / rb) if (item['tipo'] == 'Subtotal' and rb) else None
        return jsonify({'ok': True, 'modo': 'acumulado', 'estrutura': estrutura, **meta_frac})

    # Modo mensal
    totais_por_mes = []
    totais_geral = {k: 0.0 for _, _, k in DRE_LINHAS}
    for (_, _, nome, prim, ult) in meses:
        t = _calcular_dre_periodo(prim, ult)
        totais_por_mes.append({'nome': nome, 'totais': t})
        for k in totais_geral:
            totais_geral[k] += t[k]

    for item in estrutura:
        item['meses'] = []
        for tm in totais_por_mes:
            v = tm['totais'][item['key']]
            rb = tm['totais']['receita_bruta']
            item['meses'].append({
                'nome': tm['nome'],
                'valor': v,
                'pct': (v / rb) if (item['tipo'] == 'Subtotal' and rb) else None
            })
        v = totais_geral[item['key']]
        rb_total = totais_geral['receita_bruta']
        item['total'] = v
        item['total_pct'] = (v / rb_total) if (item['tipo'] == 'Subtotal' and rb_total) else None

    return jsonify({'ok': True, 'modo': 'mensal', 'meses': [m[2] for m in meses],
                    'estrutura': estrutura, **meta_frac})


def _query_despesas_periodo(start, end, grupo=None, evento=None):
    """Filtra consulta_despesas_477 via DAX por REF (YYYY/MM), opcionalmente por grupo ou evento específico."""
    from datetime import datetime
    if isinstance(start, str):
        start = datetime.strptime(start, '%Y-%m-%d').date()
    if isinstance(end, str):
        end = datetime.strptime(end, '%Y-%m-%d').date()

    refs = _gerar_refs_periodo(start, end)
    if not refs:
        return [], []

    filtro_ref = f"'public consulta_despesas_477'[REF] IN {_dax_lista_refs(refs)}"

    if evento:
        dax = (
            f'EVALUATE FILTER(\'public consulta_despesas_477\', '
            f'{filtro_ref} && \'public consulta_despesas_477\'[descr_evento] = "{evento}")'
        )
    elif grupo:
        eventos = [e for e, (g, _) in MAPA_DRE.items() if g == grupo]
        if not eventos:
            return [], []
        lista_eventos = '{ ' + ', '.join(f'"{e}"' for e in eventos) + ' }'
        dax = (
            f'EVALUATE FILTER(\'public consulta_despesas_477\', '
            f'{filtro_ref} && \'public consulta_despesas_477\'[descr_evento] IN {lista_eventos})'
        )
    else:
        dax = f'EVALUATE FILTER(\'public consulta_despesas_477\', {filtro_ref})'

    token = get_token()
    result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
    rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
    data = clean_rows(rows)
    cols = list(data[0].keys()) if data else []
    return cols, data


def _query_conhecimentos_periodo(start, end):
    """Filtra conhecimentos_emitidos via DAX por data_autorizacao."""
    from datetime import datetime
    if isinstance(start, str):
        start = datetime.strptime(start, '%Y-%m-%d').date()
    if isinstance(end, str):
        end = datetime.strptime(end, '%Y-%m-%d').date()

    dax = (
        f'EVALUATE FILTER(\'public conhecimentos_emitidos\', '
        f'\'public conhecimentos_emitidos\'[data_autorizacao] >= {_dax_data(start)} && '
        f'\'public conhecimentos_emitidos\'[data_autorizacao] <= {_dax_data(end)})'
    )
    token = get_token()
    result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
    rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
    data = clean_rows(rows)
    cols = list(data[0].keys()) if data else []
    return cols, data


@app.route('/api/dre/detalhamento')
@page_required('dre')
def api_dre_detalhamento():
    """Retorna despesas agrupadas por Subgrupo (com eventos dentro) para o período/lista de meses."""
    from datetime import datetime
    meses_param = request.args.get('meses')
    if meses_param:
        try:
            pares = _parse_meses_param(meses_param)
        except Exception:
            return jsonify({'ok': False, 'error': 'Parâmetro meses inválido'}), 400
        refs = [f"{y:04d}/{m:02d}" for y, m in pares]
    else:
        try:
            start = datetime.strptime(request.args.get('start'), '%Y-%m-%d').date()
            end   = datetime.strptime(request.args.get('end'), '%Y-%m-%d').date()
        except (TypeError, ValueError):
            return jsonify({'ok': False, 'error': 'Datas inválidas'}), 400
        refs = _gerar_refs_periodo(start, end)

    if not refs:
        return jsonify({'ok': True, 'subgrupos': []})

    try:
        dax = (
            f'EVALUATE CALCULATETABLE('
            f'SUMMARIZE(\'public consulta_despesas_477\', '
            f'\'public consulta_despesas_477\'[descr_evento], '
            f'"Total", SUM(\'public consulta_despesas_477\'[vlr_final])), '
            f'\'public consulta_despesas_477\'[REF] IN {_dax_lista_refs(refs)})'
        )
        token = get_token()
        result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

    # Agrupa por Subgrupo
    subgrupos = {}
    for r in data:
        descr = r.get('descr_evento')
        valor = float(r.get('Total') or 0)
        if descr in MAPA_DRE:
            grupo, sub = MAPA_DRE[descr]
            if sub not in subgrupos:
                subgrupos[sub] = {'nome': sub, 'grupo': grupo, 'total': 0.0, 'eventos': []}
            subgrupos[sub]['total'] += valor
            subgrupos[sub]['eventos'].append({'descr_evento': descr, 'total': valor})

    # Ordena: subgrupos por nome, eventos por valor desc
    lista = sorted(subgrupos.values(), key=lambda x: x['nome'])
    for s in lista:
        s['eventos'].sort(key=lambda e: e['total'], reverse=True)

    total_geral = sum(s['total'] for s in lista)
    return jsonify({'ok': True, 'subgrupos': lista, 'total_geral': total_geral})


@app.route('/api/dre/despesas')
@page_required('despesas')
def api_dre_despesas():
    start = request.args.get('start')
    end = request.args.get('end')
    grupo = request.args.get('grupo')
    evento = request.args.get('evento')
    if not start or not end:
        return jsonify({'ok': False, 'error': 'Informe start e end (YYYY-MM-DD)'}), 400
    try:
        cols, data = _query_despesas_periodo(start, end, grupo, evento)
        return jsonify({'ok': True, 'columns': cols, 'data': data, 'count': len(data),
                        'grupo': grupo, 'evento': evento})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/dre/conhecimentos')
@page_required('conhecimentos')
def api_dre_conhecimentos():
    start = request.args.get('start')
    end = request.args.get('end')
    if not start or not end:
        return jsonify({'ok': False, 'error': 'Informe start e end (YYYY-MM-DD)'}), 400
    try:
        cols, data = _query_conhecimentos_periodo(start, end)
        return jsonify({'ok': True, 'columns': cols, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Grupo econômico (raízes de CNPJ que são o mesmo cliente) ──
# A consolidação padrão de cliente é a raiz do CNPJ (8 primeiros dígitos), o que já junta as
# filiais. Quando o mesmo grupo fatura por mais de uma raiz, a raiz secundária é mapeada aqui
# para a principal — senão o cliente aparece partido em duas linhas. Vale para o Faturamento
# por Tomador e para o recorte Cliente da Análise por Veículo (ambos passam por _raiz_cnpj).
GRUPOS_ECONOMICOS = {
    '43214055': '18485037',   # MARTINS COM E SERV DE DISTR → MARTINS URN-MG DISTRIBUICAO
}


def _raiz_cnpj(cnpj):
    """Raiz do CNPJ (8 dígitos) já resolvida para a raiz do grupo econômico.
    Devolve None quando não há 8 dígitos — aí o chamador cai no fallback por nome.
    Em viagem mista o `cnpj_pagador` vem como lista; mantém-se o comportamento
    histórico de considerar o primeiro CNPJ da lista."""
    import re as _re
    d = _re.sub(r'\D', '', str(cnpj or ''))
    if len(d) < 8:
        return None
    raiz = d[:8]
    return GRUPOS_ECONOMICOS.get(raiz, raiz)


@app.route('/api/faturamento/tomadores')
@page_required('faturamento')
def api_faturamento_tomadores():
    """Matriz faturamento por tomador (cliente_pagador) × mês, consolidado por raiz de CNPJ.
    Fonte: conhecimentos_emitidos (valor_frete + distinct primeiro_manifesto), filtrado por ano."""
    try:
        ano = int(request.args.get('ano', 2026))
    except Exception:
        ano = 2026

    # DAX agregado: 1 linha por (cnpj_pagador, cliente_pagador, mês). Mês derivado via ADDCOLUMNS.
    dax = (
        "EVALUATE SUMMARIZE("
        "ADDCOLUMNS("
        f"FILTER('public conhecimentos_emitidos', "
        f"'public conhecimentos_emitidos'[data_autorizacao] >= DATE({ano},1,1) && "
        f"'public conhecimentos_emitidos'[data_autorizacao] <= DATE({ano},12,31)), "
        "\"@mes\", FORMAT('public conhecimentos_emitidos'[data_autorizacao], \"MM\")), "
        "'public conhecimentos_emitidos'[cnpj_pagador], "
        "'public conhecimentos_emitidos'[cliente_pagador], "
        "[@mes], "
        "\"faturamento\", SUM('public conhecimentos_emitidos'[valor_frete]), "
        # cargas = viagens reais: NÃO conta CTe sem manifesto (cobrança que não é viagem)
        "\"cargas\", CALCULATE(DISTINCTCOUNT('public conhecimentos_emitidos'[primeiro_manifesto]), "
        "NOT(ISBLANK('public conhecimentos_emitidos'[primeiro_manifesto])) && "
        "'public conhecimentos_emitidos'[primeiro_manifesto] <> \"\"))"
    )

    try:
        token = get_token()
        result = execute_dax(token, dax)
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        linhas = clean_rows(rows)
    except requests.exceptions.HTTPError as e:
        detail = ''
        try: detail = e.response.json()
        except Exception: detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

    # Consolida por raiz de CNPJ (+ grupo econômico). Sem CNPJ → cai no nome (fallback).
    tomadores = {}        # chave -> {cnpj_raiz, nome, _nomes:{nome:freq}, meses:{1..12:{fat,cargas}}, total_*}
    totais_mes = {m: {'faturamento': 0.0, 'cargas': 0} for m in range(1, 13)}

    for r in linhas:
        nome = (str(r.get('cliente_pagador') or '').strip()) or '(sem nome)'
        raiz = _raiz_cnpj(r.get('cnpj_pagador'))   # já resolve o grupo econômico
        chave = raiz or f"nome::{nome.upper()}"
        try:
            mes = int(r.get('@mes'))
        except Exception:
            continue
        fat = float(r.get('faturamento') or 0)
        carg = int(r.get('cargas') or 0)

        t = tomadores.get(chave)
        if not t:
            t = {'cnpj_raiz': raiz, 'nome': nome, '_nomes': {},
                 'meses': {m: {'faturamento': 0.0, 'cargas': 0} for m in range(1, 13)},
                 'total_faturamento': 0.0, 'total_cargas': 0}
            tomadores[chave] = t
        t['_nomes'][nome] = t['_nomes'].get(nome, 0) + 1
        if 1 <= mes <= 12:
            t['meses'][mes]['faturamento'] += fat
            t['meses'][mes]['cargas'] += carg
            t['total_faturamento'] += fat
            t['total_cargas'] += carg
            totais_mes[mes]['faturamento'] += fat
            totais_mes[mes]['cargas'] += carg

    saida = []
    for t in tomadores.values():
        # rótulo = nome mais frequente da raiz
        t['nome'] = max(t['_nomes'].items(), key=lambda kv: kv[1])[0]
        del t['_nomes']
        t['total_faturamento'] = round(t['total_faturamento'], 2)
        for m in range(1, 13):
            t['meses'][m]['faturamento'] = round(t['meses'][m]['faturamento'], 2)
        saida.append(t)
    saida.sort(key=lambda x: x['total_faturamento'], reverse=True)
    for m in range(1, 13):
        totais_mes[m]['faturamento'] = round(totais_mes[m]['faturamento'], 2)

    return jsonify({'ok': True, 'ano': ano, 'tomadores': saida, 'totais_mes': totais_mes,
                    'count': len(saida)})


# A implementação vive em placas.py para ser compartilhada com o PGR sem import
# circular (server importa pgr) e sem duplicar a lógica. Aliases mantidos para
# não mexer em nenhum ponto de chamada existente.
_placa_mercosul = placas.mercosul
_placa_grafias = placas.grafias


# Regra interna: placas que continuam no cadastro como Rizza mas NÃO são mais frota
# (ex.: cavalo vendido ainda no nome da Rizza). Normalizadas em Mercosul.
PLACAS_VENDIDAS = {'AZM6E29'}


def _cadastro_veiculos(token):
    """Cadastro de veículos (veiculos_045) com placa normalizada em Mercosul e deduplicado.
    Retorna {placa_norm: {'proprietario','tipo','disponivel','modelo'}}.

    Colisão de grafia: a conversão antiga→Mercosul (LLL+4díg → troca o 5º char) pode gerar
    uma string idêntica à placa Mercosul REAL de outro veículo (ex.: HOA0466→HOA0E66, que é a
    Mercosul real de outra carreta). Nesses casos, prefere a entrada cuja placa CRUA já está em
    Mercosul (identidade atual) em vez da antiga convertida. Entre formatos iguais, mantém a 1ª."""
    res = execute_dax(token, "EVALUATE 'public veiculos_045'")
    linhas = clean_rows(res.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))
    cad = {}
    cad_merc = {}  # placa_norm -> a entrada guardada veio de placa crua já-Mercosul?
    for r in linhas:
        raw = placas.limpar(r.get('placa'))
        p = _placa_mercosul(raw)
        if not p:
            continue
        eh_merc = placas.eh_mercosul(raw)  # placa crua já é Mercosul
        if p in cad and not (eh_merc and not cad_merc.get(p)):
            continue  # mantém a atual, salvo quando a nova é Mercosul genuína e a atual não era
        cad[p] = {'proprietario': r.get('proprietario'), 'tipo': r.get('tipo'),
                  'disponivel': r.get('disponivel'), 'modelo': r.get('modelo')}
        cad_merc[p] = eh_merc
    return cad


def _km_hodometro(hods_em_ordem):
    """KM real pelo hodômetro do abastecimento: soma de deltas consecutivos válidos,
    ignorando retrocesso (<=0) e saltos absurdos (>3000 km = erro de digitação na bomba).
    Espera a lista de leituras já em ordem cronológica."""
    h = [x for x in hods_em_ordem if x and x > 0]
    return sum(d for d in (h[i] - h[i - 1] for i in range(1, len(h))) if 0 < d < 3000)


def _pneus_por_veiculo(modelo, tipo):
    """Nº de pneus do veículo a partir da config de eixo no modelo.
    Carreta=12; truck=6 (pelo tipo); cavalo: 6X4=10, 6X2=8, 4X2=6;
    cavalo sem config reconhecível → 8 (fallback)."""
    import re
    t = (tipo or '').upper()
    if t == 'CARRETA':
        return 12
    if t == 'TRUCK':
        return 6
    m = re.sub(r'\s+', '', str(modelo or '')).upper()
    if '6X4' in m:
        return 10
    if '6X2' in m:
        return 8
    if '4X2' in m:
        return 6
    return 8


def _dax_rows(token, dax_q):
    """Executa DAX e devolve as linhas já limpas (clean_rows)."""
    res = execute_dax(token, dax_q)
    return clean_rows(res.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))


def _dax_val(token, dax_q):
    """Executa DAX de uma linha/coluna 'v' e devolve o float."""
    r = _dax_rows(token, dax_q)
    return float((r[0] if r else {}).get('v') or 0)


def _pool_pneu_split(token, meses_set, cadastro):
    """Pool de PNEU (eventos 5411/5412, sem placa) dividido cavalo×carreta pelo nº de pneus dos
    veículos Rizza ATIVOS no período. Retorna (pool_pneu, pool_pneu_cav, pool_pneu_car)."""
    DZ = "'public consulta_despesas_477'"
    AR = "'Auditoria Receita'"
    anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'
    pool_pneu = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento] IN {{\"5411\",\"5412\"}} && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
    rizza_cav_cad = {p: v for p, v in cadastro.items()
                     if v.get('proprietario') == 'RIZZA TRANSPORTES LTDA' and v.get('tipo') != 'CARRETA'}
    rizza_car_cad = {p for p, v in cadastro.items()
                     if v.get('proprietario') == 'RIZZA TRANSPORTES LTDA' and v.get('tipo') == 'CARRETA'}

    def _ativas(col):
        chave = col.strip('[]')
        rr = _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({AR}, "
                       f"FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && "
                       f"{AR}[Tipo Operacao] IN {{\"FROTA\",\"AGREGADO\"}} && NOT(ISBLANK({AR}{col}))), {AR}{col})")
        return {_placa_mercosul(x.get(chave)) for x in rr}

    cav_ativos = _ativas('[placa_cavalo]') & set(rizza_cav_cad)
    car_ativos = _ativas('[placa_carreta]') & rizza_car_cad
    tires_cav = sum(_pneus_por_veiculo(rizza_cav_cad[p].get('modelo'), rizza_cav_cad[p].get('tipo')) for p in cav_ativos)
    tires_car = 12 * len(car_ativos)
    tot = (tires_cav + tires_car) or 1
    pool_cav = pool_pneu * tires_cav / tot
    return pool_pneu, pool_cav, pool_pneu - pool_cav


# Componentes de custo da frota por cavalo (chaves do dict por placa)
_COST_KEYS = ('pedagio', 'combustivel', 'litros', 'km_hodometro', 'arla', 'litros_arla',
              'pessoal', 'manut_cavalo', 'seguro', 'rastreador', 'pneu')


def _custo_frota_por_cavalo(token, meses_set, cadastro, fat_por_placa=None):
    """Custo real por cavalo frota no período (mesma lógica da visão Cavalo+Frota), reutilizável.

    `fat_por_placa` ({placa_norm: faturamento}) define o share/n_cav e quais cavalos recebem; se None,
    é calculado a partir das viagens FROTA. Retorna (custos_por_placa, totais), onde custos[placa] tem
    os componentes (_COST_KEYS) + `custo_total`. Mantém os mesmos valores da tela quando o `fat_por_placa`
    vem do `saida` da visão frota."""
    AR = "'Auditoria Receita'"; SP = "'public semparar_lancamentos'"
    VC = "'public abastecimentos_valecard'"; DZ = "'public consulta_despesas_477'"
    anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'

    if fat_por_placa is None:
        fat_por_placa = {}
        for r in _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({AR}, "
                           f"FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && {AR}[Tipo Operacao]=\"FROTA\" "
                           f"&& NOT(ISBLANK({AR}[placa_cavalo]))), {AR}[placa_cavalo], \"f\", SUM({AR}[receita_rateada]))"):
            p = _placa_mercosul(r.get('placa_cavalo'))
            if not p or p in PLACAS_VENDIDAS:
                continue
            fat_por_placa[p] = round(fat_por_placa.get(p, 0.0) + float(r.get('f') or 0), 2)

    periodo_sp = f"(RIGHT({SP}[data],4) & \"-\" & MID({SP}[data],4,2)) IN {meses_set}"
    periodo_vc = f"FORMAT({VC}[dch_data], \"YYYY-MM\") IN {meses_set}"

    ped = {}
    for r in _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({SP}, {periodo_sp}), {SP}[placa_veiculo], \"v\", SUM({SP}[valor]))"):
        p = _placa_mercosul(r.get('placa_veiculo'))
        ped[p] = ped.get(p, 0.0) + float(r.get('v') or 0)

    comb = {}
    for cat, filtro in (('diesel', f"SEARCH(\"ARLA\",{VC}[produto],1,0)=0"),
                        ('arla',   f"SEARCH(\"ARLA\",{VC}[produto],1,0)>0")):
        for r in _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({VC}, {periodo_vc} && {filtro}), {VC}[placa], "
                           f"\"v\", SUM({VC}[mcd_valor_total]), \"lt\", SUM({VC}[ncd_quantidade]))"):
            p = _placa_mercosul(r.get('placa'))
            d = comb.setdefault(p, {'diesel': 0.0, 'litros': 0.0, 'arla': 0.0, 'litros_arla': 0.0})
            if cat == 'arla':
                d['arla'] += float(r.get('v') or 0); d['litros_arla'] += float(r.get('lt') or 0)
            else:
                d['diesel'] += float(r.get('v') or 0); d['litros'] += float(r.get('lt') or 0)

    km_hod = {}; _fills = {}
    for r in _dax_rows(token, f"EVALUATE SELECTCOLUMNS(FILTER({VC}, {periodo_vc}), "
                       f"\"p\",{VC}[placa],\"dt\",{VC}[dch_data],\"hod\",{VC}[nsd_hodometro])"):
        _fills.setdefault(_placa_mercosul(r.get('p')), []).append((str(r.get('dt') or ''), float(r.get('hod') or 0)))
    for p, lst in _fills.items():
        lst.sort()
        km_hod[p] = _km_hodometro([h for _, h in lst])

    pessoal_total = _dax_val(token, f"EVALUATE ROW(\"v\", CALCULATE(SUM('public custo_pessoal'[total_mes]), 'public custo_pessoal'[competencia] IN {meses_set}))")
    manut_cavalo_total = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento] IN {{\"5150\",\"5154\"}} && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
    seguro_total = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento]=\"5402\" && SEARCH(\"BVIX\",{DZ}[nome_fornecedor],1,0)=0 && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
    rastreador_total = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, SEARCH(\"AUTOTRAC\",{DZ}[nome_fornecedor],1,0)>0 && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
    pool_pneu, pool_pneu_cav, _ = _pool_pneu_split(token, meses_set, cadastro)

    sum_fat = sum(fat_por_placa.values()) or 1.0
    n_cav = len(fat_por_placa) or 1
    custos = {}
    for p, fat in fat_por_placa.items():
        c = comb.get(p, {})
        share = fat / sum_fat
        d = {
            'pedagio': round(ped.get(p, 0.0), 2),
            'combustivel': round(c.get('diesel', 0.0), 2),
            'litros': round(c.get('litros', 0.0), 1),
            'km_hodometro': round(km_hod.get(p, 0.0), 0),
            'arla': round(c.get('arla', 0.0), 2),
            'litros_arla': round(c.get('litros_arla', 0.0), 1),
            'pessoal': round(pessoal_total * share, 2),
            'manut_cavalo': round(manut_cavalo_total * share, 2),
            'seguro': round(seguro_total / n_cav, 2),
            'rastreador': round(rastreador_total / n_cav, 2),
            'pneu': round(pool_pneu_cav * share, 2),
        }
        d['custo_total'] = round(d['pedagio'] + d['combustivel'] + d['arla'] + d['pessoal']
                                 + d['manut_cavalo'] + d['seguro'] + d['rastreador'] + d['pneu'], 2)
        custos[p] = d
    totais = {'pessoal': round(pessoal_total, 2), 'manut_cavalo': round(manut_cavalo_total, 2),
              'seguro': round(seguro_total, 2), 'rastreador': round(rastreador_total, 2),
              'pneu': round(pool_pneu_cav, 2), 'pneu_total': round(pool_pneu, 2)}
    return custos, totais


def _norm_manifesto(s):
    """Normaliza nº de manifesto p/ casar Auditoria (UDI027978-1) × conhecimentos (UDI 026011-8)."""
    import re as _re
    return _re.sub(r'[^A-Za-z0-9]', '', str(s or '')).upper()


def _ctrc_tomador_map(token, meses_comp):
    """Mapa CTRC → tomador (raiz de CNPJ, nome) e valor_frete, do conhecimentos_emitidos.

    O CTRC é o elo confiável entre a viagem (Auditoria) e o cliente — o nº de manifesto da Auditoria
    nem sempre é o `primeiro_manifesto` do conhecimentos. Janela = ano(s) dos meses pedidos."""
    CE = "'public conhecimentos_emitidos'"
    anos = sorted({m[:4] for m in meses_comp})
    anos_set = '{' + ','.join(f'"{a}"' for a in anos) + '}'
    rows = _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({CE}, FORMAT({CE}[data_autorizacao],\"YYYY\") IN {anos_set}), "
        f"{CE}[serie_numero_ctrc], {CE}[cnpj_pagador], {CE}[cliente_pagador], \"vf\", SUM({CE}[valor_frete]))")
    mp = {}
    for r in rows:
        c = _norm_manifesto(r.get('serie_numero_ctrc'))   # mesmo normalizador (tira espaço/pontuação)
        if not c:
            continue
        raiz = _raiz_cnpj(r.get('cnpj_pagador'))   # já resolve o grupo econômico
        nome = (str(r.get('cliente_pagador') or '').strip()) or '(sem nome)'
        e = mp.get(c)
        if e is None:
            mp[c] = {'raiz': raiz, 'chave': (raiz or f"nome::{nome.upper()}"), 'nome': nome, 'vf': float(r.get('vf') or 0)}
        else:
            e['vf'] += float(r.get('vf') or 0)
    return mp


def _veiculos_margem_cliente(token, meses_comp, meses_set, tipos, tipos_dax):
    """Margem por cliente (consolidado por raiz de CNPJ), trabalhando POR VIAGEM (CTRB).

    A 'Auditoria Receita' traz 1 linha por viagem com 1 pagador; viagens mistas (vários tomadores)
    são divididas aqui na **proporção do valor_frete dos CTRCs** (conhecimentos_emitidos), casando pelo
    nº do manifesto. Cada viagem distribui sua receita_rateada + custo (cavalo+carreta) entre os tomadores;
    nº de viagens/clientes passa a refletir a participação real. Custo só da operação (sem deduções)."""
    AR = "'Auditoria Receita'"
    DZ = "'public consulta_despesas_477'"
    anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'

    cadastro = _cadastro_veiculos(token)

    # Custo total da frota (cavalos). É diluído proporcional à RECEITA DE FROTA de cada cliente
    # (taxa única = custo total / receita total de frota) — não por cavalo específico nem por faturamento total.
    custos_cav, _t = _custo_frota_por_cavalo(token, meses_set, cadastro)
    total_custo_cavalo = sum(c.get('custo_total', 0.0) for c in custos_cav.values())

    # Taxa de custo de carreta (manut+pneu) por R$ de receita de carreta Rizza
    rizza_carretas = {p for p, v in cadastro.items()
                      if v.get('proprietario') == 'RIZZA TRANSPORTES LTDA' and v.get('tipo') == 'CARRETA'}
    manut_carreta_total = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento] IN {{\"5153\",\"5155\"}} && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
    _, _, pool_pneu_car = _pool_pneu_split(token, meses_set, cadastro)
    univ = _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({AR}, FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} "
        f"&& {AR}[Tipo Operacao] IN {{\"FROTA\",\"AGREGADO\"}} && NOT(ISBLANK({AR}[placa_carreta]))), {AR}[placa_carreta], \"f\", SUM({AR}[receita_rateada]))")
    base_fat = sum(float(r.get('f') or 0) for r in univ if _placa_mercosul(r.get('placa_carreta')) in rizza_carretas)
    taxa_car = (manut_carreta_total + pool_pneu_car) / (base_fat or 1.0)

    # Viagens (1 linha por CTRB) da Auditoria no período/tipos
    viagens = _dax_rows(token, f"EVALUATE SELECTCOLUMNS(FILTER({AR}, "
        f"FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && {AR}[Tipo Operacao] IN {tipos_dax}), "
        f"\"ctrb\",{AR}[CTRB],\"ctrc\",{AR}[CTRC],\"cnpj\",{AR}[cnpj_pagador],\"nome\",{AR}[cliente_pagador],"
        f"\"tipo\",{AR}[Tipo Operacao],\"cav\",{AR}[placa_cavalo],\"car\",{AR}[placa_carreta],"
        f"\"rec\",{AR}[receita_rateada],\"frete\",{AR}[frete_motorista_total],"
        f"\"km\",COALESCE(LOOKUPVALUE('public rotas_km'[km],'public rotas_km'[cidade_uf_origem],{AR}[cidade_uf_origem],"
        f"'public rotas_km'[cidade_uf_destino],{AR}[cidade_uf_destino]),{AR}[distancia_km]))")

    ctrc_map = _ctrc_tomador_map(token, meses_comp)

    def _chave(cnpj_raw, nome):
        raiz = _raiz_cnpj(cnpj_raw)   # já resolve o grupo econômico
        return (raiz or f"nome::{(nome or '').strip().upper()}"), raiz

    agg = {}
    n_casados = n_total = 0

    def _bucket(chave, raiz):
        a = agg.get(chave)
        if a is None:
            a = {'cnpj_raiz': raiz, '_nomes': {}, 'faturamento': 0.0, 'rec_frota': 0.0, 'rec_agregado': 0.0,
                 'rec_carreteiro': 0.0, 'frete_terceiros': 0.0, 'custo_cavalo': 0.0, 'custo_carreta': 0.0,
                 'km': 0.0, '_viagens': set()}
            agg[chave] = a
        return a

    for v in viagens:
        n_total += 1
        rec = float(v.get('rec') or 0)
        tipo = v.get('tipo') or '—'
        frete = float(v.get('frete') or 0) if tipo != 'FROTA' else 0.0
        km = float(v.get('km') or 0)
        ctrb = v.get('ctrb')
        car = _placa_mercosul(v.get('car'))
        # custo de carreta da viagem (o cavalo é diluído depois, proporcional à receita de frota)
        custo_car_v = taxa_car * rec if car in rizza_carretas else 0.0
        # tomadores da viagem pela proporção do valor_frete dos CTRCs (liga via CTRC, não manifesto)
        by_raiz = {}
        for c in str(v.get('ctrc') or '').split(','):
            e = ctrc_map.get(_norm_manifesto(c))
            if not e:
                continue
            b = by_raiz.get(e['chave'])
            if b is None:
                b = {'raiz': e['raiz'], 'vf': 0.0, 'nomes': {}}
                by_raiz[e['chave']] = b
            b['vf'] += e['vf']; b['nomes'][e['nome']] = b['nomes'].get(e['nome'], 0) + 1
        tot_vf = sum(b['vf'] for b in by_raiz.values())
        if by_raiz and tot_vf > 0:
            n_casados += 1
            dist = [(ch, b['raiz'], b['vf'] / tot_vf, max(b['nomes'], key=b['nomes'].get)) for ch, b in by_raiz.items()]
        else:
            nome = (str(v.get('nome') or '').strip()) or '(sem nome)'
            ch, rz = _chave(v.get('cnpj'), nome)
            dist = [(ch, rz, 1.0, nome)]
        for ch, rz, frac, nome in dist:
            a = _bucket(ch, rz)
            a['_nomes'][nome] = a['_nomes'].get(nome, 0) + 1
            a['faturamento'] += rec * frac
            if tipo == 'FROTA':
                a['rec_frota'] += rec * frac
            elif tipo == 'AGREGADO':
                a['rec_agregado'] += rec * frac
            else:
                a['rec_carreteiro'] += rec * frac
            a['frete_terceiros'] += frete * frac
            a['custo_carreta'] += custo_car_v * frac
            a['km'] += km * frac
            if ctrb:
                a['_viagens'].add(ctrb)

    # Custo de cavalo diluído proporcional à receita de frota de cada cliente (taxa única)
    total_rec_frota = sum(a['rec_frota'] for a in agg.values()) or 1.0
    taxa_cav = total_custo_cavalo / total_rec_frota
    saida = []
    for a in agg.values():
        a['custo_cavalo'] = taxa_cav * a['rec_frota']
        a['dim'] = max(a['_nomes'].items(), key=lambda kv: kv[1])[0] if a['_nomes'] else '(sem nome)'
        a['viagens'] = len(a.pop('_viagens'))
        a.pop('_nomes', None)
        for k in ('faturamento', 'rec_frota', 'rec_agregado', 'rec_carreteiro', 'frete_terceiros',
                  'custo_cavalo', 'custo_carreta'):
            a[k] = round(a[k], 2)
        a['custo_frota'] = round(a['custo_cavalo'] + a['custo_carreta'], 2)
        a['km'] = round(a['km'], 1)
        a['resultado'] = round(a['faturamento'] - a['frete_terceiros'] - a['custo_frota'], 2)
        a['margem'] = (a['resultado'] / a['faturamento']) if a['faturamento'] else 0.0
        saida.append(a)
    saida.sort(key=lambda x: x['faturamento'], reverse=True)
    totais_custo = {'custo_cavalo': round(sum(a['custo_cavalo'] for a in saida), 2),
                    'custo_carreta': round(sum(a['custo_carreta'] for a in saida), 2),
                    'custo_frota': round(sum(a['custo_frota'] for a in saida), 2),
                    'frete_terceiros': round(sum(a['frete_terceiros'] for a in saida), 2),
                    'viagens_casadas': n_casados, 'viagens_total': n_total}
    return {'ok': True, 'dim': 'cliente', 'meses': meses_comp, 'tipos': tipos,
            'rows': saida, 'count': len(saida), 'custos_frota': False, 'custos_carreta': False,
            'custos_cliente': True, 'totais_custo': totais_custo}


@app.route('/api/veiculos/analise')
@page_required('veiculos')
def api_veiculos_analise():
    """Análise por veículo/pessoa sobre 'Auditoria Receita' (rateio por KM já pronto no dataset).
    Agrega Faturamento (receita_rateada) × Pagamento de frete (frete_motorista_total) + KM, no grão da
    dimensão escolhida. Fase 1 — sem custo de frota (na FROTA o pagamento é 0)."""
    from datetime import date, datetime

    # Eixo da análise. 'proprietario' usa placa_cavalo no DAX e resolve o dono em Python (cadastro normalizado).
    DIMS = {
        'cavalo':       "'Auditoria Receita'[placa_cavalo]",
        'carreta':      "'Auditoria Receita'[placa_carreta]",
        'motorista':    "'Auditoria Receita'[motorista]",
        'proprietario': "'Auditoria Receita'[placa_cavalo]",
    }
    dim = (request.args.get('dim') or 'cavalo').lower()
    if dim not in DIMS and dim != 'cliente':
        dim = 'cavalo'

    def _parse(s, default):
        try:
            return datetime.strptime(s, '%Y-%m-%d').date()
        except Exception:
            return default
    import re as _re_mes
    hoje = date.today()
    # Filtro por competência. 'meses' (lista YYYY-MM, multi-mês) é o preferido; aceita 'mes' único e de/ate (legado).
    meses_param = request.args.get('meses')
    mes = request.args.get('mes')
    meses_comp = []
    if meses_param:
        meses_comp = sorted({m.strip() for m in meses_param.split(',') if _re_mes.fullmatch(r'\d{4}-\d{2}', m.strip())})
    elif mes and _re_mes.fullmatch(r'\d{4}-\d{2}', mes):
        meses_comp = [mes]
    else:
        de = _parse(request.args.get('de'), date(hoje.year, hoje.month, 1))
        ate = _parse(request.args.get('ate'), hoje)
        _cy, _cm = de.year, de.month
        while (_cy, _cm) <= (ate.year, ate.month):
            meses_comp.append(f"{_cy:04d}-{_cm:02d}"); _cm += 1
            if _cm > 12: _cm = 1; _cy += 1
    if not meses_comp:
        meses_comp = [f"{hoje.year:04d}-{hoje.month:02d}"]
    meses_set = '{' + ','.join(f'"{m}"' for m in meses_comp) + '}'

    tipos = [t.strip().upper() for t in (request.args.get('tipos') or '').split(',') if t.strip()]
    tipos = [t for t in tipos if t in ('FROTA', 'AGREGADO', 'CARRETEIRO')] or ['FROTA', 'AGREGADO', 'CARRETEIRO']
    tipos_dax = '{' + ','.join(f'"{t}"' for t in tipos) + '}'

    # Recorte CLIENTE: margem por cliente (consolidado por raiz de CNPJ) — caminho dedicado.
    if dim == 'cliente':
        try:
            return jsonify(_veiculos_margem_cliente(get_token(), meses_comp, meses_set, tipos, tipos_dax))
        except requests.exceptions.HTTPError as e:
            try:
                detail = e.response.json()
            except Exception:
                detail = e.response.text
            return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    dim_expr = DIMS[dim]
    dax = (
        "EVALUATE SUMMARIZE(ADDCOLUMNS(FILTER('Auditoria Receita', "
        f"FORMAT('Auditoria Receita'[data_ref_ctrc], \"YYYY-MM\") IN {meses_set} && "
        f"'Auditoria Receita'[Tipo Operacao] IN {tipos_dax}), "
        f"\"dim\", {dim_expr}), "
        "[dim], 'Auditoria Receita'[Tipo Operacao], "
        "\"faturamento\", SUM('Auditoria Receita'[receita_rateada]), "
        "\"pagamento\", SUM('Auditoria Receita'[frete_motorista_total]), "
        # Vale-pedágio (repasse) e retenção (SEST/INSS/IRRF) p/ desdobrar o pagamento em colunas visuais.
        # frete_motorista_total = valor_a_pagar + vale_pedagio; retenção vem do ctrbs_oss (join 1:1 por CTRB).
        "\"vale_pedagio\", SUM('Auditoria Receita'[vale_pedagio]), "
        "\"retencao\", SUMX(VALUES('Auditoria Receita'[CTRB]), "
        "COALESCE(LOOKUPVALUE('public ctrbs_oss'[total_retencoes], 'public ctrbs_oss'[ctrb], 'Auditoria Receita'[CTRB]), 0)), "
        # KM igual ao BI: prefere rotas_km (distância de rota), cai no distancia_km cru só se não achar
        "\"km\", SUMX('Auditoria Receita', COALESCE(LOOKUPVALUE('public rotas_km'[km],"
        "'public rotas_km'[cidade_uf_origem],'Auditoria Receita'[cidade_uf_origem],"
        "'public rotas_km'[cidade_uf_destino],'Auditoria Receita'[cidade_uf_destino]),"
        "'Auditoria Receita'[distancia_km])), "
        "\"viagens\", DISTINCTCOUNT('Auditoria Receita'[CTRB]))"
    )

    try:
        token = get_token()
        result = execute_dax(token, dax)
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        linhas = clean_rows(rows)
    except requests.exceptions.HTTPError as e:
        try:
            detail = e.response.json()
        except Exception:
            detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

    # Unifica grafia antiga + Mercosul (rótulo Mercosul), mantendo o grão (dim, tipo).
    # Para 'proprietario', resolve o dono pelo cadastro normalizado (placa_cavalo → proprietário).
    cadastro = (_cadastro_veiculos(token)
                if dim in ('proprietario', 'carreta', 'cavalo')
                else {})
    agg = {}
    for r in linhas:
        raw = r.get('dim')
        # No recorte por carreta, truck/toco (veículo rígido) não tem carreta → fora da análise
        if dim == 'carreta' and (raw is None or str(raw).strip() == ''):
            continue
        if dim == 'proprietario':
            p = _placa_mercosul(str(raw)) if raw not in (None, '') else ''
            nome = cadastro.get(p, {}).get('proprietario') or '(sem identificação)'
        elif dim in ('cavalo', 'carreta'):
            nome = _placa_mercosul(str(raw)) if raw not in (None, '') else '(sem identificação)'
        else:  # motorista
            nome = str(raw) if raw not in (None, '') else '(sem identificação)'
        tipo = r.get('Tipo Operacao') or '—'
        # Regra interna: placa vendida (ainda no nome da Rizza) não conta como frota
        if dim == 'cavalo' and tipo == 'FROTA' and nome in PLACAS_VENDIDAS:
            continue
        key = (nome, tipo)
        a = agg.get(key)
        if a is None:
            a = {'dim': nome, 'tipo': tipo, 'faturamento': 0.0, 'frete_total': 0.0,
                 'vale_pedagio': 0.0, 'retencao': 0.0, 'km': 0.0, 'viagens': 0}
            agg[key] = a
        frota = (tipo == 'FROTA')
        a['faturamento'] += float(r.get('faturamento') or 0)
        # Frota = veículo próprio: não há pagamento de frete a terceiro (o frete_motorista aí é comissão do motorista próprio)
        # frete_total = pagamento de frete cheio (bruto + vale-pedágio) — base do Resultado (inalterado).
        a['frete_total'] += 0.0 if frota else float(r.get('pagamento') or 0)
        # Vale-pedágio (repasse) e retenção (SEST/INSS/IRRF): só agregado/carreteiro; exibidos separados do Pagamento.
        a['vale_pedagio'] += 0.0 if frota else float(r.get('vale_pedagio') or 0)
        a['retencao'] += 0.0 if frota else float(r.get('retencao') or 0)
        a['km'] += float(r.get('km') or 0)
        a['viagens'] += int(r.get('viagens') or 0)

    saida = []
    for a in agg.values():
        a['faturamento'] = round(a['faturamento'], 2)
        # Desdobra o frete cheio em Pagamento (líquido) + Retenção + Pedágio (somam de volta o frete_total).
        # Resultado usa o frete_total → permanece idêntico ao de antes (é só uma separação visual).
        frete_total = a.pop('frete_total')
        a['vale_pedagio'] = round(a['vale_pedagio'], 2)
        a['retencao'] = round(a['retencao'], 2)
        a['pagamento'] = round(max(frete_total - a['vale_pedagio'] - a['retencao'], 0.0), 2)
        a['km'] = round(a['km'], 1)
        a['resultado'] = round(a['faturamento'] - frete_total, 2)
        saida.append(a)
    saida.sort(key=lambda x: x['faturamento'], reverse=True)

    # ── Custos mensais rateados proporcional ao faturamento ──
    def _q(dax_q):
        res = execute_dax(token, dax_q)
        return clean_rows(res.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))

    def _q1(dax_q):
        r = _q(dax_q)
        return float((r[0] if r else {}).get('v') or 0)

    custos_frota = (dim == 'cavalo' and tipos == ['FROTA'])
    custos_carreta = (dim == 'carreta')
    totais_custo = {}

    DZ = "'public consulta_despesas_477'"
    anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'

    if custos_frota:
        # Custo real por cavalo via helper (mesmos valores de sempre: usa o faturamento do saida).
        fat_por_placa = {a['dim']: a['faturamento'] for a in saida}
        custos, totais_custo = _custo_frota_por_cavalo(token, meses_set, cadastro, fat_por_placa)
        for a in saida:
            c = custos.get(a['dim'], {})
            for k in _COST_KEYS:
                a[k] = c.get(k, 0.0)

    elif custos_carreta:
        # Manutenção carreta + pneu rateados entre as carretas Rizza (frota + agregado)
        rizza_carretas = {p for p, v in cadastro.items()
                          if v.get('proprietario') == 'RIZZA TRANSPORTES LTDA' and v.get('tipo') == 'CARRETA'}
        manut_carreta_total = _q1(f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento] IN {{\"5153\",\"5155\"}} && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
        pool_pneu, _, pool_pneu_car = _pool_pneu_split(token, meses_set, cadastro)
        # Base de rateio = faturamento de TODAS as carretas Rizza (frota+agregado) no mês,
        # independente do filtro de tipo da tela → taxa fixa por R$ de faturamento de carreta.
        AR = "'Auditoria Receita'"
        univ = _q(f"EVALUATE SUMMARIZE(FILTER({AR}, "
                  f"FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && "
                  f"{AR}[Tipo Operacao] IN {{\"FROTA\",\"AGREGADO\"}} && NOT(ISBLANK({AR}[placa_carreta]))), "
                  f"{AR}[placa_carreta], \"f\", SUM({AR}[receita_rateada]))")
        base_fat = 0.0
        for r in univ:
            if _placa_mercosul(r.get('placa_carreta')) in rizza_carretas:
                base_fat += float(r.get('f') or 0)
        taxa = manut_carreta_total / (base_fat or 1.0)
        taxa_pneu = pool_pneu_car / (base_fat or 1.0)  # base fixa = todas carretas Rizza
        for a in saida:
            a['rizza'] = a['dim'] in rizza_carretas
            a['manut_carreta'] = round(taxa * a['faturamento'], 2) if a['rizza'] else 0.0
            a['pneu'] = round(taxa_pneu * a['faturamento'], 2) if a['rizza'] else 0.0
        totais_custo = {'manut_carreta': round(manut_carreta_total, 2), 'base_fat_carretas': round(base_fat, 2),
                        'pneu': round(pool_pneu_car, 2), 'pneu_total': round(pool_pneu, 2)}

    # ── Proprietário do CAVALO no recorte cavalo (exceto a visão de custo da frota) ──
    # É o dono do próprio cavalo da linha (1:1 no cadastro); frota não traz.
    if dim == 'cavalo' and not custos_frota:
        for a in saida:
            a['prop_cavalo'] = '' if a['tipo'] == 'FROTA' else (cadastro.get(a['dim'], {}).get('proprietario') or '')

    return jsonify({'ok': True, 'dim': dim, 'meses': meses_comp,
                    'tipos': tipos, 'rows': saida, 'count': len(saida),
                    'custos_frota': custos_frota, 'custos_carreta': custos_carreta,
                    'custos_cliente': False, 'totais_custo': totais_custo})


_COST_MONEY = ('pedagio', 'combustivel', 'arla', 'pessoal', 'manut_cavalo', 'seguro', 'rastreador', 'pneu')
_COST_LABEL = {'pedagio': 'Pedágio', 'combustivel': 'Combustível', 'arla': 'ARLA', 'pessoal': 'Pessoal',
               'manut_cavalo': 'Manut. Cavalo', 'seguro': 'Seguro', 'rastreador': 'Rastreador', 'pneu': 'Pneu (cavalo)'}


def _veiculos_detalhe_cliente(token, valor, meses_comp, meses_set, tipos, tipos_dax):
    """Detalhe (drawer) de um cliente (raiz de CNPJ ou nome): KPIs, custo de frota por componente
    (cavalo + carreta) alocado, quebra por tipo, top rotas, cavalos que atenderam, evolução mensal e cargas."""
    AR = "'Auditoria Receita'"; DZ = "'public consulta_despesas_477'"
    anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'
    # `valor` é a raiz vinda da linha (já canônica); passa pelo mapa p/ aceitar também a raiz secundária
    chave_alvo = _raiz_cnpj(valor) or f"nome::{valor.strip().upper()}"

    # Participação do cliente em cada viagem vem da proporção dos CTRCs (igual à tabela) — não por cnpj inteiro
    ctrc_map = _ctrc_tomador_map(token, meses_comp)
    vrows = _dax_rows(token, f"EVALUATE SELECTCOLUMNS(FILTER({AR}, FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && {AR}[Tipo Operacao] IN {tipos_dax}), "
        f"\"data\",{AR}[data_ref_ctrc],\"mes\",FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\"),\"ctrb\",{AR}[CTRB],\"ctrc\",{AR}[CTRC],"
        f"\"origem\",{AR}[cidade_uf_origem],\"destino\",{AR}[cidade_uf_destino],\"tipo\",{AR}[Tipo Operacao],"
        f"\"receita\",{AR}[receita_rateada],\"frete\",{AR}[frete_motorista_total],\"km\",{AR}[distancia_km],"
        f"\"cavalo\",{AR}[placa_cavalo],\"carreta\",{AR}[placa_carreta],\"cnpj\",{AR}[cnpj_pagador],\"nome\",{AR}[cliente_pagador])")

    nome_disp = valor
    cargas = []
    for v in vrows:
        by_ch = {}
        for c in str(v.get('ctrc') or '').split(','):
            e = ctrc_map.get(_norm_manifesto(c))
            if e:
                by_ch[e['chave']] = by_ch.get(e['chave'], 0.0) + e['vf']
                if e['chave'] == chave_alvo and e.get('nome'):
                    nome_disp = e['nome']
        tot_vf = sum(by_ch.values())
        if tot_vf > 0:
            share = by_ch.get(chave_alvo, 0.0) / tot_vf
        else:  # viagem sem CTRC casado → cai no cnpj único da Auditoria
            rz = _raiz_cnpj(v.get('cnpj')) or f"nome::{(str(v.get('nome') or '').strip()).upper()}"
            share = 1.0 if rz == chave_alvo else 0.0
            if share and str(v.get('nome') or '').strip():
                nome_disp = str(v.get('nome')).strip()
        if share <= 0:
            continue
        cargas.append({'data': v.get('data'), 'mes': v.get('mes'), 'ctrb': v.get('ctrb'), 'ctrc': v.get('ctrc'),
                       'origem': v.get('origem'), 'destino': v.get('destino'), 'tipo': v.get('tipo'),
                       'receita': float(v.get('receita') or 0) * share, 'frete': float(v.get('frete') or 0) * share,
                       'km': float(v.get('km') or 0) * share, 'cavalo': v.get('cavalo'), 'carreta': v.get('carreta')})

    receita = sum(c['receita'] for c in cargas)
    frete_terceiros = sum(c['frete'] for c in cargas if (c.get('tipo') or '') != 'FROTA')
    km = sum(c['km'] for c in cargas)
    viagens = len({c.get('ctrb') for c in cargas if c.get('ctrb')})

    cadastro = _cadastro_veiculos(token)
    comp = {k: 0.0 for k in _COST_MONEY}
    cavalos_rows = []
    custo_cavalo = 0.0
    # Custo de cavalo diluído proporcional à RECEITA DE FROTA do cliente (taxa única da frota),
    # não por cavalo específico. Componentes seguem a mesma proporção.
    if 'FROTA' in tipos:
        custos_cav, _t = _custo_frota_por_cavalo(token, meses_set, cadastro)
        total_custo_cavalo = sum(c.get('custo_total', 0.0) for c in custos_cav.values())
        total_comp = {k: sum(c.get(k, 0.0) for c in custos_cav.values()) for k in _COST_MONEY}
        total_rec_frota = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({AR}, {AR}[Tipo Operacao]=\"FROTA\" "
            f"&& FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set}), {AR}[receita_rateada]))") or 1.0
        taxa_cav = total_custo_cavalo / total_rec_frota
        rec_cli_cav = {}
        for c in cargas:
            if (c.get('tipo') or '') == 'FROTA':
                p = _placa_mercosul(c.get('cavalo'))
                rec_cli_cav[p] = rec_cli_cav.get(p, 0.0) + float(c.get('receita') or 0)
        rec_cli_frota = sum(rec_cli_cav.values())
        for k in _COST_MONEY:
            comp[k] = total_comp[k] / total_rec_frota * rec_cli_frota
        for p, rc in rec_cli_cav.items():
            cavalos_rows.append({'placa': p, 'receita': round(rc, 2), 'custo': round(taxa_cav * rc, 2)})
        custo_cavalo = round(taxa_cav * rec_cli_frota, 2)

    # Custo da CARRETA (manut + pneu), taxa fixa × receita do cliente em carretas Rizza
    manut_carreta_aloc = pneu_carreta_aloc = 0.0
    rizza_carretas = set(); taxa_car = 0.0   # taxa única de custo de carreta por R$ de receita (manut+pneu)
    if 'FROTA' in tipos or 'AGREGADO' in tipos:
        rizza_carretas = {p for p, v in cadastro.items()
                          if v.get('proprietario') == 'RIZZA TRANSPORTES LTDA' and v.get('tipo') == 'CARRETA'}
        manut_carreta_total = _dax_val(token, f"EVALUATE ROW(\"v\", SUMX(FILTER({DZ}, {DZ}[evento] IN {{\"5153\",\"5155\"}} && {anomes} IN {meses_set}), {DZ}[vlr_final]))")
        _, _, pool_pneu_car = _pool_pneu_split(token, meses_set, cadastro)
        univ = _dax_rows(token, f"EVALUATE SUMMARIZE(FILTER({AR}, FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} "
            f"&& {AR}[Tipo Operacao] IN {{\"FROTA\",\"AGREGADO\"}} && NOT(ISBLANK({AR}[placa_carreta]))), {AR}[placa_carreta], \"f\", SUM({AR}[receita_rateada]))")
        base_fat = sum(float(r.get('f') or 0) for r in univ if _placa_mercosul(r.get('placa_carreta')) in rizza_carretas)
        rec_cli_carr = sum(float(c.get('receita') or 0) for c in cargas if _placa_mercosul(c.get('carreta')) in rizza_carretas)
        manut_carreta_aloc = round((manut_carreta_total / (base_fat or 1.0)) * rec_cli_carr, 2)
        pneu_carreta_aloc = round((pool_pneu_car / (base_fat or 1.0)) * rec_cli_carr, 2)
        taxa_car = (manut_carreta_total + pool_pneu_car) / (base_fat or 1.0)
    custo_carreta = round(manut_carreta_aloc + pneu_carreta_aloc, 2)

    resultado = round(receita - frete_terceiros - custo_cavalo - custo_carreta, 2)

    # Componentes (cavalo + carreta) p/ a seção de custo
    componentes = [{'nome': _COST_LABEL[k], 'valor': round(comp[k], 2)} for k in _COST_MONEY if round(comp[k], 2)]
    if manut_carreta_aloc:
        componentes.append({'nome': 'Manut. Carreta', 'valor': manut_carreta_aloc})
    if pneu_carreta_aloc:
        componentes.append({'nome': 'Pneu (carreta)', 'valor': pneu_carreta_aloc})

    # Quebra por tipo (mini-DRE): receita, frete, custo (cavalo+carreta), resultado, margem
    por_tipo = {}
    for c in cargas:
        t = c.get('tipo') or '—'
        a = por_tipo.setdefault(t, {'tipo': t, 'receita': 0.0, 'frete': 0.0, 'custo': 0.0, 'viagens': set()})
        a['receita'] += float(c.get('receita') or 0)
        if t != 'FROTA':
            a['frete'] += float(c.get('frete') or 0)
        if _placa_mercosul(c.get('carreta')) in rizza_carretas:   # custo de carreta (manut+pneu) alocado
            a['custo'] += taxa_car * float(c.get('receita') or 0)
        if c.get('ctrb'):
            a['viagens'].add(c.get('ctrb'))
    if 'FROTA' in por_tipo:   # custo de cavalo é 100% da frota
        por_tipo['FROTA']['custo'] += custo_cavalo
    por_tipo = [{'tipo': v['tipo'], 'receita': round(v['receita'], 2), 'frete': round(v['frete'], 2),
                 'custo': round(v['custo'], 2), 'resultado': round(v['receita'] - v['frete'] - v['custo'], 2),
                 'margem': ((v['receita'] - v['frete'] - v['custo']) / v['receita']) if v['receita'] else 0.0,
                 'viagens': len(v['viagens'])} for v in por_tipo.values()]
    por_tipo.sort(key=lambda x: -x['receita'])

    # Top rotas
    rotas = {}
    for c in cargas:
        k = f"{c.get('origem') or '—'} → {c.get('destino') or '—'}"
        a = rotas.setdefault(k, {'nome': k, 'receita': 0.0, 'viagens': set()})
        a['receita'] += float(c.get('receita') or 0)
        if c.get('ctrb'):
            a['viagens'].add(c.get('ctrb'))
    rotas = sorted(({'nome': v['nome'], 'receita': round(v['receita'], 2), 'viagens': len(v['viagens'])} for v in rotas.values()),
                   key=lambda x: -x['receita'])[:15]

    # Evolução mensal
    evol = {}
    for c in cargas:
        m = c.get('mes') or '—'
        a = evol.setdefault(m, {'mes': m, 'receita': 0.0, 'viagens': set()})
        a['receita'] += float(c.get('receita') or 0)
        if c.get('ctrb'):
            a['viagens'].add(c.get('ctrb'))
    evolucao = sorted(({'mes': v['mes'], 'receita': round(v['receita'], 2), 'viagens': len(v['viagens'])} for v in evol.values()),
                      key=lambda x: x['mes'])

    cargas_out = sorted(({'data': c.get('data'), 'ctrc': c.get('ctrc'), 'tipo': c.get('tipo'),
                          'origem': c.get('origem'), 'destino': c.get('destino'), 'cliente': nome_disp,
                          'receita': round(float(c.get('receita') or 0), 2), 'km': round(float(c.get('km') or 0), 1)}
                         for c in cargas), key=lambda x: str(x.get('data') or ''))

    return {'ok': True, 'dim': 'cliente', 'valor': valor, 'nome': nome_disp, 'meses': meses_comp,
            'kpis': {'receita': round(receita, 2), 'frete': round(frete_terceiros, 2),
                     'custo_cavalo': custo_cavalo, 'custo_carreta': custo_carreta,
                     'resultado': resultado, 'margem': (resultado / receita) if receita else 0.0,
                     'km': round(km, 1), 'viagens': viagens,
                     'ticket': round(receita / viagens, 2) if viagens else 0.0},
            'custo_componentes': componentes, 'por_tipo': por_tipo, 'rotas': rotas,
            'cavalos': sorted(cavalos_rows, key=lambda x: -x['receita']), 'evolucao': evolucao,
            'cargas': cargas_out}


@app.route('/api/veiculos/detalhe')
@page_required('veiculos')
def api_veiculos_detalhe():
    """Detalhe (painel lateral) de um veículo/pessoa: cargas, abastecimentos, pedágios,
    manutenção real e relacionamentos, cruzando Auditoria/ValeCard/SemParar/Despesas/cadastro."""
    import re as _re
    dim = (request.args.get('dim') or 'cavalo').lower()
    valor = (request.args.get('valor') or '').strip()
    meses_comp = sorted({m.strip() for m in (request.args.get('meses') or '').split(',') if _re.fullmatch(r'\d{4}-\d{2}', m.strip())})
    if not valor or not meses_comp:
        return jsonify({'ok': False, 'error': 'Informe valor e meses'}), 400
    meses_set = '{' + ','.join(f'"{m}"' for m in meses_comp) + '}'
    # mesmo filtro de tipo da tela, para o detalhe reconciliar com a linha clicada
    tipos = [t.strip().upper() for t in (request.args.get('tipos') or '').split(',') if t.strip()]
    tipos = [t for t in tipos if t in ('FROTA', 'AGREGADO', 'CARRETEIRO')] or ['FROTA', 'AGREGADO', 'CARRETEIRO']

    # Recorte CLIENTE: caminho dedicado (consolida por raiz de CNPJ)
    if dim == 'cliente':
        tipos_dax = '{' + ','.join(f'"{t}"' for t in tipos) + '}'
        try:
            return jsonify(_veiculos_detalhe_cliente(get_token(), valor, meses_comp, meses_set, tipos, tipos_dax))
        except requests.exceptions.HTTPError as e:
            try:
                detail = e.response.json()
            except Exception:
                detail = e.response.text
            return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)}), 500

    AR = "'Auditoria Receita'"; VC = "'public abastecimentos_valecard'"
    SP = "'public semparar_lancamentos'"; DZ = "'public consulta_despesas_477'"
    tipo_clause = f"{AR}[Tipo Operacao] IN {{{','.join(chr(34) + t + chr(34) for t in tipos)}}}"

    try:
        token = get_token()

        def _q(dax_q):
            res = execute_dax(token, dax_q)
            return clean_rows(res.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))

        def _set(vals):
            return '{' + ','.join(f'"{v}"' for v in vals) + '}'

        grafias = _placa_grafias(valor) if dim in ('cavalo', 'carreta') else [valor]
        out = {'ok': True, 'dim': dim, 'valor': valor, 'meses': meses_comp}

        # ── CARGAS (Auditoria) ──
        col_placa = '[placa_cavalo]' if dim in ('cavalo', 'proprietario') else '[placa_carreta]' if dim == 'carreta' else None
        if dim == 'motorista':
            filtro_aud = f"{AR}[motorista] = \"{valor}\""
        elif dim == 'proprietario':
            filtro_aud = None  # tratado abaixo (quebra por veículo)
        else:
            filtro_aud = f"{AR}{col_placa} IN {_set(grafias)}"

        cargas = []
        if filtro_aud:
            cargas = _q(
                f"EVALUATE SELECTCOLUMNS(FILTER({AR}, FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && {tipo_clause} && {filtro_aud}), "
                f"\"data\",{AR}[data_ref_ctrc],\"ctrc\",{AR}[CTRC],\"manifesto\",{AR}[Manifesto],"
                f"\"origem\",{AR}[cidade_uf_origem],\"destino\",{AR}[cidade_uf_destino],\"cliente\",{AR}[cliente_pagador],"
                f"\"receita\",{AR}[receita_rateada],\"frete\",{AR}[frete_motorista_total],\"km\",{AR}[distancia_km],"
                f"\"status\",{AR}[status_auditoria_frete],\"motorista\",{AR}[motorista],"
                f"\"cavalo\",{AR}[placa_cavalo],\"carreta\",{AR}[placa_carreta])")
            cargas.sort(key=lambda r: str(r.get('data') or ''))   # data crescente
            out['cargas'] = cargas
            out['kpis'] = {
                'receita': round(sum(float(c.get('receita') or 0) for c in cargas), 2),
                'frete': round(sum(float(c.get('frete') or 0) for c in cargas), 2),
                'km': round(sum(float(c.get('km') or 0) for c in cargas), 1),
                'viagens': len(cargas),
            }
            # Relacionamentos (agregados em Python sobre as cargas)
            def _agrupa(chave):
                acc = {}
                for c in cargas:
                    k = c.get(chave) or '—'
                    a = acc.setdefault(k, {'nome': k, 'viagens': 0, 'receita': 0.0})
                    a['viagens'] += 1; a['receita'] += float(c.get('receita') or 0)
                return sorted(acc.values(), key=lambda x: -x['receita'])
            rotas = {}
            for c in cargas:
                k = f"{c.get('origem') or '—'} → {c.get('destino') or '—'}"
                a = rotas.setdefault(k, {'nome': k, 'viagens': 0, 'receita': 0.0})
                a['viagens'] += 1; a['receita'] += float(c.get('receita') or 0)
            carretas_rel = _agrupa('carreta') if dim != 'carreta' else []
            cavalos_rel = _agrupa('cavalo') if dim == 'carreta' else []
            # Proprietário (2º nível): resolve placa → dono pelo cadastro.
            # No recorte carreta é o dono dos CAVALOS que puxaram (pode ser vários).
            if carretas_rel or cavalos_rel:
                cadr = _cadastro_veiculos(token)
                for c in carretas_rel + cavalos_rel:
                    c['prop'] = cadr.get(_placa_mercosul(c['nome']), {}).get('proprietario') or ''
            out['relacionamentos'] = {
                'rotas': sorted(rotas.values(), key=lambda x: -x['receita'])[:15],
                'motoristas': _agrupa('motorista'),
                'carretas': carretas_rel,
                'cavalos': cavalos_rel,
            }

        # ── ABASTECIMENTOS (ValeCard) — cavalo e motorista ──
        if dim in ('cavalo', 'motorista'):
            filtro_vc = (f"{VC}[placa] IN {_set(grafias)}" if dim == 'cavalo'
                         else f"{VC}[motorista] = \"{valor}\"")
            ab = _q(
                f"EVALUATE SELECTCOLUMNS(FILTER({VC}, FORMAT({VC}[dch_data],\"YYYY-MM\") IN {meses_set} && {filtro_vc}), "
                f"\"data\",{VC}[dch_data],\"posto\",{VC}[estabelecimento],\"cidade\",{VC}[cidade],\"uf\",{VC}[uf],"
                f"\"produto\",{VC}[produto],\"litros\",{VC}[ncd_quantidade],\"vunit\",{VC}[mcd_valor_unitario],"
                f"\"valor\",{VC}[mcd_valor_total],\"hodometro\",{VC}[nsd_hodometro],\"motorista\",{VC}[motorista])")
            ab.sort(key=lambda r: (str(r.get('data') or ''), float(r.get('hodometro') or 0)))
            out['abastecimentos'] = ab
            # consumo real pelo km do hodômetro (mesmo método do 1º nível)
            litros_diesel = sum(float(r.get('litros') or 0) for r in ab if 'ARLA' not in str(r.get('produto') or '').upper())
            km_hod = _km_hodometro([float(r.get('hodometro') or 0) for r in ab])
            out['consumo'] = {
                'km_hodometro': round(km_hod, 0),
                'litros_diesel': round(litros_diesel, 1),
                'km_por_litro': round(km_hod / litros_diesel, 2) if litros_diesel and km_hod else 0,
                'gasto': round(sum(float(r.get('valor') or 0) for r in ab), 2),
            }

        # ── PEDÁGIOS (Sem Parar) — cavalo ──
        if dim == 'cavalo':
            ped = _q(
                f"EVALUATE SELECTCOLUMNS(FILTER({SP}, (RIGHT({SP}[data],4) & \"-\" & MID({SP}[data],4,2)) IN {meses_set} && {SP}[placa_veiculo] IN {_set(grafias)}), "
                f"\"data\",{SP}[data],\"hora\",{SP}[horario],\"sentido\",{SP}[sentido_praca],\"tipo\",{SP}[tipo_uso],"
                f"\"valor\",{SP}[valor],\"dc\",{SP}[debito_credito],\"embarcador\",{SP}[embarcador])")
            ped.sort(key=lambda r: str(r.get('data') or ''))
            out['pedagios'] = ped
            out['pedagio_total'] = round(sum(float(r.get('valor') or 0) for r in ped), 2)

        # ── MANUTENÇÃO real (despesas via placa no histórico) — cavalo e carreta ──
        if dim in ('cavalo', 'carreta'):
            eventos = '{"5150","5154"}' if dim == 'cavalo' else '{"5153","5155"}'
            anomes = f'("20" & RIGHT({DZ}[mes_competencia],2) & "-" & LEFT({DZ}[mes_competencia],2))'
            buscas = ' || '.join(f"SEARCH(\"{g}\",{DZ}[historico_despesa],1,0)>0" for g in grafias)
            man = _q(
                f"EVALUATE SELECTCOLUMNS(FILTER({DZ}, {DZ}[evento] IN {eventos} && {anomes} IN {meses_set} && ({buscas})), "
                f"\"emissao\",{DZ}[emissao],\"fornecedor\",{DZ}[nome_fornecedor],\"descricao\",{DZ}[historico_despesa],"
                f"\"valor\",{DZ}[vlr_final],\"evento\",{DZ}[evento])")
            man.sort(key=lambda r: str(r.get('emissao') or ''))
            out['manutencao'] = man
            out['manutencao_total'] = round(sum(float(r.get('valor') or 0) for r in man), 2)

        # ── PERFIL (cadastro) — placa dims ──
        if dim in ('cavalo', 'carreta'):
            cad = _cadastro_veiculos(token)
            out['perfil'] = cad.get(valor) or {}

        # ── QUEBRA POR VEÍCULO — motorista e proprietário ──
        if dim in ('motorista', 'proprietario'):
            def _por_placa(campo, extra):
                return _q(f"EVALUATE SUMMARIZE(FILTER({AR}, FORMAT({AR}[data_ref_ctrc],\"YYYY-MM\") IN {meses_set} && {tipo_clause} && {extra}), "
                          f"{AR}{campo}, \"rec\",SUM({AR}[receita_rateada]),\"km\",SUM({AR}[distancia_km]),\"v\",DISTINCTCOUNT({AR}[CTRB]))")
            if dim == 'motorista':
                veic = []
                for campo, tp in (('[placa_cavalo]', 'CAVALO'), ('[placa_carreta]', 'CARRETA')):
                    for r in _por_placa(campo, f"{AR}[motorista] = \"{valor}\""):
                        raw = r.get('placa_cavalo') or r.get('placa_carreta')
                        if not raw:
                            continue
                        veic.append({'placa': _placa_mercosul(raw), 'tipo': tp,
                                     'receita': round(float(r.get('rec') or 0), 2), 'km': round(float(r.get('km') or 0), 1),
                                     'viagens': int(r.get('v') or 0)})
                out['veiculos'] = sorted(veic, key=lambda x: -x['receita'])
            else:  # proprietario — base CAVALO (igual à tabela, que resolve o dono pelo cavalo); sem dupla contagem cavalo+carreta
                cad = _cadastro_veiculos(token)
                donos = {p: v for p, v in cad.items() if (v.get('proprietario') or '') == valor}
                ativ_cav, ativ_car = {}, {}
                for campo, dest in (('[placa_cavalo]', ativ_cav), ('[placa_carreta]', ativ_car)):
                    for r in _por_placa(campo, f"NOT(ISBLANK({AR}{campo}))"):
                        raw = r.get('placa_cavalo') or r.get('placa_carreta')
                        if not raw:
                            continue
                        dest[_placa_mercosul(raw)] = {'rec': float(r.get('rec') or 0), 'km': float(r.get('km') or 0), 'v': int(r.get('v') or 0)}
                veic = []
                for p, v in donos.items():
                    # carreta usa atividade como carreta; cavalo/cavalo trucado/truck usam atividade como placa_cavalo
                    src = ativ_car if v.get('tipo') == 'CARRETA' else ativ_cav
                    a = src.get(p, {})
                    veic.append({'placa': p, 'tipo': v.get('tipo'), 'modelo': v.get('modelo'), 'disponivel': v.get('disponivel'),
                                 'receita': round(a.get('rec', 0.0), 2), 'km': round(a.get('km', 0.0), 1), 'viagens': int(a.get('v', 0))})
                out['veiculos'] = sorted(veic, key=lambda x: -x['receita'])
                # headline = soma dos veículos-tração do dono (cavalo/truck), igual à tabela que resolve pelo placa_cavalo
                cav = [p for p, v in donos.items() if v.get('tipo') != 'CARRETA']
                out['kpis'] = {'receita': round(sum(ativ_cav.get(p, {}).get('rec', 0) for p in cav), 2),
                               'km': round(sum(ativ_cav.get(p, {}).get('km', 0) for p in cav), 1),
                               'viagens': sum(ativ_cav.get(p, {}).get('v', 0) for p in cav),
                               'frota': len(veic)}

        return jsonify(out)

    except requests.exceptions.HTTPError as e:
        try:
            detail = e.response.json()
        except Exception:
            detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _gerar_csv(cols, data):
    import csv
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=cols, delimiter=';')
    writer.writeheader()
    for row in data:
        writer.writerow({c: ('' if row.get(c) is None else row.get(c)) for c in cols})
    csv_bytes = '﻿'.encode('utf-8') + buf.getvalue().encode('utf-8')
    return io.BytesIO(csv_bytes)


def _csv_linha(valores):
    """Formata lista de valores em uma linha CSV com separador ; e BOM-safe."""
    out = []
    for v in valores:
        if v is None:
            out.append('')
        else:
            s = str(v).replace('"', '""')
            if ';' in s or '"' in s or '\n' in s:
                s = f'"{s}"'
            out.append(s)
    return ';'.join(out) + '\n'


@app.route('/api/dre/despesas/csv')
@page_required('despesas')
def api_dre_despesas_csv():
    from datetime import datetime
    start = request.args.get('start')
    end = request.args.get('end')
    meses_param = request.args.get('meses')
    grupo = request.args.get('grupo')
    evento = request.args.get('evento')

    if meses_param:
        pares = _parse_meses_param(meses_param)
        if not pares:
            return jsonify({'ok': False, 'error': 'Parâmetro meses inválido'}), 400
        meses = _meses_para_periodos(pares)
        sorted_meses = sorted(pares)
        start = f"{sorted_meses[0][0]}-{sorted_meses[0][1]:02d}-01"
        end = f"{sorted_meses[-1][0]}-{sorted_meses[-1][1]:02d}"
    elif start and end:
        start_d = datetime.strptime(start, '%Y-%m-%d').date()
        end_d   = datetime.strptime(end, '%Y-%m-%d').date()
        meses = list(_iterar_meses(start_d, end_d))
    else:
        return jsonify({'ok': False, 'error': 'Informe meses ou start/end'}), 400

    sufixo = ('_' + evento.replace(' ', '_')[:30]) if evento else (('_' + grupo) if grupo else '')
    nome = f"despesas_{start}_{end}{sufixo}.csv"

    def gerar():
        yield '﻿'  # BOM para Excel reconhecer UTF-8
        token = get_token()
        cols = None
        for (y, m, _, _, _) in meses:
            ref = f"{y:04d}/{m:02d}"
            dax = f'EVALUATE FILTER(\'public consulta_despesas_477\', \'public consulta_despesas_477\'[REF] = "{ref}"'
            if evento:
                dax += f' && \'public consulta_despesas_477\'[descr_evento] = "{evento}"'
            elif grupo:
                eventos = [e for e, (g, _) in MAPA_DRE.items() if g == grupo]
                if eventos:
                    lista = '{ ' + ', '.join(f'"{e}"' for e in eventos) + ' }'
                    dax += f' && \'public consulta_despesas_477\'[descr_evento] IN {lista}'
            dax += ')'

            try:
                result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
            except Exception:
                # Renovar token e tentar de novo
                token = get_token()
                result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])

            rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
            data = clean_rows(rows)
            if not cols and data:
                cols = list(data[0].keys())
                yield _csv_linha(cols)
            for row in data:
                yield _csv_linha([row.get(c) for c in cols] if cols else [])

    return Response(
        stream_with_context(gerar()),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{nome}"'}
    )


@app.route('/api/dre/conhecimentos/csv')
@page_required('conhecimentos')
def api_dre_conhecimentos_csv():
    from datetime import datetime
    start = request.args.get('start')
    end = request.args.get('end')
    meses_param = request.args.get('meses')

    if meses_param:
        pares = _parse_meses_param(meses_param)
        if not pares:
            return jsonify({'ok': False, 'error': 'Parâmetro meses inválido'}), 400
        meses = _meses_para_periodos(pares)
        sorted_meses = sorted(pares)
        start = f"{sorted_meses[0][0]}-{sorted_meses[0][1]:02d}-01"
        end = f"{sorted_meses[-1][0]}-{sorted_meses[-1][1]:02d}"
    elif start and end:
        start_d = datetime.strptime(start, '%Y-%m-%d').date()
        end_d   = datetime.strptime(end, '%Y-%m-%d').date()
        meses = list(_iterar_meses(start_d, end_d))
    else:
        return jsonify({'ok': False, 'error': 'Informe meses ou start/end'}), 400

    nome = f"conhecimentos_{start}_{end}.csv"

    def gerar():
        yield '﻿'
        token = get_token()
        cols = None
        for (_, _, _, prim, ult) in meses:
            dax = (
                f'EVALUATE FILTER(\'public conhecimentos_emitidos\', '
                f'\'public conhecimentos_emitidos\'[data_autorizacao] >= {_dax_data(prim)} && '
                f'\'public conhecimentos_emitidos\'[data_autorizacao] <= {_dax_data(ult)})'
            )
            try:
                result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
            except Exception:
                token = get_token()
                result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])

            rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
            data = clean_rows(rows)
            if not cols and data:
                cols = list(data[0].keys())
                yield _csv_linha(cols)
            for row in data:
                yield _csv_linha([row.get(c) for c in cols] if cols else [])

    return Response(
        stream_with_context(gerar()),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{nome}"'}
    )


# ════════════════════════════════════════════════════════════════════════════
# CONTÁBIL — dataset tabelas.contabil (456 extrato · 441 faturamento ·
#            571 ACNI · 479 eventos)
#
# Fonte das medidas: Rizza/contabil_pbi.py (dict MEDIDAS), validado por
# Rizza/contabil_testes.py. Não reescrever DAX aqui de cabeça — copiar de lá.
# ════════════════════════════════════════════════════════════════════════════

_T456     = "'public extrato_bancario_456'"
_T456_TOT = "'public extrato_bancario_456_totais'"
_T441     = "'public faturas_441'"
_T441_CTR = "'public faturas_441_ctrcs'"
_T571     = "'public acni_571'"
_T479     = "'public eventos_479'"

# Conta contábil de cada conta do 456, chaveada por (banco, agencia, conta).
# As nove primeiras casam pelo número da conta no plano da PERSETO; as três
# marcadas saem de 1.1.1.02 porque NÃO são disponibilidade — foi deduzido do
# comportamento do movimento, não lido de cadastro, então está comentado:
#   BB GARANTIDA  — 6 lançamentos, todos saque do limite p/ a conta corrente do
#                   BB. É conta garantida = empréstimo, logo passivo.
#   D.D SOLAR     — 210 lançamentos transferindo o produto do desconto p/ o
#                   Bradesco. DD = duplicatas descontadas, conta redutora.
#   TRIBANCO      — conta bancária de verdade (entra por FAT/ACN, sai por
#                   transferência) que não existe no plano. Precisa ser criada.
#   CAIXA PAMBANK — instituição de pagamento: 4.418 lançamentos CPG pagando
#                   frete/pedágio contra 634 MAN de abastecimento. Não é banco.
def _norm_conta(v):
    """'0001067129' e '1067129' são a mesma conta: o SSW preenche com zero à
    esquerda em um lugar e sem em outro."""
    return str(v or '').strip().lstrip('0') or '0'


def chave_banco(banco, agencia, conta):
    return f'BANCO:{_norm_conta(banco)}/{_norm_conta(agencia)}/{_norm_conta(conta)}'


# As contas fixas moram em `contabil_conta_fixa`, editável pela tela — abrir
# conta bancária não pode exigir deploy, e duas (TRIBANCO, CAIXA PAMBANK) já
# nasceram pendentes de criação no plano. Cache curto porque _quadro_bancos
# consulta 13 vezes por requisição e a tabela muda uma vez por ano.
_cache_fixas = {'valor': None, 'expira': 0}


def contas_fixas(forcar=False):
    """{chave: (classificacao, codigo_reduzido, descricao, observacao)}"""
    if not forcar and _cache_fixas['valor'] and time.time() < _cache_fixas['expira']:
        return _cache_fixas['valor']
    conn = get_db()
    cur = conn.cursor()
    cur.execute("""
        SELECT f.chave, f.classificacao, p.codigo_reduzido, f.descricao, f.observacao
          FROM contabil_conta_fixa f
          LEFT JOIN contabil_plano_contas p ON p.classificacao = f.classificacao
    """)
    d = {r[0]: (r[1], r[2], r[3], r[4]) for r in cur.fetchall()}
    cur.close(); conn.close()
    _cache_fixas['valor'] = d
    _cache_fixas['expira'] = time.time() + 60
    return d


def _conta_contabil_456(banco, agencia, conta):
    """(classificacao, codigo_reduzido, observacao) da conta bancária."""
    cls, cod, desc, obs = contas_fixas().get(
        chave_banco(banco, agencia, conta), (None, None, '', ''))
    return cls, cod, (obs or desc or '')


def contabil_dax(dax):
    """Executa DAX no dataset contábil e devolve linhas já sem o prefixo da tabela.

    O executeQueries corta a resposta e devolve HTTP 200 assim mesmo. São dois
    tetos diferentes: 100.000 linhas e **15 MiB de payload** — este último morde
    muito antes, porque depende da largura da linha. Medido: a mesma consulta
    devolve 10.127 linhas com 29 colunas e as 22.835 completas com 3.

    Resultado cortado é pior que erro, porque parece resposta — foi assim que a
    ponte 456×477 mediu 8,67% de cobertura quando a real é 100%.

    A boa notícia é que o serviço avisa: `results[0]['error']` só existe quando
    houve corte (`DaxByteCountNotSupported`). É essa a guarda; a contagem de
    linhas fica como cinto de segurança."""
    token = get_token()
    result = execute_dax(token, dax, dataset_id=CONFIG['contabil_dataset_id'])
    bloco = (result.get('results') or [{}])[0]
    erro = bloco.get('error')
    if erro:
        raise RuntimeError(
            f"Consulta cortada pelo Power BI ({erro.get('code')}): {erro.get('message')} "
            f"— reduza as colunas ou fatie o período.")
    rows = (bloco.get('tables') or [{}])[0].get('rows', [])
    if len(rows) >= 100000:
        raise RuntimeError(
            f'Consulta devolveu {len(rows):,} linhas — teto do executeQueries. '
            f'O resultado está cortado; agregue no DAX ou fatie a consulta.')
    return clean_rows(rows)


# Colunas do extrato que interessam ao fechamento. Selecionar explicitamente não
# é só higiene: puxar a tabela inteira (29 colunas) estoura os 15 MiB e o
# serviço devolve meia resposta. Ficam de fora id, RowNumber, ano/ordem e os
# carimbos de carga, que não dizem nada para quem concilia.
_COLS_EXTRATO = [
    ('conta_apelido', 'conta_apelido'), ('data', 'data'), ('orig', 'orig'),
    ('nlanca', 'nlanca'), ('documento', 'documento'), ('cnpj', 'cnpj'),
    ('cliente_fornecedor', 'cliente_fornecedor'), ('historico', 'historico'),
    ('valor', 'valor'), ('sit', 'sit'), ('saldo', 'saldo'),
    ('transferencia_interna', 'transferencia_interna'),
    ('ref_477_uni', 'ref_477_uni'), ('ref_477_numlancto', 'ref_477_numlancto'),
    ('ref_477_parcela', 'ref_477_parcela'),
]


def _select_extrato(filtro):
    """SELECTCOLUMNS do extrato já com a regra de classificação como coluna."""
    cols = ', '.join(f'"{apelido}", {_T456}[{col}]' for apelido, col in _COLS_EXTRATO)
    return (f'EVALUATE SELECTCOLUMNS(ADDCOLUMNS(FILTER({_T456}, {filtro}), '
            f'"regra", {_SWITCH_REGRA_456}), {cols}, "regra", [regra])')


# Classificação do movimento do 456 conforme a especificação da contadora
# (PARA GABRIEL.xlsx). A ordem do SWITCH importa: "VIA RET BCO" é testado sem
# amarrar à origem, porque o mesmo histórico aparece em BCO **e** em MAN — a
# regra dela ficou presa a BCO e deixava 13 linhas (R$ 2,7 mi) sem classificação.
_SWITCH_REGRA_456 = f"""SWITCH(TRUE(),
    {_T456}[transferencia_interna] = TRUE(), "TRANSFERENCIA",
    {_T456}[orig] = "CPG" && NOT ISBLANK({_T456}[ref_477_numlancto]), "477 (evento)",
    {_T456}[orig] = "FAT", "CLIENTE",
    {_T456}[orig] = "ACN", "ADTO CLIENTE",
    SEARCH("VIA RET BCO", {_T456}[historico], 1, 0) > 0, "RET BCO",
    LEFT({_T456}[historico], 4) = "FDBI", "DESC DUPLICATA",
    LEFT({_T456}[historico], 7) = "ESTORNO", "ESTORNO",
    "SEM REGRA")"""


def _quadro_bancos():
    """Uma linha por banco — a tela que a contadora desenhou.

    Vem de extrato_bancario_456_totais, que é o rodapé do próprio extrato do SSW
    gravado na carga. Não é soma nossa: por isso a coluna 'confere'."""
    tot = contabil_dax(
        f'EVALUATE SELECTCOLUMNS({_T456_TOT}, '
        f'"banco", [banco], "agencia", [agencia], "conta", [conta], '
        f'"conta_apelido", [conta_apelido], "banco_nome", [banco_nome], '
        f'"periodo_ini", [periodo_ini], "periodo_fim", [periodo_fim], '
        f'"saldo_inicial", [saldo_inicial], "total_creditos", [total_creditos], '
        f'"total_debitos", [total_debitos], "saldo_final", [saldo_final], '
        f'"movimentos", [movimentos])')

    # Conferência: a contagem carregada bate com a do rodapé?
    carregado = {r['conta_apelido']: r['n'] for r in contabil_dax(
        f'EVALUATE SUMMARIZECOLUMNS({_T456}[conta_apelido], "n", COUNTROWS({_T456}))')}

    for r in tot:
        cls, cod, obs = _conta_contabil_456(r['banco'], r['agencia'], r['conta'])
        r['conta_contabil'] = cls
        r['codigo_reduzido'] = cod
        r['observacao'] = obs
        r['confere'] = carregado.get(r['conta_apelido']) == r['movimentos']
    return tot


def _classificacao_456(mes=None):
    """Quantos movimentos ainda não têm regra de classificação, por banco."""
    filtro = f'{_T456}[realizado] = TRUE()'
    if mes:
        filtro += f' && FORMAT({_T456}[data], "YYYY-MM") = "{mes}"'
    return contabil_dax(
        f'EVALUATE SUMMARIZE(ADDCOLUMNS(FILTER({_T456}, {filtro}), '
        f'"@regra", {_SWITCH_REGRA_456}), '
        f'{_T456}[conta_apelido], [@regra], '
        f'"n", COUNTROWS({_T456}), "valor", SUM({_T456}[valor]))')


def _movimento_por_mes():
    """Agregado (conta × mês) do 456 — 13 × nº de meses, algumas dezenas de linhas.

    Traz o total COM programado e o total só do realizado. O primeiro existe
    porque é ele que reproduz o rodapé do SSW (conferido: o crédito do rodapé é
    o de tudo, não o do realizado); o segundo é o que a contabilidade usa."""
    return contabil_dax(f"""EVALUATE SUMMARIZE(
        ADDCOLUMNS({_T456}, "@mes", FORMAT({_T456}[data], "YYYY-MM")),
        {_T456}[conta_apelido], [@mes],
        "n", COUNTROWS({_T456}),
        "cred", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] > 0),
        "deb", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] < 0),
        "n_real", CALCULATE(COUNTROWS({_T456}), {_T456}[realizado] = TRUE()),
        "cred_real", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] > 0, {_T456}[realizado] = TRUE()),
        "deb_real", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] < 0, {_T456}[realizado] = TRUE()),
        "prog", CALCULATE(SUM({_T456}[valor]), {_T456}[realizado] = FALSE()))""")


def _saldos_por_mes(bancos):
    """Encadeia o saldo mês a mês a partir do saldo inicial do rodapé.

    Não precisa reextrair o SSW por mês: a identidade
    `saldo_inicial + créditos + débitos = saldo_final` foi conferida e fecha ao
    centavo nas 13 contas. Encadeando os meses, o fechamento de cada mês fica
    derivado — e o último mês tem que reproduzir o saldo_final do rodapé, que é
    a prova de que a corrente inteira está certa.
    """
    mov = _movimento_por_mes()
    por_conta = {}
    for m in mov:
        por_conta.setdefault(m['conta_apelido'], []).append(m)

    ini_rodape = {b['conta_apelido']: (b['saldo_inicial'] or 0) for b in bancos}
    fim_rodape = {b['conta_apelido']: (b['saldo_final'] or 0) for b in bancos}

    # Todo mês entra para TODA conta, mesmo sem movimento. Sem isso a conta some
    # do consolidado no mês parado e o total de abertura de um mês deixa de ser
    # o total de fechamento do anterior — o saldo dela continua existindo.
    meses = sorted({m['@mes'] for m in mov})

    saldos, prova = {}, {}
    for conta in ini_rodape:
        do_mes = {l['@mes']: l for l in por_conta.get(conta, [])}
        corrente = ini_rodape.get(conta, 0)
        for mes in meses:
            l = do_mes.get(mes)
            abertura = corrente
            corrente = abertura + ((l['cred'] or 0) + (l['deb'] or 0) if l else 0)
            saldos[(conta, mes)] = {
                'saldo_inicial': abertura, 'saldo_final': corrente,
                'total_creditos': (l['cred'] or 0) if l else 0,
                'total_debitos': -(l['deb'] or 0) if l else 0,
                'movimentos': (l['n'] or 0) if l else 0,
                'creditos_realizado': (l['cred_real'] or 0) if l else 0,
                'debitos_realizado': -(l['deb_real'] or 0) if l else 0,
                'programados_valor': (l['prog'] or 0) if l else 0,
                'movimentos_realizado': (l['n_real'] or 0) if l else 0,
            }
        # A corrente fechou onde o rodapé do SSW disse que fecharia?
        prova[conta] = abs(corrente - fim_rodape.get(conta, 0)) < 0.01
    return saldos, meses, prova


@app.route('/api/contabil/quadro')
@page_required('contabil')
def api_contabil_quadro():
    """Quadro por banco. Sem `mes` = a janela inteira que o robô extraiu (os
    números são o rodapé do SSW). Com `mes=YYYY-MM` = o fechamento daquele mês,
    derivado dos movimentos e provado contra o rodapé."""
    mes = (request.args.get('mes') or '').strip() or None
    try:
        bancos = _quadro_bancos()
        saldos, meses, prova = _saldos_por_mes(bancos)

        if mes:
            for b in bancos:
                s = saldos.get((b['conta_apelido'], mes))
                b.update(s or {'saldo_inicial': None, 'saldo_final': None,
                               'total_creditos': 0, 'total_debitos': 0, 'movimentos': 0})
                # No mês o saldo é derivado, não é o rodapé impresso. O que
                # sustenta o número é a corrente fechar no fim da janela.
                b['confere'] = prova.get(b['conta_apelido'], False)
                b['origem_saldo'] = 'corrente conferida contra o rodapé'
        else:
            for b in bancos:
                b['origem_saldo'] = 'rodapé do extrato do SSW'

        regras = _classificacao_456(mes)

        # Guardas: transferência entre contas próprias aparece 2× (crédito no
        # destino, débito na origem) e programado é previsão, não extrato.
        fm = f', FORMAT({_T456}[data], "YYYY-MM") = "{mes}"' if mes else ''
        m = contabil_dax(f"""EVALUATE ROW(
            "movimentos", CALCULATE(COUNTROWS({_T456}){fm}),
            "programados", CALCULATE(COUNTROWS({_T456}), {_T456}[realizado] = FALSE(){fm}),
            "transf_linhas", CALCULATE(COUNTROWS({_T456}), {_T456}[transferencia_interna] = TRUE(){fm}),
            "transf_valor", CALCULATE(SUMX({_T456}, ABS({_T456}[valor])), {_T456}[transferencia_interna] = TRUE(){fm}),
            "creditos", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] > 0,
                        {_T456}[transferencia_interna] = FALSE(), {_T456}[realizado] = TRUE(){fm}),
            "creditos_brutos", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] > 0{fm}),
            "debitos", CALCULATE(SUM({_T456}[valor]), {_T456}[valor] < 0,
                       {_T456}[transferencia_interna] = FALSE(), {_T456}[realizado] = TRUE(){fm}),
            "programados_valor", CALCULATE(SUM({_T456}[valor]), {_T456}[realizado] = FALSE(){fm}))""")[0]

        sem_regra = [r for r in regras if r.get('@regra') == 'SEM REGRA']
        m['sem_regra_linhas'] = sum(r['n'] or 0 for r in sem_regra)
        m['sem_regra_valor'] = sum(r['valor'] or 0 for r in sem_regra)
        # O saldo_final do rodapé é o número do SSW e **inclui os programados**
        # (conferido: o total de crédito do rodapé é o de tudo, não o do
        # realizado). Como os demais cards aplicam a guarda, mostrar só ele
        # deixaria dois números de bases diferentes lado a lado. Vão os dois.
        m['saldo_consolidado'] = sum(b['saldo_final'] or 0 for b in bancos)
        m['saldo_realizado'] = m['saldo_consolidado'] - (m.get('programados_valor') or 0)
        m['sem_conta_contabil'] = sum(1 for b in bancos if not b['conta_contabil'])

        return jsonify({'ok': True, 'bancos': bancos, 'regras': regras, 'totais': m,
                        'mes': mes, 'meses': meses,
                        'prova_corrente': all(prova.values()),
                        'contas_provadas': sum(1 for v in prova.values() if v),
                        'contas_total': len(prova)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/banco')
@page_required('contabil')
def api_contabil_banco():
    """Detalhe de uma conta (drawer): quebra por origem, por regra e maiores movimentos."""
    apelido = (request.args.get('conta') or '').strip()
    if not apelido:
        return jsonify({'ok': False, 'error': 'Informe conta'}), 400
    alvo = apelido.replace('"', '""')
    try:
        filtro = f'FILTER({_T456}, {_T456}[conta_apelido] = "{alvo}")'
        origem = contabil_dax(
            f'EVALUATE SUMMARIZECOLUMNS({_T456}[orig], {filtro}, '
            f'"n", COUNTROWS({_T456}), "valor", SUM({_T456}[valor]))')
        regra = contabil_dax(
            f'EVALUATE SUMMARIZE(ADDCOLUMNS(FILTER({filtro}, {_T456}[realizado] = TRUE()), '
            f'"@regra", {_SWITCH_REGRA_456}), [@regra], '
            f'"n", COUNTROWS({_T456}), "valor", SUM({_T456}[valor]))')
        maiores = contabil_dax(
            f'EVALUATE TOPN(12, SELECTCOLUMNS({filtro}, "data", [data], "orig", [orig], '
            f'"documento", [documento], "cliente_fornecedor", [cliente_fornecedor], '
            f'"historico", [historico], "valor", [valor], "sit", [sit]), ABS([valor]), DESC)')
        return jsonify({'ok': True, 'origem': origem, 'regra': regra, 'maiores': maiores})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _filtro_periodo_456(start, end, banco=None, transferencias=False, programados=False):
    partes = [f'{_T456}[data] >= {_dax_data_str(start)}', f'{_T456}[data] <= {_dax_data_str(end)}']
    if banco:
        partes.append(f'{_T456}[conta_apelido] = "{banco.replace(chr(34), chr(34) * 2)}"')
    if not transferencias:
        partes.append(f'{_T456}[transferencia_interna] = FALSE()')
    if not programados:
        partes.append(f'{_T456}[realizado] = TRUE()')
    return ' && '.join(partes)


def _dax_data_str(s):
    """'2026-01-31' → DATE(2026,1,31)."""
    y, m, d = str(s).split('-')
    return f'DATE({int(y)},{int(m)},{int(d)})'


@app.route('/api/contabil/extrato')
@page_required('contabil')
def api_contabil_extrato():
    start, end = request.args.get('start'), request.args.get('end')
    if not start or not end:
        return jsonify({'ok': False, 'error': 'Informe start e end (YYYY-MM-DD)'}), 400
    banco = request.args.get('banco') or None
    transf = request.args.get('transferencias') == '1'
    prog = request.args.get('programados') == '1'
    try:
        filtro = _filtro_periodo_456(start, end, banco, transf, prog)
        data = contabil_dax(_select_extrato(filtro))
        # Conferência de integridade: o que veio tem que bater com o que o
        # próprio modelo conta. Sem isso um corte de payload passa por resposta.
        esperado = contabil_dax(
            f'EVALUATE ROW("n", CALCULATE(COUNTROWS({_T456}), {filtro}))')[0]['n'] or 0
        if len(data) != esperado:
            raise RuntimeError(
                f'Vieram {len(data):,} linhas mas o modelo tem {esperado:,} no filtro — '
                f'resposta cortada. Reduza o período.')
        # O que ficou de fora precisa aparecer na tela — nunca esconder calado.
        oculto = contabil_dax(f"""EVALUATE ROW(
            "transf_linhas", CALCULATE(COUNTROWS({_T456}),
                {_T456}[data] >= {_dax_data_str(start)}, {_T456}[data] <= {_dax_data_str(end)},
                {_T456}[transferencia_interna] = TRUE()),
            "transf_valor", CALCULATE(SUMX({_T456}, ABS({_T456}[valor])),
                {_T456}[data] >= {_dax_data_str(start)}, {_T456}[data] <= {_dax_data_str(end)},
                {_T456}[transferencia_interna] = TRUE()),
            "programados", CALCULATE(COUNTROWS({_T456}),
                {_T456}[data] >= {_dax_data_str(start)}, {_T456}[data] <= {_dax_data_str(end)},
                {_T456}[realizado] = FALSE()))""")[0]
        cols = list(data[0].keys()) if data else []
        return jsonify({'ok': True, 'columns': cols, 'data': data, 'count': len(data),
                        'oculto': oculto, 'guardas': {'transferencias': transf, 'programados': prog}})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/faturas')
@page_required('contabil')
def api_contabil_faturas():
    """441 em dois grãos. `grao=fatura` (padrão) ou `grao=ctrc`.

    ATENÇÃO ao grão CTRC: valor_frete é o valor CHEIO do CTRC, e um CTRC
    repartido entre faturas aparece inteiro em cada uma (45 faturas,
    R$ 172.850,39). Para valor por fatura, use vlr_ctrcs do grão fatura."""
    grao = request.args.get('grao', 'fatura')
    start, end = request.args.get('start'), request.args.get('end')
    try:
        dax = _dax_faturas(grao, start, end)
        data = contabil_dax(dax)
        cols = list(data[0].keys()) if data else []
        return jsonify({'ok': True, 'grao': grao, 'columns': cols, 'data': data,
                        'count': len(data), 'data_filtro': _DATA_441[grao]})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# A data de recorte da fatura é a que a contadora especificou no PARA
# GABRIEL.xlsx: "SOMENTE DOCUMENTOS COM JUROS/DESCONTOS LANÇADOS PELA DATA DA
# LIQUIDAÇÃO DA PARCELA" — ou seja, `pagamento`.
_DATA_441 = {'fatura': 'pagamento', 'ctrc': 'faturas do período'}

# Teto de faturas para montar a lista literal no DAX. Acima disso a consulta
# fica gigante; melhor avisar do que devolver meia resposta.
_MAX_FATURAS_IN = 3000


def _faturas_do_periodo(start, end):
    """Números de fatura liquidados no período, no formato do grão CTRC.

    `faturas_441` grava '0038407-1' e `faturas_441_ctrcs` grava '0038407' — é o
    mesmo número sem o dígito verificador, zeros à esquerda preservados.
    """
    linhas = contabil_dax(
        f'EVALUATE SELECTCOLUMNS(FILTER({_T441}, '
        f'{_T441}[pagamento] >= {_dax_data_str(start)} && '
        f'{_T441}[pagamento] <= {_dax_data_str(end)}), "fatura", [fatura])')
    return sorted({str(l['fatura']).split('-')[0].strip()
                   for l in linhas if l.get('fatura')})


def _dax_faturas(grao, start, end):
    if grao != 'ctrc':
        if start and end:
            return (f'EVALUATE FILTER({_T441}, {_T441}[pagamento] >= {_dax_data_str(start)} '
                    f'&& {_T441}[pagamento] <= {_dax_data_str(end)})')
        return f'EVALUATE FILTER({_T441}, TRUE())'

    # O grão CTRC é DETALHE DA FATURA. Recortá-lo pela emissão do próprio CTRC
    # devolve outro conjunto — uma fatura liquidada em agosto carrega CTRC
    # emitido em junho. Medido: 264 faturas de agosto contra 34 CTRCs emitidos
    # em agosto. O recorte certo é "os CTRCs das faturas do período".
    if not (start and end):
        return f'EVALUATE FILTER({_T441_CTR}, TRUE())'
    numeros = _faturas_do_periodo(start, end)
    if not numeros:
        return f'EVALUATE FILTER({_T441_CTR}, FALSE())'
    if len(numeros) > _MAX_FATURAS_IN:
        raise RuntimeError(
            f'{len(numeros):,} faturas no período — lista grande demais para o '
            f'detalhe de CTRC. Reduza o período ou use "sem recorte".')
    lista = '{ ' + ', '.join(f'"{n}"' for n in numeros) + ' }'
    return f'EVALUATE FILTER({_T441_CTR}, {_T441_CTR}[fatura] IN {lista})'


@app.route('/api/contabil/acni')
@page_required('contabil')
def api_contabil_acni():
    """571. O rodapé do SSW infla 90% ao repetir o mestre por documento aplicado;
    aqui o total de adiantamento sai só das linhas mestre."""
    start, end = request.args.get('start'), request.args.get('end')
    try:
        # "CONTABILIZAR SOMENTE OS ITENS QUE TEM DATA DE LIQUIDAÇÃO — PELA DATA
        # DA LIQUIDAÇÃO" (PARA GABRIEL.xlsx). Então o recorte é `liquidac`, não
        # `data` (que é a data do crédito do adiantamento).
        if start and end:
            f = (f'{_T571}[liquidac] >= {_dax_data_str(start)} '
                 f'&& {_T571}[liquidac] <= {_dax_data_str(end)}')
            dax, filtro_tot = f'EVALUATE FILTER({_T571}, {f})', f', {f}'
        else:
            dax, filtro_tot = f'EVALUATE FILTER({_T571}, TRUE())', ''
        data = contabil_dax(dax)
        tot = contabil_dax(f"""EVALUATE ROW(
            "recebido", CALCULATE(SUM({_T571}[valor]), {_T571}[linha_mestre] = TRUE(){filtro_tot}),
            "aplicado", CALCULATE(SUM({_T571}[valor_doc]){filtro_tot}),
            "em_aberto", CALCULATE(SUM({_T571}[saldo_acni]){filtro_tot}),
            "acnis", CALCULATE(COUNTROWS({_T571}), {_T571}[linha_mestre] = TRUE(){filtro_tot}),
            "linhas", CALCULATE(COUNTROWS({_T571}){filtro_tot}),
            "com_liquidacao", CALCULATE(COUNTROWS({_T571}), NOT ISBLANK({_T571}[liquidac])),
            "sem_liquidacao", CALCULATE(COUNTROWS({_T571}), ISBLANK({_T571}[liquidac])))""")[0]
        cols = list(data[0].keys()) if data else []
        return jsonify({'ok': True, 'columns': cols, 'data': data,
                        'count': len(data), 'totais': tot, 'data_filtro': 'liquidac'})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/de-para')
@page_required('contabil')
def api_contabil_depara():
    """Status do de-para evento → conta contábil. SOMENTE LEITURA.

    A contadora preenche no cadastro de evento do SSW (tela 503); o robô do 479
    puxa às 04:00 e grava. O app não escreve: duas fontes divergem sempre.

    Hoje o 479 traz conta do plano DO SSW (5.02.02.01.0025), não do plano
    contábil (4.1.6.01.0013) — são dois planos diferentes, e no SSW o grupo 5 é
    despesa enquanto no plano contábil o grupo 5 é apuração. Por isso a coluna
    'plano_contabil' testa se a conta começa com 3 ou 4, que é o critério da
    própria contadora (PARA GABRIEL.xlsx, linha 44)."""
    try:
        uso = contabil_dax_477_eventos()
        cad = {str(e['codigo']).strip(): e for e in contabil_dax(
            f'EVALUATE SELECTCOLUMNS({_T479}, "codigo", [codigo], "evento", [evento], '
            f'"grupo", [grupo], "conta_debito", [conta_debito], "conta_credito", [conta_credito])')}

        total = sum(u['valor'] or 0 for u in uso)
        acum = 0.0
        for u in uso:
            cod = str(u['evento']).strip()
            reg = cad.get(cod) or {}
            conta = reg.get('conta_debito')
            u['conta_debito'] = conta
            u['conta_credito'] = reg.get('conta_credito')
            u['grupo_ssw'] = reg.get('grupo')
            u['cadastrado'] = cod in cad
            u['plano_contabil'] = bool(conta) and str(conta).strip()[:1] in ('3', '4')
            acum += (u['valor'] or 0)
            u['pct_acumulado'] = (acum / total) if total else 0

        mapeado = sum(u['valor'] or 0 for u in uso if u['plano_contabil'])
        return jsonify({'ok': True, 'eventos': uso, 'totais': {
            'eventos': len(uso), 'valor': total,
            'mapeados': sum(1 for u in uso if u['plano_contabil']),
            'valor_mapeado': mapeado,
            'pct_mapeado': (mapeado / total) if total else 0,
            'sem_cadastro': sum(1 for u in uso if not u['cadastrado']),
            'ref_inicial': REF_INICIAL_477,
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# Escopo do projeto contábil: JANEIRO DE 2026 EM DIANTE. O que é anterior não
# será reprocessado (decisão do Gabriel, 17/08/2026). Vale para tudo que tocar o
# 477 — priorização, de-para, arquivo de importação.
#
# Não é detalhe de filtro: muda a conclusão. No histórico completo a despesa é
# R$ 332,4 mi e os maiores eventos são os fretes; em 2026 são R$ 70,8 mi e o
# topo vira imobilizado, CDC e ICMS. Os três eventos aposentados (5213, 5216,
# 5410) têm ZERO em 2026 — deixam de ser pendência.
REF_INICIAL_477 = '2026/01'


def ref_final_477():
    """Teto do escopo = a competência corrente.

    O 477 tem lançamento com competência FUTURA — parcela a vencer de
    financiamento e consórcio. São 117 competências além de 2026/12, somando
    R$ 18,5 mi. Um filtro só com piso (`REF >= '2026/01'`) é comparação de
    texto e varre tudo isso para dentro, inflando a despesa do exercício em 26%.
    """
    from datetime import date
    h = date.today()
    return f'{h.year:04d}/{h.month:02d}'


def filtro_ref_477(tabela="'public consulta_despesas_477'"):
    """Escopo do projeto contábil, nas duas pontas e com formato validado.
    `LEN = 7` derruba REF malformado (existe '20ES/6' na base)."""
    return (f'{tabela}[REF] >= "{REF_INICIAL_477}" && '
            f'{tabela}[REF] <= "{ref_final_477()}" && '
            f'LEN({tabela}[REF]) = 7')


def contabil_dax_477_eventos():
    """Eventos usados na despesa dentro do escopo, com valor.

    Vem do dataset do DRE (o 477 mora lá), não do contábil. Ordenar por valor é
    o que faz o preenchimento começar onde está o dinheiro."""
    token = get_token()
    dax = ("EVALUATE SUMMARIZECOLUMNS("
           "'public consulta_despesas_477'[evento], "
           "'public consulta_despesas_477'[descr_evento], "
           f"FILTER(ALL('public consulta_despesas_477'), {filtro_ref_477()}), "
           "\"lancamentos\", COUNTROWS('public consulta_despesas_477'), "
           "\"valor\", SUM('public consulta_despesas_477'[vlr_final]))")
    result = execute_dax(token, dax, dataset_id=CONFIG['dre_dataset_id'])
    rows = clean_rows(result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', []))
    agg = {}
    for r in rows:
        k = (str(r.get('evento') or '').strip(), r.get('descr_evento') or '')
        a = agg.setdefault(k, {'evento': k[0], 'descr_evento': k[1],
                               'lancamentos': 0, 'valor': 0.0})
        a['lancamentos'] += r.get('lancamentos') or 0
        a['valor'] += r.get('valor') or 0.0
    return sorted(agg.values(), key=lambda x: -x['valor'])


# ── Configuração: plano de contas, eventos e contas fixas ──────────────────
# Tudo que é DECISÃO da contadora mora em tabela e é editável por ela. Só o
# encanamento — como a partida dobrada se monta, qual data recorta cada
# relatório, as guardas — fica no código. O objetivo é ela nunca precisar
# pedir deploy para mudar uma regra.

VALORES_DESPESA = ('SIM', 'NAO', 'PARCIAL')
VALORES_PROVISAO = ('SIM', 'NAO')


@app.route('/api/contabil/plano-contas')
@page_required('contabil')
def api_contabil_plano_contas():
    """Alimenta os campos de escolha. `?analitica=1` traz só quem recebe lançamento."""
    so_analitica = request.args.get('analitica') == '1'
    grupos = [g for g in (request.args.get('grupos') or '').split(',') if g]
    try:
        conn = get_db(); cur = conn.cursor()
        sql = ("SELECT classificacao, codigo_reduzido, descricao, niveis, analitica, grupo "
               "FROM contabil_plano_contas WHERE 1=1")
        p = []
        if so_analitica:
            sql += " AND analitica = TRUE"
        if grupos:
            sql += " AND grupo = ANY(%s)"; p.append(grupos)
        cur.execute(sql + " ORDER BY classificacao", p)
        data = [{'classificacao': r[0], 'codigo': r[1], 'descricao': r[2],
                 'niveis': r[3], 'analitica': r[4], 'grupo': r[5]} for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _config_eventos():
    """Valor corrente por evento = a linha mais nova. A tabela é append-only,
    então 'corrente' é uma consulta, não um campo que alguém mantém."""
    conn = get_db(); cur = conn.cursor()
    cur.execute("""
        SELECT DISTINCT ON (evento)
               evento, descricao, conta_debito, conta_credito, tem_nota,
               contabiliza_despesa, contabiliza_provisao, aproveita_credito,
               importar_fiscal, validar_simples, grupo_importacao, observacao,
               usuario_nome, criado_em
          FROM contabil_evento_conta
         ORDER BY evento, criado_em DESC, id DESC
    """)
    cols = ['evento', 'descricao', 'conta_debito', 'conta_credito', 'tem_nota',
            'contabiliza_despesa', 'contabiliza_provisao', 'aproveita_credito',
            'importar_fiscal', 'validar_simples', 'grupo_importacao', 'observacao',
            'usuario_nome', 'criado_em']
    d = {}
    for r in cur.fetchall():
        item = dict(zip(cols, r))
        item['criado_em'] = item['criado_em'].isoformat() if item['criado_em'] else None
        d[str(item['evento']).strip()] = item
    cur.close(); conn.close()
    return d


@app.route('/api/contabil/eventos')
@page_required('contabil')
def api_contabil_eventos():
    """Lista de configuração dos eventos, ordenada por R$ com % acumulado.

    A lista é a UNIÃO de: eventos usados no 477 no escopo ∪ cadastro do 479 ∪
    o que já tem configuração aqui. Só o 479 não basta: ele é substituição
    total a cada carga, e evento desativado no SSW some de lá — a linha não
    pode desaparecer da tela e levar junto a vinculação de um fechamento."""
    try:
        uso = {u['evento']: u for u in contabil_dax_477_eventos()}
        cad = {str(e['codigo']).strip(): e for e in contabil_dax(
            f'EVALUATE SELECTCOLUMNS({_T479}, "codigo", [codigo], "evento", [evento], '
            f'"grupo", [grupo], "conta_debito", [conta_debito], '
            f'"conta_credito", [conta_credito])')}
        cfg = _config_eventos()

        total = sum(u['valor'] or 0 for u in uso.values())
        chaves = sorted(set(uso) | set(cad) | set(cfg),
                        key=lambda k: -(uso.get(k, {}).get('valor') or 0))

        linhas, acum = [], 0.0
        for k in chaves:
            u = uso.get(k, {})
            c = cfg.get(k, {})
            r = cad.get(k, {})
            valor = u.get('valor') or 0
            acum += valor
            despesa = (c.get('contabiliza_despesa') or '').upper()
            provisao = (c.get('contabiliza_provisao') or '').upper()
            linhas.append({
                'evento': k,
                'descricao': c.get('descricao') or r.get('evento') or u.get('descr_evento') or '',
                'lancamentos': u.get('lancamentos') or 0,
                'valor': valor,
                'pct_acumulado': (acum / total) if total else 0,
                'no_477': k in uso, 'no_479': k in cad,
                'conta_ssw': r.get('conta_debito'),
                'grupo_ssw': r.get('grupo'),
                'conta_debito': c.get('conta_debito'),
                'conta_credito': c.get('conta_credito'),
                'tem_nota': c.get('tem_nota'),
                'contabiliza_despesa': despesa or None,
                'contabiliza_provisao': provisao or None,
                'aproveita_credito': c.get('aproveita_credito'),
                'importar_fiscal': c.get('importar_fiscal'),
                'validar_simples': c.get('validar_simples'),
                'grupo_importacao': c.get('grupo_importacao'),
                'observacao': c.get('observacao'),
                'editado_por': c.get('usuario_nome'),
                'editado_em': c.get('criado_em'),
                # Precisa de conta quando gera despesa (SIM/PARCIAL) ou quando o
                # pagamento debita a conta do evento em vez de fornecedores.
                'precisa_conta': despesa in ('SIM', 'PARCIAL') or provisao == 'NAO',
            })

        precisam = [l for l in linhas if l['precisa_conta']]
        prontos = [l for l in precisam if l['conta_debito']]
        return jsonify({'ok': True, 'eventos': linhas, 'totais': {
            'eventos': len(linhas), 'valor': total,
            'precisam_conta': len(precisam),
            'valor_precisa': sum(l['valor'] for l in precisam),
            'preenchidos': len(prontos),
            'valor_preenchido': sum(l['valor'] for l in prontos),
            'sem_configuracao': sum(1 for l in linhas if not l['contabiliza_despesa']),
            'ref_inicial': REF_INICIAL_477, 'ref_final': ref_final_477(),
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/eventos', methods=['POST'])
@page_required('contabil')
def api_contabil_eventos_salvar():
    """Grava uma LINHA NOVA. Nunca UPDATE — o histórico é a própria tabela."""
    b = request.get_json(silent=True) or {}
    evento = str(b.get('evento') or '').strip()
    if not evento:
        return jsonify({'ok': False, 'error': 'Informe o evento'}), 400

    despesa = (b.get('contabiliza_despesa') or '').strip().upper() or None
    provisao = (b.get('contabiliza_provisao') or '').strip().upper() or None
    if despesa and despesa not in VALORES_DESPESA:
        return jsonify({'ok': False, 'error':
                        f'contabiliza_despesa deve ser {" / ".join(VALORES_DESPESA)}'}), 400
    if provisao and provisao not in VALORES_PROVISAO:
        return jsonify({'ok': False, 'error':
                        f'contabiliza_provisao deve ser {" / ".join(VALORES_PROVISAO)}'}), 400

    try:
        conn = get_db(); cur = conn.cursor()

        # Conta só entra se existir no plano E receber lançamento. É a trava que
        # impede '4.1.6.01.013' digitado torto de chegar no arquivo de importação.
        for campo in ('conta_debito', 'conta_credito'):
            v = (b.get(campo) or '').strip()
            if not v:
                continue
            cur.execute("SELECT analitica FROM contabil_plano_contas WHERE classificacao = %s", (v,))
            r = cur.fetchone()
            if not r:
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error':
                                f'{campo}: a conta {v} não existe no plano'}), 400
            if not r[0]:
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error':
                                f'{campo}: {v} é conta sintética e não recebe lançamento'}), 400

        atual = _config_eventos().get(evento, {})
        cur.execute("""
            INSERT INTO contabil_evento_conta
                (evento, descricao, conta_debito, conta_credito, tem_nota,
                 contabiliza_despesa, contabiliza_provisao, aproveita_credito,
                 importar_fiscal, validar_simples, grupo_importacao, observacao,
                 usuario_id, usuario_nome)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING id, criado_em
        """, (
            evento,
            (b.get('descricao') or atual.get('descricao') or '')[:200] or None,
            (b.get('conta_debito') or '').strip() or None,
            (b.get('conta_credito') or '').strip() or None,
            b.get('tem_nota'), despesa, provisao,
            b.get('aproveita_credito'), b.get('importar_fiscal'), b.get('validar_simples'),
            (b.get('grupo_importacao') or '').strip() or None,
            (b.get('observacao') or '').strip() or None,
            session.get('user_id'), session.get('nome'),
        ))
        novo_id, criado = cur.fetchone()
        conn.commit(); cur.close(); conn.close()
        return jsonify({'ok': True, 'id': novo_id,
                        'criado_em': criado.isoformat() if criado else None})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/eventos/historico')
@page_required('contabil')
def api_contabil_eventos_historico():
    evento = (request.args.get('evento') or '').strip()
    if not evento:
        return jsonify({'ok': False, 'error': 'Informe o evento'}), 400
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT id, conta_debito, conta_credito, tem_nota, contabiliza_despesa,
                   contabiliza_provisao, observacao, usuario_nome, criado_em
              FROM contabil_evento_conta
             WHERE evento = %s ORDER BY criado_em DESC, id DESC
        """, (evento,))
        data = [{'id': r[0], 'conta_debito': r[1], 'conta_credito': r[2],
                 'tem_nota': r[3], 'contabiliza_despesa': r[4],
                 'contabiliza_provisao': r[5], 'observacao': r[6],
                 'usuario_nome': r[7],
                 'criado_em': r[8].isoformat() if r[8] else None}
                for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/contas-fixas')
@page_required('contabil')
def api_contabil_contas_fixas():
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT f.chave, f.classificacao, p.codigo_reduzido, p.descricao,
                   f.descricao, f.observacao, f.usuario_nome, f.atualizado_em
              FROM contabil_conta_fixa f
              LEFT JOIN contabil_plano_contas p ON p.classificacao = f.classificacao
             ORDER BY (f.chave = 'FORNECEDOR_PADRAO') DESC, f.descricao
        """)
        data = [{'chave': r[0], 'classificacao': r[1], 'codigo': r[2],
                 'conta_descricao': r[3], 'descricao': r[4], 'observacao': r[5],
                 'usuario_nome': r[6],
                 'atualizado_em': r[7].isoformat() if r[7] else None}
                for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data,
                        'pendentes': sum(1 for x in data if not x['classificacao'])})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/contabil/contas-fixas', methods=['POST'])
@page_required('contabil')
def api_contabil_contas_fixas_salvar():
    b = request.get_json(silent=True) or {}
    chave = (b.get('chave') or '').strip()
    classificacao = (b.get('classificacao') or '').strip() or None
    if not chave:
        return jsonify({'ok': False, 'error': 'Informe a chave'}), 400
    try:
        conn = get_db(); cur = conn.cursor()
        if classificacao:
            cur.execute("SELECT analitica FROM contabil_plano_contas WHERE classificacao = %s",
                        (classificacao,))
            r = cur.fetchone()
            if not r:
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error':
                                f'A conta {classificacao} não existe no plano'}), 400
            if not r[0]:
                cur.close(); conn.close()
                return jsonify({'ok': False, 'error':
                                f'{classificacao} é sintética e não recebe lançamento'}), 400
        cur.execute("""
            UPDATE contabil_conta_fixa
               SET classificacao = %s, observacao = %s,
                   usuario_nome = %s, atualizado_em = NOW()
             WHERE chave = %s
        """, (classificacao, (b.get('observacao') or '').strip() or None,
              session.get('nome'), chave))
        if cur.rowcount == 0:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': f'Chave {chave} não existe'}), 404
        conn.commit(); cur.close(); conn.close()
        contas_fixas(forcar=True)   # o quadro por banco lê daqui
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def _csv_contabil(nome, dax_fn):
    """Streaming CSV com BOM, mesmo padrão das Despesas/Conhecimentos."""
    def gerar():
        yield '﻿'
        data = dax_fn()
        if not data:
            yield _csv_linha(['sem dados'])
            return
        cols = list(data[0].keys())
        yield _csv_linha(cols)
        for row in data:
            yield _csv_linha([row.get(c) for c in cols])
    return Response(gerar(), mimetype='text/csv; charset=utf-8',
                    headers={'Content-Disposition': f'attachment; filename="{nome}"'})


@app.route('/api/contabil/extrato/csv')
@page_required('contabil')
def api_contabil_extrato_csv():
    start, end = request.args.get('start'), request.args.get('end')
    if not start or not end:
        return jsonify({'ok': False, 'error': 'Informe start e end'}), 400
    banco = request.args.get('banco') or None
    transf = request.args.get('transferencias') == '1'
    prog = request.args.get('programados') == '1'
    filtro = _filtro_periodo_456(start, end, banco, transf, prog)
    return _csv_contabil(f'extrato_456_{start}_{end}.csv',
                         lambda: contabil_dax(_select_extrato(filtro)))


@app.route('/api/contabil/faturas/csv')
@page_required('contabil')
def api_contabil_faturas_csv():
    grao = request.args.get('grao', 'fatura')
    start, end = request.args.get('start'), request.args.get('end')
    sufixo = f'_{start}_{end}' if start and end else ''
    return _csv_contabil(f'441_{grao}{sufixo}.csv',
                         lambda: contabil_dax(_dax_faturas(grao, start, end)))


@app.route('/api/contabil/acni/csv')
@page_required('contabil')
def api_contabil_acni_csv():
    start, end = request.args.get('start'), request.args.get('end')
    if start and end:
        dax = (f'EVALUATE FILTER({_T571}, {_T571}[liquidac] >= {_dax_data_str(start)} '
               f'&& {_T571}[liquidac] <= {_dax_data_str(end)})')
        nome = f'571_acni_{start}_{end}.csv'
    else:
        dax, nome = f'EVALUATE FILTER({_T571}, TRUE())', '571_acni.csv'
    return _csv_contabil(nome, lambda: contabil_dax(dax))


# ════════════════════════════════════════
# CHAT IA — Analista Financeiro DRE
# ════════════════════════════════════════

def _fmt_brl(v):
    try:
        v = float(v or 0)
    except (TypeError, ValueError):
        return 'R$ 0,00'
    return f"R$ {v:,.2f}".replace(',', '_').replace('.', ',').replace('_', '.')


_DESCRICAO_PADRAO = {
    'unico': 'Mês único — análise pontual.',
    'contiguo': 'Período contíguo (meses consecutivos no mesmo ano) — analise como série temporal sequencial. Cite o período inteiro e variação MoM.',
    'mesmo_mes_varios_anos': 'COMPARATIVO ANUAL: o mesmo mês selecionado em vários anos (ex: Mai/21, Mai/22, ..., Mai/26). NÃO analise como se fosse só o último mês. Faça comparação ano-a-ano: identifique tendência (melhorou/piorou ao longo dos anos), mês com melhor/pior performance, evolução das margens.',
    'multi_anos_multi_meses': 'COMPARATIVO MISTO: múltiplos meses em múltiplos anos. Analise cada bloco de ano separadamente. Identifique padrões sazonais (mesmo mês comporta-se igual entre anos?) e tendências (cada ano melhor ou pior que o anterior?).',
    'esparso_mesmo_ano': 'MESES NÃO-CONSECUTIVOS NO MESMO ANO. Analise cada mês como ponto independente, não como série contínua.',
    'esparso': 'Seleção esparsa. Analise mês a mês.',
}


def _montar_prompt_chat(contexto):
    """Monta system prompt com dados financeiros estruturados (todos pré-calculados)."""
    partes = []
    periodo = contexto.get('periodo', [])
    modo = contexto.get('modo', 'acumulado')
    padrao = contexto.get('padrao', 'contiguo')
    partes.append(f"PERÍODO ANALISADO: {', '.join(periodo) if periodo else 'não informado'}")
    partes.append(f"MODO: {modo}")
    partes.append(f"PADRÃO DE SELEÇÃO: {padrao}")
    partes.append(f"COMO INTERPRETAR: {_DESCRICAO_PADRAO.get(padrao, '')}\n")

    # Modo fracionado: receita recortada por dia, despesa do mês inteiro.
    # Sem este aviso o modelo lê a margem negativa estrutural como prejuízo real.
    if contexto.get('fracionado'):
        di, df = contexto.get('dia_ini'), contexto.get('dia_fim')
        partes.append(
            f"*** ATENÇÃO — MODO FRACIONADO (dias {di} a {df} de cada mês) ***\n"
            f"A RECEITA está recortada aos dias {di}-{df}. A DESPESA é do MÊS INTEIRO "
            f"(a competência do ERP não tem dia, não há como fracioná-la).\n"
            f"CONSEQUÊNCIA: todas as linhas abaixo de 'Receita Bruta' (Receita Líquida, EBITDA, "
            f"LAIR, Lucro Líquido, Resultado Final) e TODAS as margens comparam receita parcial "
            f"com despesa cheia — são NEGATIVAS POR CONSTRUÇÃO e NÃO significam prejuízo.\n"
            f"VOCÊ NÃO PODE: chamar esses números de prejuízo, alertar sobre margem negativa, "
            f"citar EBITDA/Resultado/margens em reais ou %, nem comparar com o mês fechado.\n"
            f"VOCÊ DEVE: analisar APENAS a Receita Bruta do recorte e a Despesa do mês, cada uma "
            f"comparada consigo mesma entre os períodos (crescimento, tendência, aceleração). "
            f"Se perguntarem sobre lucro/margem/resultado, responda que no modo fracionado só a "
            f"receita é comparável e sugira desligar o filtro de dias para ver o resultado.\n"
        )

    dre = contexto.get('dre') or {}
    if modo == 'acumulado' and dre:
        partes.append("--- DRE (ACUMULADO) ---")
        labels = [
            ('receita_bruta',            'Receita Bruta'),
            ('deducoes',                 '(-) Deduções'),
            ('receita_liquida',          '= Receita Líquida'),
            ('custo_operacional',        '(-) Custo Operacional'),
            ('despesas_administrativas', '(-) Despesas Administrativas'),
            ('ebitda',                   '= EBITDA'),
            ('despesas_financeiras',     '(-) Despesas Financeiras'),
            ('lair',                     '= LAIR'),
            ('impostos',                 '(-) Impostos'),
            ('lucro_liquido',            '= Lucro Líquido'),
            ('investimentos',            '(-) Investimentos'),
            ('pos_investimento',         '= Pós Investimento'),
            ('retiradas',                '(-) Retiradas'),
            ('resultado_final',          '= Resultado Final'),
        ]
        for key, label in labels:
            if key in dre:
                partes.append(f"{label}: {_fmt_brl(dre[key])}")

    elif modo == 'mensal':
        dre_meses = contexto.get('dre_por_mes', [])
        partes.append("--- DRE POR MÊS ---")
        for d in dre_meses:
            partes.append(f"\n[{d.get('mes', '?')}]")
            for k, v in d.items():
                if k != 'mes' and isinstance(v, (int, float)):
                    partes.append(f"  {k}: {_fmt_brl(v)}")

    margens = contexto.get('margens_agregadas') or contexto.get('margens') or {}
    if margens:
        partes.append("\n--- MARGENS AGREGADAS DO PERÍODO TOTAL (%) ---")
        for k, v in margens.items():
            try:
                partes.append(f"{k.replace('_', ' ').title()}: {float(v):.1f}%".replace('.', ','))
            except (TypeError, ValueError):
                pass

    variacoes = contexto.get('variacao_ultimo_vs_anterior') or {}
    if variacoes:
        partes.append("\n--- VARIAÇÃO ENTRE ÚLTIMO MÊS E O ANTERIOR DA SELEÇÃO ---")
        for k, v in variacoes.items():
            try:
                seta = '↑' if v > 0 else ('↓' if v < 0 else '→')
                partes.append(f"{k}: {seta} {float(v):.1f}%".replace('.', ','))
            except (TypeError, ValueError):
                pass

    top_sub = contexto.get('top_subgrupos') or []
    if top_sub:
        partes.append("\n--- TOP SUBGRUPOS DE DESPESA ---")
        for s in top_sub:
            partes.append(f"- {s.get('nome')} ({s.get('grupo')}): {_fmt_brl(s.get('valor'))} ({s.get('pct', 0):.1f}%)".replace('.', ','))

    pareto = contexto.get('pareto_80') or []
    if pareto:
        partes.append(f"\n--- PARETO 80% ({len(pareto)} subgrupos respondem por 80% das despesas) ---")
        partes.append(', '.join(pareto))

    total_desp = contexto.get('total_despesas')
    if total_desp is not None:
        partes.append(f"\nTOTAL DE DESPESAS: {_fmt_brl(total_desp)}")

    dados_texto = '\n'.join(partes)

    return f"""Você é o Analista Financeiro da Rizza Transportes — uma transportadora rodoviária de cargas com operações em SP, RJ, GO, ES, BA. Você analisa DRE e despesas.

REGRAS INVIOLÁVEIS:
1. Use APENAS os números fornecidos abaixo — todos JÁ CALCULADOS. Não recalcule.
2. Não invente nada. Se a pergunta exigir dado que não foi enviado, diga "essa informação não está no período selecionado".
3. Máximo 3 parágrafos curtos. Diretor não lê longo.
4. Formate em **markdown**: use **negrito** para destacar números/conclusões críticas.
5. Sempre cite valores em R$ (formato brasileiro) e %.
6. Tom: direto, profissional, sem rodeios.
7. Se identificar problema (margem negativa, queda, custo alto), DESTAQUE em negrito.
8. Perguntas fora de finanças/DRE → responda: "Só consigo analisar dados financeiros da DRE."

REGRA CRÍTICA SOBRE SELEÇÕES MÚLTIPLAS:
- Quando "PADRÃO DE SELEÇÃO" for DIFERENTE de 'unico' ou 'contiguo', você DEVE percorrer TODOS os meses listados em "DRE POR MÊS" — não foque só no último mês.
- Em 'mesmo_mes_varios_anos' (ex: Mai/21 a Mai/26): trate como COMPARATIVO ANUAL. Identifique evolução, melhor/pior ano, tendência.
- Em 'multi_anos_multi_meses' (ex: Mar+Abr+Mai de 24/25/26): trate como COMPARATIVO MISTO. Compare blocos de ano, identifique padrões sazonais.
- Em 'esparso_mesmo_ano': analise cada mês como ponto independente.
- NUNCA responda como se o período fosse apenas o último mês quando há vários meses na seleção.

CONTEXTO DA EMPRESA:
- Operação "fretes-pesada": terceiriza muita carga (subgrupo Fretes domina ~38%)
- Margens saudáveis para o setor: EBITDA acima de 10%, Líquida acima de 5%
- Resultado negativo é alerta vermelho

DADOS DO PERÍODO ATUAL:
{dados_texto}
"""


@app.route('/api/chat-dre', methods=['POST'])
@page_required('dre')
def chat_dre():
    from openai import OpenAI
    data = request.get_json() or {}
    pergunta = (data.get('pergunta') or '').strip()
    contexto = data.get('contexto') or {}
    historico = data.get('historico') or []

    if not pergunta:
        return jsonify({'ok': False, 'error': 'Pergunta vazia'}), 400

    system_prompt = _montar_prompt_chat(contexto)
    messages = [{'role': 'system', 'content': system_prompt}]
    for m in historico[-6:]:
        role = m.get('role')
        content = m.get('content', '')
        if role in ('user', 'assistant') and content:
            messages.append({'role': role, 'content': content})
    if not messages or messages[-1].get('content') != pergunta:
        messages.append({'role': 'user', 'content': pergunta})

    def gerar():
        try:
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            stream = client.chat.completions.create(
                model='gpt-4.1-mini',
                messages=messages,
                max_tokens=800,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                if not chunk.choices: continue
                delta = chunk.choices[0].delta.content
                if delta:
                    yield f"data: {json.dumps({'token': delta})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'erro': str(e)})}\n\n"

    return Response(stream_with_context(gerar()),
                    mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


# ════════════════════════════════════════════════════════════════════════
# EMBARQUES — Módulo operacional de lançamento de cargas
# ════════════════════════════════════════════════════════════════════════

# Cache em memória (5min) para motoristas/veículos vindos do Power BI
_EMBARQUES_CACHE = {}
_CACHE_TTL_SEG = 300

def _cache_get(key):
    entry = _EMBARQUES_CACHE.get(key)
    if entry and (time.time() - entry['ts']) < _CACHE_TTL_SEG:
        return entry['data']
    return None

def _cache_set(key, data):
    _EMBARQUES_CACHE[key] = {'data': data, 'ts': time.time()}


def _eh_rizza(proprietario):
    """Identifica se o proprietário é Rizza (busca parcial case-insensitive)."""
    return 'RIZZA' in (proprietario or '').upper()


def _norm_tipo_veiculo(tipo):
    """Classifica o TIPO cru do veiculos_045 em Cavalo/Carreta/Truck.
    Usa prefixo/palavra-chave (não match exato) porque o Winthor cadastra variações
    como 'CAVALO TRUCADO', 'TOCO', 'BALSA' que antes eram descartadas por engano —
    sumindo o veículo do lançamento de carga. Tipos desconhecidos viram None."""
    t = str(tipo or '').strip().upper()
    if not t:
        return None
    if t.startswith('CAVALO'):                       # CAVALO, CAVALO TRUCADO, ...
        return 'Cavalo'
    if t.startswith('CARRETA') or t == 'BALSA':      # BALSA é reboque (prancha)
        return 'Carreta'
    if t.startswith('TRUCK') or t == 'TOCO':         # TOCO = rígido de eixo simples
        return 'Truck'
    return None                                      # OUTROS e desconhecidos ficam fora


def _pode_editar_carga(criado_por_id):
    """Admin ou quem criou a carga pode editar."""
    if session.get('role') == 'admin':
        return True
    return session.get('user_id') == criado_por_id


def _classifica_tipo_operacao(cavalo_eh_rizza, carreta_eh_rizza, tem_carreta1=True):
    """Tipo de operação esperado conforme proprietários do cavalo e carreta1.
    Sem carreta1 (truck rígido ou cadastro incompleto), classifica só pelo cavalo
    — senão a ausência da carreta seria lida como 'não-Rizza' e daria Agregado errado."""
    if not tem_carreta1:
        return 'Frota' if cavalo_eh_rizza else 'Terceiro'
    if cavalo_eh_rizza and carreta_eh_rizza:
        return 'Frota'
    if cavalo_eh_rizza or carreta_eh_rizza:
        return 'Agregado'
    return 'Terceiro'


def _pick(row, *keys):
    """Pega o primeiro valor não-vazio entre variações de nome de coluna."""
    for k in keys:
        v = row.get(k)
        if v not in (None, ''):
            return v
    return None


def _csv_linha_embarques(valores):
    """Wrapper local em torno de _csv_linha para clareza."""
    return _csv_linha(valores)


def _buscar_conflitos(cpf, placas, exclude_id=0):
    """Retorna lista de conflitos com cargas ativas.
       'placas' é lista de strings uppercase (cavalo, carreta1, carreta2 — sem nulls).
       'exclude_id' permite ignorar a própria carga ao editar.

       Assimetria do desengate (status 'Desengatada' = carreta carregada parada no
       destino, cavalo+motorista liberados):
         - motorista (CPF) e CAVALO → ativos só em Aberta/Em rota/No destino
           (liberados quando a carga está 'Desengatada').
         - CARRETA → ainda comprometida; bloqueia também em 'Desengatada'.
    """
    placas = [p for p in (placas or []) if p]
    cpf = (cpf or '').strip()
    if not cpf and not placas:
        return []
    ativas_cav = ('Aberta', 'Em rota', 'No destino')
    ativas_carreta = ('Aberta', 'Em rota', 'No destino', 'Desengatada')
    conn = get_db(); cur = conn.cursor()
    try:
        conflitos = []

        if cpf:
            # Compara SÓ OS DÍGITOS: o lançamento manual grava '04455930671' e o
            # robô grava '044.559.306-71' (é o formato que vem do manifesto do
            # SSW). Medido: 12/12 das cargas manuais sem pontuação, 30/30 das do
            # robô com — a igualdade exata nunca casava os dois mundos, e o mesmo
            # motorista podia ficar em duas viagens ativas sem ninguém ser avisado.
            cur.execute("""
                SELECT id, numero, status, data_carregamento, motorista_nome
                FROM embarques_cargas
                WHERE regexp_replace(motorista_cpf, '[^0-9]', '', 'g') =
                      regexp_replace(%s, '[^0-9]', '', 'g')
                  AND status IN ('Aberta', 'Em rota', 'No destino')
                  AND id <> %s
                ORDER BY data_carregamento DESC
                LIMIT 5
            """, (cpf, exclude_id))
            for r in cur.fetchall():
                conflitos.append({
                    'tipo': 'motorista',
                    'recurso': r[4] or cpf,
                    'carga_id': r[0],
                    'numero': r[1],
                    'status': r[2],
                    'data_carregamento': r[3].isoformat() if r[3] else None,
                })

        if placas:
            ph = ','.join(['%s'] * len(placas))
            cav_ph = ','.join(['%s'] * len(ativas_cav))
            car_ph = ','.join(['%s'] * len(ativas_carreta))
            # Cavalo só conflita em status "duros"; carreta conflita também em Desengatada.
            cur.execute(f"""
                SELECT id, numero, status, data_carregamento,
                       cavalo_placa, carreta1_placa, carreta2_placa
                FROM embarques_cargas
                WHERE id <> %s
                  AND (
                    (cavalo_placa IN ({ph}) AND status IN ({cav_ph}))
                    OR ((carreta1_placa IN ({ph}) OR carreta2_placa IN ({ph}))
                        AND status IN ({car_ph}))
                  )
                ORDER BY data_carregamento DESC
                LIMIT 10
            """, (exclude_id, *placas, *ativas_cav, *placas, *placas, *ativas_carreta))
            for r in cur.fetchall():
                cid, num, st, dt, cav, c1, c2 = r
                dt_iso = dt.isoformat() if dt else None
                for placa in placas:
                    # cavalo só é conflito se a carga ainda está nos status duros
                    if cav == placa and st in ativas_cav:
                        conflitos.append({'tipo': 'cavalo',  'recurso': placa, 'carga_id': cid, 'numero': num, 'status': st, 'data_carregamento': dt_iso})
                    if (c1 == placa or c2 == placa) and st in ativas_carreta:
                        conflitos.append({'tipo': 'carreta', 'recurso': placa, 'carga_id': cid, 'numero': num, 'status': st, 'data_carregamento': dt_iso})
        return conflitos
    finally:
        cur.close(); conn.close()


@app.route('/api/embarques/conflitos')
@login_required
def api_embarques_conflitos():
    cpf = (request.args.get('cpf') or '').strip()
    placas_raw = (request.args.get('placas') or '').strip()
    placas = [p.strip().upper() for p in placas_raw.split(',') if p.strip()]
    try:
        exclude_id = int(request.args.get('exclude_id') or 0)
    except (TypeError, ValueError):
        exclude_id = 0
    try:
        data = _buscar_conflitos(cpf, placas, exclude_id)
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Leitura DAX: motoristas ─────────────────────────────────────────────
@app.route('/api/embarques/motoristas')
@login_required
def api_embarques_motoristas():
    if request.args.get('refresh') != '1':
        cached = _cache_get('motoristas')
        if cached is not None:
            return jsonify({'ok': True, 'data': cached, 'count': len(cached), 'cached': True})
    try:
        token = get_token()
        result = execute_dax(token, "EVALUATE 'public motoristas_047'")
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)

        normalizados = []
        for r in data:
            nome = _pick(r, 'nome', 'Nome', 'NOME')
            cpf  = _pick(r, 'cpf', 'CPF', 'Cpf')
            tel  = _pick(r, 'telefone', 'Telefone', 'TELEFONE', 'celular', 'Celular')
            if not nome or not cpf:
                continue
            normalizados.append({
                'nome':      str(nome).strip(),
                'cpf':       str(cpf).strip(),
                'telefone':  str(tel).strip() if tel else None,
            })
        # Dedup por CPF (única chave confiável)
        vistos = {}
        for m in normalizados:
            vistos[m['cpf']] = m
        final = sorted(vistos.values(), key=lambda x: x['nome'])

        _cache_set('motoristas', final)
        return jsonify({'ok': True, 'data': final, 'count': len(final), 'cached': False})

    except requests.exceptions.HTTPError as e:
        detail = ''
        try: detail = e.response.json()
        except Exception: detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Leitura DAX: veículos ───────────────────────────────────────────────
@app.route('/api/embarques/veiculos')
@login_required
def api_embarques_veiculos():
    if request.args.get('refresh') != '1':
        cached = _cache_get('veiculos')
        if cached is not None:
            return jsonify({'ok': True, 'data': cached, 'count': len(cached), 'cached': True})
    try:
        token = get_token()
        # EVALUATE simples — filtro feito em Python (mais robusto que IN no DAX)
        result = execute_dax(token, "EVALUATE 'public veiculos_045'")
        rows = result.get('results', [{}])[0].get('tables', [{}])[0].get('rows', [])
        data = clean_rows(rows)

        normalizados = []
        for r in data:
            placa        = _pick(r, 'placa', 'PLACA', 'Placa')
            tipo         = _pick(r, 'TIPO', 'tipo', 'Tipo')
            marca        = _pick(r, 'marca', 'MARCA', 'Marca')
            modelo       = _pick(r, 'modelo', 'MODELO', 'Modelo')
            carroceria   = _pick(r, 'carroceria', 'CARROCERIA', 'Carroceria')
            proprietario = _pick(r, 'proprietario', 'PROPRIETARIO', 'Proprietario', 'proprietário', 'Proprietário')
            if not placa or not tipo:
                continue
            tipo_norm = _norm_tipo_veiculo(tipo)
            if not tipo_norm:
                continue
            partes = [str(marca or '').strip(), str(modelo or '').strip()]
            marca_modelo = ' '.join(p for p in partes if p) or None
            normalizados.append({
                'placa':        str(placa).strip().upper(),
                'tipo':         tipo_norm,
                'marca_modelo': marca_modelo,
                'carroceria':   str(carroceria).strip() if carroceria else None,
                'proprietario': str(proprietario).strip() if proprietario else None,
                'eh_rizza':     _eh_rizza(proprietario),
            })
        vistos = {}
        for v in normalizados:
            vistos[v['placa']] = v
        final = sorted(vistos.values(), key=lambda x: x['placa'])

        _cache_set('veiculos', final)
        return jsonify({'ok': True, 'data': final, 'count': len(final), 'cached': False})

    except requests.exceptions.HTTPError as e:
        detail = ''
        try: detail = e.response.json()
        except Exception: detail = e.response.text
        return jsonify({'ok': False, 'error': str(e), 'detail': detail}), 500
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Clientes (Postgres local) ───────────────────────────────────────────
@app.route('/api/embarques/clientes')
@login_required
def api_embarques_clientes_list():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT id, nome, importado_em FROM clientes ORDER BY nome")
        data = [
            {'id': r[0], 'nome': r[1], 'importado_em': r[2].isoformat() if r[2] else None}
            for r in cur.fetchall()
        ]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/embarques/clientes', methods=['POST'])
@login_required
def api_embarques_clientes_create():
    body = request.get_json(silent=True) or {}
    nome = (body.get('nome') or '').strip()
    if len(nome) < 3:
        return jsonify({'ok': False, 'error': 'Nome inválido (mínimo 3 caracteres)'}), 400
    try:
        conn = get_db()
        cur = conn.cursor()
        # Dedup case-insensitive: se já existe, devolve o id; senão insere.
        # (Não usa ON CONFLICT pra não depender de constraint nomeada — o nome único
        #  de clientes é garantido por ÍNDICE de expressão, não por constraint.)
        cur.execute(
            "SELECT id FROM clientes WHERE LOWER(TRIM(nome)) = LOWER(TRIM(%s)) LIMIT 1",
            (nome,)
        )
        r = cur.fetchone()
        if r:
            new_id = r[0]
            ja_existia = True
        else:
            cur.execute("INSERT INTO clientes (nome) VALUES (%s) RETURNING id", (nome,))
            new_id = cur.fetchone()[0]
            ja_existia = False
        conn.commit()
        cur.close(); conn.close()
        return jsonify({'ok': True, 'id': new_id, 'ja_existia': ja_existia})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Cargas: criação ─────────────────────────────────────────────────────
def _validar_carga_payload(b):
    erros = []
    obrig = ['tipo_operacao', 'origem', 'destinos', 'motorista', 'cavalo', 'data_carregamento']
    # Viagem vazia (sem carga) não exige cliente.
    if not b.get('viagem_vazia'):
        obrig = ['cliente_id', 'cliente_nome'] + obrig
    for c in obrig:
        if c not in b or b.get(c) in (None, '', []):
            erros.append(f'Campo obrigatório ausente: {c}')
    if not erros:
        if b['tipo_operacao'] not in ('Frota', 'Agregado', 'Terceiro'):
            erros.append('tipo_operacao inválido')
        if not isinstance(b['destinos'], list) or len(b['destinos']) < 1:
            erros.append('Informe pelo menos 1 destino')
        ori = b.get('origem') or {}
        if not ori.get('cidade') or not ori.get('uf'):
            erros.append('Origem incompleta (cidade + uf)')
        for i, d in enumerate(b.get('destinos') or []):
            if not d.get('cidade') or not d.get('uf'):
                erros.append(f'Destino {i+1} incompleto')
        mot = b.get('motorista') or {}
        if not mot.get('nome') or not mot.get('cpf'):
            erros.append('Motorista incompleto (nome + cpf)')
        cav = b.get('cavalo') or {}
        if not cav.get('placa') or not cav.get('tipo'):
            erros.append('Cavalo incompleto (placa + tipo)')
        if cav.get('tipo') == 'Cavalo' and not (b.get('carreta1') or {}).get('placa'):
            erros.append('Carreta 1 obrigatória quando o tipo do veículo é Cavalo')
    return erros


@app.route('/api/embarques/cargas', methods=['POST'])
@login_required
def api_embarques_cargas_create():
    b = request.get_json(silent=True) or {}
    erros = _validar_carga_payload(b)
    if erros:
        return jsonify({'ok': False, 'error': 'Validação falhou', 'detail': erros}), 400

    warnings = []
    cav = b.get('cavalo') or {}
    c1  = b.get('carreta1') or {}
    c2  = b.get('carreta2') or {}
    esperado = _classifica_tipo_operacao(bool(cav.get('eh_rizza')), bool(c1.get('eh_rizza')),
                                         tem_carreta1=bool((c1.get('placa') or '').strip()))
    if b['tipo_operacao'] != esperado:
        warnings.append(f"tipo_operacao '{b['tipo_operacao']}' diverge do esperado '{esperado}' pelos proprietários.")

    # Bloqueio de conflito (motorista/veículos em carga ativa)
    mot_pre = b.get('motorista') or {}
    placas_check = [
        (cav.get('placa') or '').upper().strip(),
        (c1.get('placa') or '').upper().strip(),
        (c2.get('placa') or '').upper().strip(),
    ]
    conflitos = _buscar_conflitos(mot_pre.get('cpf'), [p for p in placas_check if p])
    if conflitos:
        # Constrói mensagem detalhada (útil mesmo se o frontend ignorar o campo 'conflitos')
        nums = sorted({c['numero'] for c in conflitos if c.get('numero')})
        msg = 'Recurso já em uso na(s) carga(s) ativa(s): ' + ', '.join(nums)
        return jsonify({
            'ok': False,
            'error': msg,
            'tipo': 'conflito',
            'conflitos': conflitos
        }), 409

    try:
        conn = get_db()
        cur = conn.cursor()
        mot = b['motorista']
        ori = b['origem']
        viagem_vazia = bool(b.get('viagem_vazia'))
        cliente_id = None if viagem_vazia else b.get('cliente_id')
        cliente_nome = None if viagem_vazia else b.get('cliente_nome')
        cur.execute("""
            INSERT INTO embarques_cargas (
                tipo_operacao, status, viagem_vazia,
                cliente_id, cliente_nome,
                origem_cidade, origem_uf,
                motorista_nome, motorista_cpf, motorista_telefone,
                cavalo_placa, cavalo_tipo, cavalo_marca_modelo, cavalo_carroceria, cavalo_proprietario, cavalo_eh_rizza,
                carreta1_placa, carreta1_marca_modelo, carreta1_carroceria, carreta1_proprietario, carreta1_eh_rizza,
                carreta2_placa, carreta2_marca_modelo, carreta2_carroceria, carreta2_proprietario, carreta2_eh_rizza,
                data_carregamento, previsao_entrega, observacoes,
                criado_por_id, criado_por_nome
            ) VALUES (
                %s, 'Aberta', %s,
                %s, %s,
                %s, %s,
                %s, %s, %s,
                %s, %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s,
                %s, %s, %s,
                %s, %s
            ) RETURNING id, criado_em
        """, (
            b['tipo_operacao'], viagem_vazia,
            cliente_id, cliente_nome,
            ori['cidade'], ori['uf'],
            mot['nome'], mot['cpf'], mot.get('telefone'),
            cav['placa'], cav['tipo'], cav.get('marca_modelo'), cav.get('carroceria'), cav.get('proprietario'), bool(cav.get('eh_rizza')),
            c1.get('placa'), c1.get('marca_modelo'), c1.get('carroceria'), c1.get('proprietario'), bool(c1.get('eh_rizza')),
            c2.get('placa'), c2.get('marca_modelo'), c2.get('carroceria'), c2.get('proprietario'), bool(c2.get('eh_rizza')),
            b['data_carregamento'], b.get('previsao_entrega'), b.get('observacoes'),
            session.get('user_id'), session.get('nome'),
        ))
        carga_id, criado_em = cur.fetchone()

        # Destinos + geocoding (centroide IBGE)
        import geocoding
        destinos_inseridos = []
        for i, d in enumerate(b['destinos'], start=1):
            dlat, dlng = geocoding.geocoder_municipio(d['cidade'], d['uf'], conn=conn)
            cur.execute(
                "INSERT INTO embarques_cargas_destinos (carga_id, ordem, cidade, uf, latitude, longitude, data_agendamento) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s)",
                (carga_id, i, d['cidade'], d['uf'], dlat, dlng, d.get('data_agendamento') or None)
            )
            destinos_inseridos.append({'cidade': d['cidade'], 'uf': d['uf'], 'lat': dlat, 'lng': dlng})

        # Cidades de rota (passagem; moldam o caminho, não são entrega)
        rota_inseridas = []
        for i, r in enumerate(b.get('rota') or [], start=1):
            if not r.get('cidade') or not r.get('uf'):
                continue
            rlat, rlng = geocoding.geocoder_municipio(r['cidade'], r['uf'], conn=conn)
            cur.execute(
                "INSERT INTO embarques_cargas_rota (carga_id, ordem, cidade, uf, latitude, longitude) "
                "VALUES (%s, %s, %s, %s, %s, %s)",
                (carga_id, i, r['cidade'], r['uf'], rlat, rlng)
            )
            rota_inseridas.append({'cidade': r['cidade'], 'uf': r['uf'], 'lat': rlat, 'lng': rlng})

        # Geocoding origem
        olat, olng = geocoding.geocoder_municipio(ori['cidade'], ori['uf'], conn=conn)
        cur.execute(
            "UPDATE embarques_cargas SET origem_latitude=%s, origem_longitude=%s WHERE id=%s",
            (olat, olng, carga_id)
        )

        # Gera numero
        ano = criado_em.year
        numero = f"C-{ano}-{carga_id:06d}"
        cur.execute("UPDATE embarques_cargas SET numero = %s WHERE id = %s", (numero, carga_id))

        conn.commit()
        cur.close(); conn.close()

        # Calcula rota planejada via ORS (após commit; falha não derruba o POST).
        # Passa por origem -> cidades de rota -> TODOS os destinos, na ordem.
        ors_warn = None
        pontos = [{'lat': olat, 'lng': olng}] + rota_inseridas + destinos_inseridos
        pontos = [p for p in pontos if p.get('lat') is not None and p.get('lng') is not None]
        if olat is not None and len(pontos) >= 2:
            try:
                import ors_client
                rota = ors_client.tracar_rota_multi(pontos)
                conn2 = get_db()
                cur2 = conn2.cursor()
                cur2.execute("""
                    UPDATE embarques_cargas SET
                        rota_planejada_polyline=%s,
                        distancia_planejada_km=%s,
                        duracao_estimada_min=%s,
                        rota_recalculada_em=NOW()
                    WHERE id=%s
                """, (rota['polyline'], rota['distancia_km'], rota['duracao_min'], carga_id))
                conn2.commit()
                cur2.close(); conn2.close()
            except Exception as e:
                ors_warn = f'ORS falhou: {e}'
        elif olat is None:
            ors_warn = 'Origem sem coordenadas IBGE (cidade não encontrada)'
        else:
            ors_warn = 'Sem pontos suficientes com coordenadas IBGE para a rota'

        if ors_warn:
            warnings.append(ors_warn)

        return jsonify({'ok': True, 'id': carga_id, 'numero': numero, 'warnings': warnings})

    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({'ok': False, 'error': str(e)}), 500


def _aplica_periodo_where(args, where, params):
    """Adiciona filtro de período em where/params conforme data_campo:
    'carregamento'|'previsao' filtram colunas da carga; 'agendamento' faz EXISTS
    em embarques_cargas_destinos. Datas em UTC; date-only vira [00:00, 23:59:59]."""
    campo = args.get('data_campo', 'carregamento')
    start, end = args.get('start'), args.get('end')
    if campo == 'agendamento':
        cond = ["d.carga_id = c.id", "d.data_agendamento IS NOT NULL"]
        if start: cond.append("d.data_agendamento >= %s")
        if end:   cond.append("d.data_agendamento <= %s")
        where.append("EXISTS (SELECT 1 FROM embarques_cargas_destinos d WHERE " + " AND ".join(cond) + ")")
        if start: params.append(start if len(start) > 10 else start + ' 00:00:00')
        if end:   params.append(end if len(end) > 10 else end + ' 23:59:59')
    else:
        col = 'data_carregamento' if campo == 'carregamento' else 'previsao_entrega'
        if start:
            where.append(f"c.{col} >= %s"); params.append(start)
        if end:
            where.append(f"c.{col} <= %s"); params.append(end)


# ── Cargas: listagem com filtros ────────────────────────────────────────
@app.route('/api/embarques/cargas')
@login_required
def api_embarques_cargas_list():
    args = request.args
    where = ["1=1"]
    params = []

    _aplica_periodo_where(args, where, params)
    if args.get('tipo_operacao'):
        where.append("c.tipo_operacao = %s"); params.append(args['tipo_operacao'])
    if args.get('cliente_id'):
        where.append("c.cliente_id = %s"); params.append(args['cliente_id'])
    if args.get('criado_por_id'):
        where.append("c.criado_por_id = %s"); params.append(args['criado_por_id'])
    if args.get('motorista'):
        where.append("c.motorista_nome ILIKE %s"); params.append(f"%{args['motorista']}%")
    if args.get('origem_uf'):
        where.append("c.origem_uf = %s"); params.append(args['origem_uf'])
    if args.get('destino_uf'):
        where.append("EXISTS (SELECT 1 FROM embarques_cargas_destinos d WHERE d.carga_id = c.id AND d.uf = %s)")
        params.append(args['destino_uf'])
    if args.get('status'):
        where.append("c.status = %s"); params.append(args['status'])
    if args.get('viagem_vazia') in ('1', '0'):
        where.append("c.viagem_vazia = %s"); params.append(args['viagem_vazia'] == '1')
    if args.get('q'):
        q = f"%{args['q']}%"
        where.append("(c.numero ILIKE %s OR c.motorista_nome ILIKE %s OR c.cliente_nome ILIKE %s OR c.cavalo_placa ILIKE %s OR c.carreta1_placa ILIKE %s OR c.carreta2_placa ILIKE %s)")
        params.extend([q, q, q, q, q, q])

    try:
        limite = int(args.get('limit', 1000))
    except Exception:
        limite = 1000
    limite = max(1, min(limite, 1000))

    sql = f"""
        SELECT c.id, c.numero, c.status, c.tipo_operacao, c.viagem_vazia,
               c.cliente_id, c.cliente_nome,
               c.origem_cidade, c.origem_uf,
               c.motorista_nome, c.motorista_cpf,
               c.cavalo_placa, c.cavalo_tipo, c.cavalo_marca_modelo, c.cavalo_proprietario,
               c.carreta1_placa, c.carreta2_placa,
               c.data_carregamento, c.previsao_entrega, c.data_conclusao,
               c.observacoes,
               c.criado_em, c.criado_por_id, c.criado_por_nome, c.atualizado_em,
               c.no_local_desde, c.saida_auto, c.entregue_auto, c.data_saida_real,
               c.distancia_planejada_km, c.duracao_estimada_min,
               c.desengatada_em, c.descarga_motorista_nome, c.descarga_cavalo_placa,
               (SELECT EXTRACT(EPOCH FROM ((NOW() AT TIME ZONE 'UTC') - pa.data_posicao)) / 3600.0
                  FROM embarques_posicoes_atuais pa WHERE pa.placa = c.carreta1_placa) AS rastreio_carreta_idade_h,
               (
                 SELECT string_agg(d.cidade || '/' || d.uf, '; ' ORDER BY d.ordem)
                 FROM embarques_cargas_destinos d WHERE d.carga_id = c.id
               ) AS destinos,
               (
                 SELECT string_agg(rt.cidade || '/' || rt.uf, ', ' ORDER BY rt.ordem)
                 FROM embarques_cargas_rota rt WHERE rt.carga_id = c.id
               ) AS rota_resumo,
               (SELECT d.data_agendamento FROM embarques_cargas_destinos d
                 WHERE d.carga_id = c.id ORDER BY d.ordem DESC LIMIT 1) AS agendamento_final,
               (c.status IN ('Aberta','Em rota') AND EXISTS (
                 SELECT 1 FROM embarques_cargas_destinos d
                 WHERE d.carga_id = c.id AND d.data_agendamento IS NOT NULL
                   AND d.data_agendamento < (NOW() AT TIME ZONE 'UTC')
               )) AS tem_agendamento_vencido,
               (SELECT json_agg(json_build_object(
                   'ordem', d.ordem, 'cidade', d.cidade, 'uf', d.uf,
                   'data_agendamento', d.data_agendamento) ORDER BY d.ordem)
                 FROM embarques_cargas_destinos d WHERE d.carga_id = c.id) AS agendamentos_destinos
        FROM embarques_cargas c
        WHERE {' AND '.join(where)}
        ORDER BY c.data_carregamento DESC, c.id DESC
        LIMIT {limite}
    """

    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute(sql, params)
        cols = [d[0] for d in cur.description]
        rows = cur.fetchall()
        data = []
        for r in rows:
            obj = dict(zip(cols, r))
            for k in ('data_carregamento', 'previsao_entrega'):
                if obj.get(k): obj[k] = obj[k].isoformat()
            for k in ('data_conclusao', 'criado_em', 'atualizado_em'):
                if obj.get(k): obj[k] = obj[k].isoformat()
            if obj.get('data_saida_real'):
                obj['data_saida_real'] = obj['data_saida_real'].isoformat() + 'Z'
            if obj.get('desengatada_em'):
                obj['desengatada_em'] = obj['desengatada_em'].isoformat() + 'Z'
            # Agendamento (UTC) — marca com 'Z' p/ o front converter pra local
            if obj.get('agendamento_final'):
                obj['agendamento_final'] = obj['agendamento_final'].isoformat() + 'Z'
            ag = obj.get('agendamentos_destinos')
            if isinstance(ag, str):
                ag = json.loads(ag); obj['agendamentos_destinos'] = ag
            if ag:
                for dd in ag:
                    if dd.get('data_agendamento'):
                        dd['data_agendamento'] = str(dd['data_agendamento']).replace(' ', 'T').rstrip('Z') + 'Z'
            obj['pode_editar'] = _pode_editar_carga(obj.get('criado_por_id'))
            # Rastreio defasado: carga ativa cuja carreta está sem posição há +X dias.
            idade_h = obj.get('rastreio_carreta_idade_h')
            obj['rastreio_carreta_idade_h'] = round(float(idade_h), 1) if idade_h is not None else None
            obj['rastreio_defasado'] = bool(
                obj.get('status') in ('Em rota', 'No destino', 'Desengatada')
                and idade_h is not None
                and float(idade_h) > RASTREIO_ALERTA_SEM_GPS_DIAS * 24
            )
            data.append(obj)
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Carga: detalhe ──────────────────────────────────────────────────────
@app.route('/api/embarques/cargas/<int:carga_id>')
@login_required
def api_embarques_carga_detail(carga_id):
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("SELECT * FROM embarques_cargas WHERE id = %s", (carga_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Carga não encontrada'}), 404
        cols = [d[0] for d in cur.description]
        carga = dict(zip(cols, row))
        for k in ('data_carregamento', 'previsao_entrega'):
            if carga.get(k): carga[k] = carga[k].isoformat()
        for k in ('data_conclusao', 'criado_em', 'atualizado_em'):
            if carga.get(k): carga[k] = carga[k].isoformat()

        cur.execute("SELECT id, ordem, cidade, uf, data_agendamento FROM embarques_cargas_destinos WHERE carga_id = %s ORDER BY ordem", (carga_id,))
        destinos = [{'id': r[0], 'ordem': r[1], 'cidade': r[2], 'uf': r[3],
                     'data_agendamento': (r[4].isoformat() + 'Z') if r[4] else None} for r in cur.fetchall()]
        carga['destinos'] = destinos
        cur.execute("SELECT ordem, cidade, uf FROM embarques_cargas_rota WHERE carga_id = %s ORDER BY ordem", (carga_id,))
        carga['rota'] = [{'ordem': r[0], 'cidade': r[1], 'uf': r[2]} for r in cur.fetchall()]
        carga['pode_editar'] = _pode_editar_carga(carga.get('criado_por_id'))
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': carga})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Carga: edição com log ───────────────────────────────────────────────
_PATCH_WHITELIST = (
    'status', 'observacoes', 'previsao_entrega', 'data_carregamento', 'viagem_vazia',
    'cliente_id', 'cliente_nome', 'tipo_operacao',
    'motorista_nome', 'motorista_cpf', 'motorista_telefone',
    'cavalo_placa', 'cavalo_tipo', 'cavalo_marca_modelo', 'cavalo_carroceria', 'cavalo_proprietario', 'cavalo_eh_rizza',
    'carreta1_placa', 'carreta1_marca_modelo', 'carreta1_carroceria', 'carreta1_proprietario', 'carreta1_eh_rizza',
    'carreta2_placa', 'carreta2_marca_modelo', 'carreta2_carroceria', 'carreta2_proprietario', 'carreta2_eh_rizza',
    'origem_cidade', 'origem_uf',
)


@app.route('/api/embarques/cargas/<int:carga_id>', methods=['PATCH'])
@login_required
def api_embarques_carga_patch(carga_id):
    b = request.get_json(silent=True) or {}
    campos = {k: b[k] for k in b if k in _PATCH_WHITELIST}
    if not campos and 'destinos' not in b and 'rota' not in b:
        return jsonify({'ok': False, 'error': 'Nada a atualizar'}), 400

    try:
        conn = get_db(); cur = conn.cursor()
        # Estado atual
        cur.execute("SELECT * FROM embarques_cargas WHERE id = %s", (carga_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Carga não encontrada'}), 404
        cols_atuais = [d[0] for d in cur.description]
        atual = dict(zip(cols_atuais, row))

        # Permissão: admin ou criador da carga
        if not _pode_editar_carga(atual.get('criado_por_id')):
            cur.close(); conn.close()
            return jsonify({
                'ok': False,
                'error': 'Você não pode editar esta carga. Apenas quem lançou ou um administrador.'
            }), 403

        # Bloqueio de conflito quando o novo estado fica/continua ativo
        novo_status = campos.get('status', atual.get('status'))
        if novo_status in ('Aberta', 'Em rota', 'No destino'):
            novo_cpf = campos.get('motorista_cpf', atual.get('motorista_cpf'))
            placas_novas = [
                (campos.get('cavalo_placa',   atual.get('cavalo_placa'))   or '').upper().strip(),
                (campos.get('carreta1_placa', atual.get('carreta1_placa')) or '').upper().strip(),
                (campos.get('carreta2_placa', atual.get('carreta2_placa')) or '').upper().strip(),
            ]
            conflitos = _buscar_conflitos(novo_cpf, [p for p in placas_novas if p], exclude_id=carga_id)
            if conflitos:
                cur.close(); conn.close()
                nums = sorted({c['numero'] for c in conflitos if c.get('numero')})
                msg = 'Recurso já em uso na(s) carga(s) ativa(s): ' + ', '.join(nums)
                return jsonify({'ok': False, 'error': msg, 'tipo': 'conflito', 'conflitos': conflitos}), 409

        # Diff: só campos cujo valor mudou
        diffs = []
        sets = []
        params = []
        for k, v in campos.items():
            antigo = atual.get(k)
            if isinstance(antigo, (bool,)):
                novo_norm = bool(v)
            elif hasattr(antigo, 'isoformat'):
                novo_norm = v  # comparar como veio
                antigo = antigo.isoformat() if antigo else None
            else:
                novo_norm = v
            if str(antigo) != str(novo_norm) and not (antigo is None and novo_norm in (None, '')):
                diffs.append((k, antigo, novo_norm))
                sets.append(f"{k} = %s")
                params.append(novo_norm)

        # Auto data_conclusao quando muda para Entregue
        if campos.get('status') == 'Entregue' and atual.get('status') != 'Entregue':
            sets.append("data_conclusao = NOW()")

        # Destinos — substituir lista inteira se fornecida
        destinos_mudaram = False
        novos_destinos = b.get('destinos')
        if isinstance(novos_destinos, list):
            cur.execute(
                "SELECT ordem, cidade, uf, data_agendamento FROM embarques_cargas_destinos WHERE carga_id = %s ORDER BY ordem",
                (carga_id,)
            )
            atuais = [{'ordem': r[0], 'cidade': r[1], 'uf': r[2], 'data_agendamento': r[3]} for r in cur.fetchall()]
            # repr inclui agendamento (granularidade de minuto, ambos em UTC) p/ detectar mudança e logar
            def _ag(v):
                if not v:
                    return ''
                return (v.isoformat() if hasattr(v, 'isoformat') else str(v))[:16]
            atuais_repr = '; '.join(f"{d['cidade']}/{d['uf']}@{_ag(d['data_agendamento'])}" for d in atuais)
            novos_repr  = '; '.join(f"{d.get('cidade','?')}/{d.get('uf','?')}@{_ag(d.get('data_agendamento'))}" for d in novos_destinos)
            if atuais_repr != novos_repr:
                destinos_mudaram = True
                import geocoding
                cur.execute("DELETE FROM embarques_cargas_destinos WHERE carga_id = %s", (carga_id,))
                for i, d in enumerate(novos_destinos, start=1):
                    if not d.get('cidade') or not d.get('uf'):
                        continue
                    # Re-geocoda (senão lat/lng ficariam NULL e o tracking quebra)
                    dlat, dlng = geocoding.geocoder_municipio(d['cidade'], d['uf'], conn=conn)
                    cur.execute(
                        "INSERT INTO embarques_cargas_destinos (carga_id, ordem, cidade, uf, latitude, longitude, data_agendamento) "
                        "VALUES (%s, %s, %s, %s, %s, %s, %s)",
                        (carga_id, i, d['cidade'], d['uf'], dlat, dlng, d.get('data_agendamento') or None)
                    )

        # Cidades de rota — substituir lista inteira se fornecida
        rota_mudou = False
        rota_ant_repr = rota_nov_repr = ''
        nova_rota = b.get('rota')
        if isinstance(nova_rota, list):
            cur.execute("SELECT cidade, uf FROM embarques_cargas_rota WHERE carga_id = %s ORDER BY ordem", (carga_id,))
            rota_ant_repr = ', '.join(f"{r[0]}/{r[1]}" for r in cur.fetchall())
            rota_nov_repr = ', '.join(f"{r.get('cidade','?')}/{r.get('uf','?')}" for r in nova_rota if r.get('cidade') and r.get('uf'))
            if rota_ant_repr != rota_nov_repr:
                rota_mudou = True
                import geocoding
                cur.execute("DELETE FROM embarques_cargas_rota WHERE carga_id = %s", (carga_id,))
                ordem = 0
                for r in nova_rota:
                    if not r.get('cidade') or not r.get('uf'):
                        continue
                    ordem += 1
                    rlat, rlng = geocoding.geocoder_municipio(r['cidade'], r['uf'], conn=conn)
                    cur.execute(
                        "INSERT INTO embarques_cargas_rota (carga_id, ordem, cidade, uf, latitude, longitude) "
                        "VALUES (%s, %s, %s, %s, %s, %s)",
                        (carga_id, ordem, r['cidade'], r['uf'], rlat, rlng)
                    )

        # Mudou rota/destinos? zera a polyline pro worker recalcular a rota completa.
        precisa_recalc = destinos_mudaram or rota_mudou
        if precisa_recalc:
            sets.append("rota_planejada_polyline = NULL")
            sets.append("rota_recalculada_em = NULL")

        if sets:
            sets.append("atualizado_em = NOW()")
            params.append(carga_id)
            cur.execute(f"UPDATE embarques_cargas SET {', '.join(sets)} WHERE id = %s", params)
            for campo, va, vn in diffs:
                cur.execute("""
                    INSERT INTO embarques_cargas_log
                    (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """, (carga_id, session.get('user_id'), session.get('nome'),
                      campo,
                      None if va is None else str(va),
                      None if vn is None else str(vn)))
            if destinos_mudaram:
                cur.execute("""
                    INSERT INTO embarques_cargas_log
                    (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
                    VALUES (%s, %s, %s, 'destinos', %s, %s)
                """, (carga_id, session.get('user_id'), session.get('nome'),
                      atuais_repr, novos_repr))
            if rota_mudou:
                cur.execute("""
                    INSERT INTO embarques_cargas_log
                    (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
                    VALUES (%s, %s, %s, 'rota', %s, %s)
                """, (carga_id, session.get('user_id'), session.get('nome'),
                      rota_ant_repr, rota_nov_repr))

        conn.commit()
        cur.close(); conn.close()
        total_alt = len(diffs) + (1 if destinos_mudaram else 0) + (1 if rota_mudou else 0)
        return jsonify({'ok': True, 'alteracoes': total_alt})
    except Exception as e:
        try: conn.rollback()
        except Exception: pass
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Carga: log de edição ────────────────────────────────────────────────
@app.route('/api/embarques/cargas/<int:carga_id>/log')
@login_required
def api_embarques_carga_log(carga_id):
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT id, usuario_id, usuario_nome, editado_em, campo, valor_anterior, valor_novo
            FROM embarques_cargas_log
            WHERE carga_id = %s
            ORDER BY editado_em DESC, id DESC
        """, (carga_id,))
        data = [{
            'id': r[0], 'usuario_id': r[1], 'usuario_nome': r[2],
            'editado_em': r[3].isoformat() if r[3] else None,
            'campo': r[4], 'valor_anterior': r[5], 'valor_novo': r[6],
        } for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Carga: desengate de carreta carregada (drop-and-hook) ───────────────
@app.route('/api/embarques/cargas/<int:carga_id>/desengatar', methods=['POST'])
@login_required
def api_embarques_carga_desengatar(carga_id):
    """Desengata o cavalo+motorista; a carreta carregada segue no destino aguardando
    descarga. Libera cavalo+motorista para nova carga (conflito passa a ignorá-los),
    mantém a carreta comprometida e deixa a carga pronta p/ finalizar automático quando
    a carreta sair do destino. Substituto (cavalo/motorista) é opcional, p/ registro."""
    b = request.get_json(silent=True) or {}
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT status, criado_por_id, no_local_desde
            FROM embarques_cargas WHERE id = %s
        """, (carga_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Carga não encontrada'}), 404
        status_atual, criado_por_id, no_local_desde = row

        if not _pode_editar_carga(criado_por_id):
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Você não pode desengatar esta carga. Apenas quem lançou ou um administrador.'}), 403

        if status_atual not in ('Em rota', 'No destino'):
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': f'Só é possível desengatar uma carga "Em rota" ou "No destino" (status atual: {status_atual}).'}), 400

        desc_mot = (b.get('descarga_motorista_nome') or '').strip() or None
        desc_cav = (b.get('descarga_cavalo_placa') or '').strip().upper() or None
        obs = (b.get('observacao') or '').strip()

        # Marca o desengate. NÃO força no_local_desde: se a carreta ainda não chegou
        # (desengate "Em rota" com GPS atrasado), deixa o worker detectar a chegada e só
        # então finalizar na saída — evita finalização falsa quando a carreta está longe.
        # Timestamps gravados como UTC naive (AT TIME ZONE 'UTC') — corretos qualquer que
        # seja o fuso da sessão do Postgres (local em São Paulo, produção em UTC).
        cur.execute("""
            UPDATE embarques_cargas
            SET status = 'Desengatada',
                desengatada_em = (NOW() AT TIME ZONE 'UTC'),
                desengatada_por_id = %s,
                desengatada_por_nome = %s,
                descarga_motorista_nome = %s,
                descarga_cavalo_placa = %s,
                atualizado_em = NOW()
            WHERE id = %s
        """, (session.get('user_id'), session.get('nome'), desc_mot, desc_cav, carga_id))

        # Log do evento (campo 'desengate' p/ aparecer no histórico)
        partes = []
        if desc_cav: partes.append(f'cavalo {desc_cav}')
        if desc_mot: partes.append(f'motorista {desc_mot}')
        if obs: partes.append(obs)
        detalhe = '; '.join(partes) if partes else 'cavalo+motorista liberados'
        cur.execute("""
            INSERT INTO embarques_cargas_log
            (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
            VALUES (%s, %s, %s, 'desengate', %s, %s)
        """, (carga_id, session.get('user_id'), session.get('nome'),
              status_atual, detalhe))

        conn.commit()
        cur.close(); conn.close()
        return jsonify({'ok': True, 'status': 'Desengatada'})
    except Exception as e:
        try: conn.rollback(); conn.close()
        except Exception: pass
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Cargas: CSV streaming ───────────────────────────────────────────────
@app.route('/api/embarques/cargas/csv')
@login_required
def api_embarques_cargas_csv():
    args = request.args
    where = ["1=1"]
    params = []
    _aplica_periodo_where(args, where, params)
    if args.get('tipo_operacao'):
        where.append("c.tipo_operacao = %s"); params.append(args['tipo_operacao'])
    if args.get('cliente_id'):
        where.append("c.cliente_id = %s"); params.append(args['cliente_id'])
    if args.get('criado_por_id'):
        where.append("c.criado_por_id = %s"); params.append(args['criado_por_id'])
    if args.get('motorista'):
        where.append("c.motorista_nome ILIKE %s"); params.append(f"%{args['motorista']}%")
    if args.get('origem_uf'):
        where.append("c.origem_uf = %s"); params.append(args['origem_uf'])
    if args.get('destino_uf'):
        where.append("EXISTS (SELECT 1 FROM embarques_cargas_destinos d WHERE d.carga_id = c.id AND d.uf = %s)")
        params.append(args['destino_uf'])
    if args.get('status'):
        where.append("c.status = %s"); params.append(args['status'])
    if args.get('viagem_vazia') in ('1', '0'):
        where.append("c.viagem_vazia = %s"); params.append(args['viagem_vazia'] == '1')

    nome = f"cargas_{args.get('start','')}_{args.get('end','')}.csv".strip('_')

    sql = f"""
        SELECT c.numero, c.data_carregamento, c.previsao_entrega, c.status, c.tipo_operacao,
               CASE WHEN c.viagem_vazia THEN 'Sim' ELSE 'Não' END AS viagem_vazia,
               c.cliente_nome,
               c.origem_cidade || '/' || c.origem_uf AS origem,
               (SELECT string_agg(d.cidade || '/' || d.uf, '; ' ORDER BY d.ordem)
                FROM embarques_cargas_destinos d WHERE d.carga_id = c.id) AS destinos,
               (SELECT string_agg(rt.cidade || '/' || rt.uf, '; ' ORDER BY rt.ordem)
                FROM embarques_cargas_rota rt WHERE rt.carga_id = c.id) AS rota_via,
               c.motorista_nome, c.motorista_cpf,
               c.cavalo_placa, c.cavalo_marca_modelo, c.cavalo_proprietario,
               c.carreta1_placa, c.carreta1_proprietario,
               c.carreta2_placa, c.carreta2_proprietario,
               c.observacoes,
               c.criado_por_nome, c.criado_em
        FROM embarques_cargas c
        WHERE {' AND '.join(where)}
        ORDER BY c.data_carregamento DESC, c.id DESC
    """
    headers_csv = [
        'numero', 'data_carregamento', 'previsao_entrega', 'status', 'tipo_operacao',
        'viagem_vazia', 'cliente', 'origem', 'destinos', 'rota_via',
        'motorista', 'motorista_cpf',
        'cavalo_placa', 'cavalo_marca_modelo', 'cavalo_proprietario',
        'carreta1_placa', 'carreta1_proprietario',
        'carreta2_placa', 'carreta2_proprietario',
        'observacoes', 'lancado_por', 'lancado_em'
    ]

    def gerar():
        yield '﻿'
        yield _csv_linha(headers_csv)
        conn = get_db(); cur = conn.cursor()
        cur.execute(sql, params)
        for r in cur.fetchall():
            valores = []
            for v in r:
                if hasattr(v, 'isoformat'):
                    valores.append(v.isoformat())
                else:
                    valores.append(v)
            yield _csv_linha(valores)
        cur.close(); conn.close()

    return Response(
        stream_with_context(gerar()),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="{nome}"'}
    )


# ── KPIs da landing ─────────────────────────────────────────────────────
@app.route('/api/embarques/kpis')
@login_required
def api_embarques_kpis():
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT
              COUNT(*) FILTER (WHERE data_carregamento = (NOW() AT TIME ZONE 'America/Sao_Paulo')::date) AS hoje,
              COUNT(*) FILTER (WHERE status = 'Em rota')                 AS em_rota,
              COUNT(*) FILTER (WHERE status = 'No destino')              AS no_destino,
              COUNT(*) FILTER (WHERE status = 'Entregue'
                               AND date_trunc('month', data_conclusao) = date_trunc('month', (NOW() AT TIME ZONE 'America/Sao_Paulo')::date)) AS entregues_mes,
              COUNT(*) FILTER (WHERE status = 'Aberta')                 AS abertas,
              COUNT(*) FILTER (WHERE status = 'Desengatada')            AS desengatadas
            FROM embarques_cargas
        """)
        r = cur.fetchone()
        cur.close(); conn.close()
        return jsonify({
            'ok': True,
            'data': {
                'hoje':           r[0] or 0,
                'em_rota':        r[1] or 0,
                'no_destino':     r[2] or 0,
                'entregues_mes':  r[3] or 0,
                'abertas':        r[4] or 0,
                'desengatadas':   r[5] or 0,
            }
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ── Lista de embarcadores (usuários do sistema) p/ filtro do relatório ─
@app.route('/api/embarques/embarcadores')
@login_required
def api_embarques_embarcadores():
    try:
        conn = get_db(); cur = conn.cursor()
        cur.execute("""
            SELECT DISTINCT criado_por_id, criado_por_nome
            FROM embarques_cargas
            WHERE criado_por_id IS NOT NULL AND criado_por_nome IS NOT NULL
            ORDER BY criado_por_nome
        """)
        data = [{'id': r[0], 'nome': r[1]} for r in cur.fetchall()]
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


# ════════════════════════════════════════════════════════════════════════════
# RASTREAMENTO — Endpoints API (todos @login_required)
# ════════════════════════════════════════════════════════════════════════════

import tres_s_client
import rastreamento_worker
import embarques_auto

KM_DIA_PADRAO = int(os.getenv('KM_DIA_PADRAO', '600'))
# Alerta de rastreio defasado: carga ativa cuja carreta está sem posição há +X dias.
RASTREIO_ALERTA_SEM_GPS_DIAS = float(os.getenv('RASTREAMENTO_ALERTA_SEM_GPS_DIAS', '2'))


def eta_realista(distancia_km, partida_dt, duracao_ors_min=None, km_dia=KM_DIA_PADRAO):
    """Chegada estimada com o ritmo do motorista, RATEADO (sem arredondar):
    km_dia (600) / 24h = 25 km/h  ->  horas = distancia / (km_dia/24).
    Ex.: 606 km -> 24,2h; 2789 km -> ~111,6h. Trabalha em UTC (partida_dt naive UTC).
    (duracao_ors_min fica na assinatura por compatibilidade; não é mais usado.)"""
    from datetime import timedelta
    if not distancia_km or not partida_dt or not km_dia:
        return None
    horas = float(distancia_km) / (km_dia / 24.0)   # 600 km/dia = 25 km/h
    return partida_dt + timedelta(hours=horas)


def _decode_polyline(s, precision=5):
    """Decodifica polyline (algoritmo Google/ORS) → lista de (lat, lng)."""
    if not s:
        return []
    coords = []; index = lat = lng = 0; factor = 10 ** precision
    while index < len(s):
        for alvo in range(2):
            shift = result = 0
            while True:
                b = ord(s[index]) - 63; index += 1
                result |= (b & 0x1f) << shift; shift += 5
                if b < 0x20:
                    break
            delta = ~(result >> 1) if (result & 1) else (result >> 1)
            if alvo == 0: lat += delta
            else: lng += delta
        coords.append((lat / factor, lng / factor))
    return coords


def _km_restante(polyline_enc, pos_lat, pos_lng):
    """Distância restante (km) ao longo da rota a partir da posição atual:
    projeta no vértice mais próximo da rota e soma os segmentos até o destino."""
    if pos_lat is None or pos_lng is None:
        return None
    pts = _decode_polyline(polyline_enc)
    if len(pts) < 2:
        return None
    import geocoding
    best_i, best_d = 0, None
    for i, (la, ln) in enumerate(pts):
        d = geocoding.km_entre(pos_lat, pos_lng, la, ln)
        if d is not None and (best_d is None or d < best_d):
            best_d, best_i = d, i
    rem = best_d or 0.0
    for i in range(best_i, len(pts) - 1):
        seg = geocoding.km_entre(pts[i][0], pts[i][1], pts[i + 1][0], pts[i + 1][1])
        if seg:
            rem += seg
    return round(rem, 1)


# Raio (km) que conta como "dentro da cidade do destino" — por POSIÇÃO, nunca por nome.
RAIO_CHEGADA_DESTINO_KM = float(os.getenv('RASTREAMENTO_RAIO_CHEGADA_DESTINO', '20'))


def _indice_chegada_destino(traj, dest_lat, dest_lng):
    """Índice do ponto de CHEGADA no destino, decidido por POSIÇÃO (ignora o nome que o 3S
    reporta). Escada:
      1) primeiro ponto dentro do raio que INICIA uma parada de >= CHEGADA_MIN_PARADO
         (parado = vel <= PARADO_KMH), permanecendo dentro do raio → o local da descarga;
      2) senão, o ponto de MAIOR APROXIMAÇÃO do destino dentro do raio (descarga rápida <1h
         ou buraco de GPS no destino).
    Retorna None se nenhum ponto entrou no raio (não corta → deixa o trajeto inteiro)."""
    if not traj or dest_lat is None or dest_lng is None:
        return None
    import geocoding as _geo
    from datetime import datetime as _dtc
    def _t(s):
        try:
            return _dtc.fromisoformat(str(s).replace('Z', ''))
        except Exception:
            return None
    parado_kmh = getattr(rastreamento_worker, 'PARADO_KMH', 3)
    parado_min = getattr(rastreamento_worker, 'CHEGADA_MIN_PARADO', 60)
    dest_lat, dest_lng = float(dest_lat), float(dest_lng)
    dentro = []  # (idx, distância_km) dos pontos dentro do raio do destino
    for i, p in enumerate(traj):
        d = _geo.km_entre(p['lat'], p['lng'], dest_lat, dest_lng)
        if d is not None and d <= RAIO_CHEGADA_DESTINO_KM:
            dentro.append((i, d))
    if not dentro:
        return None
    dentro_set = {i for i, _ in dentro}
    n = len(traj)
    # Regra 1: primeira parada de >= parado_min dentro do raio.
    for i, _d in dentro:
        if (traj[i].get('velocidade') or 0) > parado_kmh:
            continue
        t0 = _t(traj[i].get('data'))
        if t0 is None:
            continue
        j = i
        while j < n and j in dentro_set and (traj[j].get('velocidade') or 0) <= parado_kmh:
            tj = _t(traj[j].get('data'))
            if tj is not None and (tj - t0).total_seconds() >= parado_min * 60:
                return i  # onde ele parou (início da parada)
            j += 1
    # Regra 2: maior aproximação do destino dentro do raio.
    return min(dentro, key=lambda x: x[1])[0]


def _kpi_ao_vivo(traj):
    """KPIs calculados ao vivo a partir dos pontos do trajeto da viagem (só da viagem,
    pois traj já vem filtrado por data_saida_real). Espelha _consolidar_kpi do worker."""
    base = {'distancia_km': 0.0, 'velocidade_max': 0, 'velocidade_media': None,
            'tempo_movimento_seg': 0, 'tempo_parado_seg': 0, 'consolidado_final': False}
    if not traj:
        return base
    import geocoding
    from datetime import datetime as _d

    def _pt(s):
        try: return _d.fromisoformat(str(s).replace('Z', ''))
        except Exception: return None

    total_m = 0.0; vmax = 0; vsum = 0; vn = 0; tmov = 0; tpar = 0
    for i in range(len(traj) - 1):
        a, b = traj[i], traj[i + 1]
        seg = geocoding.km_entre(a['lat'], a['lng'], b['lat'], b['lng'])
        if seg is not None:
            total_m += seg * 1000
        av = a.get('velocidade')
        if av is not None:
            vmax = max(vmax, int(av)); vsum += int(av); vn += 1
        da, db = _pt(a['data']), _pt(b['data'])
        delta = (db - da).total_seconds() if (da and db) else 0
        if (av or 0) > 3: tmov += delta
        else: tpar += delta
    # vel do último ponto também conta pro máximo
    lv = traj[-1].get('velocidade')
    if lv is not None:
        vmax = max(vmax, int(lv))
    base.update({
        'distancia_km': round(total_m / 1000, 1),
        'velocidade_max': vmax,
        'velocidade_media': round(vsum / vn, 1) if vn else None,
        'tempo_movimento_seg': int(tmov),
        'tempo_parado_seg': int(tpar),
    })
    return base


# Normalização de placa em SQL — 5º caractere dígito vira letra (0=A … 9=J), a
# mesma regra de `placas.mercosul()`. Existe porque o JOIN do mapa geral compara
# placa-da-carga com placa-da-3S dentro do SQL, e ali não dá para chamar Python:
# o sync grava a grafia CRUA da 3S (42 das 94 vêm na grafia antiga) enquanto a
# carga do robô está em Mercosul, então a igualdade exata perdia o vínculo e o
# veículo aparecia como "sem carga" no mapa.
def _pn(coluna):
    """Expressão SQL que normaliza `coluna` para a grafia Mercosul."""
    return ("CASE WHEN substring(upper(trim({c})) from 5 for 1) BETWEEN '0' AND '9' "
            "THEN overlay(upper(trim({c})) placing "
            "chr(65 + substring(upper(trim({c})) from 5 for 1)::int) from 5 for 1) "
            "ELSE upper(trim({c})) END").format(c=coluna)


@app.route('/api/rastreamento/posicoes')
@login_required
def api_rastreamento_posicoes():
    """Lista posições atuais com info da carga ativa (se houver).
    Filtros: carregado=1|0, eh_rizza=1, q (placa/motorista).
    """
    args = request.args
    where = ["1=1"]
    params = []

    carregado = args.get('carregado')
    eh_rizza = args.get('eh_rizza')
    q = (args.get('q') or '').strip()

    base_join = f"""
        FROM embarques_posicoes_atuais p
        LEFT JOIN embarques_veiculos_rastreio v ON v.placa = p.placa
        LEFT JOIN LATERAL (
            SELECT id, numero, status, cliente_nome, motorista_nome, cavalo_proprietario, cavalo_eh_rizza,
                   no_local_desde, saida_auto, entregue_auto, data_carregamento, origem_cidade, origem_uf,
                   cavalo_placa, carreta1_placa, carreta2_placa
            FROM embarques_cargas c
            WHERE ({_pn('c.cavalo_placa')} = {_pn('p.placa')}
                OR {_pn('c.carreta1_placa')} = {_pn('p.placa')}
                OR {_pn('c.carreta2_placa')} = {_pn('p.placa')})
              AND c.status IN ('Aberta','Em rota','No destino','Desengatada')
            ORDER BY c.id DESC LIMIT 1
        ) ca ON true
    """

    if carregado == '1':
        where.append("ca.id IS NOT NULL")
    elif carregado == '0':
        where.append("ca.id IS NULL")

    if eh_rizza == '1':
        where.append("(ca.cavalo_eh_rizza = TRUE OR v.frota ILIKE %s)")
        params.append('%RIZZA%')

    if q:
        where.append("(p.placa ILIKE %s OR ca.motorista_nome ILIKE %s)")
        params.extend([f'%{q}%', f'%{q}%'])

    sql = f"""
        SELECT p.placa, p.latitude, p.longitude, p.velocidade, p.ignicao, p.direcao,
               p.cidade, p.uf, p.data_posicao, p.bloqueio, p.atualizado_em,
               v.frota, v.modelo, v.tipo,
               ca.id AS carga_id, ca.numero, ca.status, ca.cliente_nome, ca.motorista_nome,
               ca.cavalo_proprietario, ca.cavalo_eh_rizza, ca.no_local_desde, ca.saida_auto,
               ca.origem_cidade, ca.origem_uf,
               CASE
                 WHEN {_pn('ca.cavalo_placa')} = {_pn('p.placa')} THEN 'Cavalo'
                 WHEN {_pn('ca.carreta1_placa')} = {_pn('p.placa')} THEN 'Carreta 1'
                 WHEN {_pn('ca.carreta2_placa')} = {_pn('p.placa')} THEN 'Carreta 2'
                 ELSE v.tipo
               END AS papel
        {base_join}
        WHERE {' AND '.join(where)}
        ORDER BY p.placa
        LIMIT 2000
    """

    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute(sql, params)
        rows = cur.fetchall()
        cols = [c[0] for c in cur.description]
        data = [dict(zip(cols, r)) for r in rows]
        # Normalização: lat/lng → float, data → ISO UTC (com Z)
        for d in data:
            d['latitude'] = float(d['latitude']) if d['latitude'] is not None else None
            d['longitude'] = float(d['longitude']) if d['longitude'] is not None else None
            for k in ('data_posicao', 'atualizado_em', 'no_local_desde'):
                if d.get(k) is not None:
                    d[k] = d[k].isoformat() + 'Z'
            d['carregado'] = d.get('carga_id') is not None
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/rastreamento/cargas/<int:carga_id>/trajeto')
@login_required
def api_rastreamento_trajeto(carga_id):
    """Retorna trajeto + rota planejada + KPIs + raios."""
    try:
        conn = get_db()
        cur = conn.cursor()

        cur.execute("""
            SELECT id, numero, status, cliente_nome, motorista_nome, cavalo_placa,
                   carreta1_placa, carreta2_placa,
                   origem_cidade, origem_uf, origem_latitude, origem_longitude,
                   data_carregamento, data_saida_real, data_conclusao,
                   no_local_desde, saida_auto, entregue_auto,
                   rota_planejada_polyline, distancia_planejada_km, duracao_estimada_min,
                   rota_recalculada_em, inicio_viagem,
                   desengatada_em, descarga_motorista_nome, descarga_cavalo_placa
            FROM embarques_cargas WHERE id=%s
        """, (carga_id,))
        r = cur.fetchone()
        if not r:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Carga não encontrada'}), 404
        cols = [c[0] for c in cur.description]
        carga = dict(zip(cols, r))

        cur.execute("""
            SELECT ordem, cidade, uf, latitude, longitude, data_agendamento
            FROM embarques_cargas_destinos WHERE carga_id=%s ORDER BY ordem
        """, (carga_id,))
        destinos = []
        for ord_, cidade, uf, lat, lng, ag in cur.fetchall():
            destinos.append({
                'ordem': ord_, 'cidade': cidade, 'uf': uf,
                'latitude': float(lat) if lat is not None else None,
                'longitude': float(lng) if lng is not None else None,
                'data_agendamento': (ag.isoformat() + 'Z') if ag else None,
            })

        # Período: busca LARGO por DATA (data_carregamento c/ folga) — NÃO por inicio_viagem,
        # que é por nome de cidade e o 3S mente (etiqueta a origem a 100+ km). O recorte por
        # distância (abaixo) define o começo real perto do pátio.
        from datetime import datetime as _dt, timedelta as _td, time as _time
        _dcarr = carga.get('data_carregamento')
        if _dcarr:
            _base = _dcarr if isinstance(_dcarr, _dt) else _dt.combine(_dcarr, _time())
            inicio = _base - _td(hours=12)
        else:
            inicio = carga.get('inicio_viagem') or (_dt.utcnow() - _td(days=15))
        # Entregue: busca LARGO (até a conclusão). O recorte "até a chegada" é feito
        # depois por DISTÂNCIA ao destino — NÃO por no_local_desde, que o 3S pode ter
        # cravado cedo (etiqueta uma posição a 100+ km com o nome da cidade-destino), o
        # que fazia o trajeto "voltar" ao finalizar. Em andamento segue ao vivo (agora).
        if carga.get('status') == 'Entregue':
            fim = carga.get('data_conclusao') or carga.get('no_local_desde') or _dt.utcnow()
        else:
            fim = _dt.utcnow()

        def _buscar_trajeto(placa):
            if not placa:
                return []
            # ANY(grafias): a posição é gravada na grafia CRUA da 3S (42 das 94
            # placas vêm na antiga) e a carga pode estar em Mercosul. Com
            # igualdade exata o trajeto vinha vazio e o mapa mostrava 0 km.
            cur.execute("""
                SELECT data_posicao, latitude, longitude, velocidade, ignicao, cidade, uf
                FROM embarques_posicoes_historico
                WHERE placa = ANY(%s) AND data_posicao BETWEEN %s AND %s
                ORDER BY data_posicao
            """, (placas.grafias(placa), inicio, fim))
            return [{
                'data': dp.isoformat() + 'Z',
                'lat': float(la),
                'lng': float(ln),
                'velocidade': vel,
                'ignicao': ig,
                'cidade': cid,
                'uf': uff,
            } for (dp, la, ln, vel, ig, cid, uff) in cur.fetchall()]

        traj_cavalo = _buscar_trajeto(carga['cavalo_placa'])
        traj_c1 = _buscar_trajeto(carga.get('carreta1_placa')) if carga.get('carreta1_placa') else []
        traj_c2 = _buscar_trajeto(carga.get('carreta2_placa')) if carga.get('carreta2_placa') else []

        # Recorta o trecho PRÉ-origem (caminhão já rodando antes do lançamento) — a linha e o
        # KPI passam a começar na saída da origem, não antes.
        _olat, _olng = carga.get('origem_latitude'), carga.get('origem_longitude')
        if _olat is not None and _olng is not None:
            import geocoding as _geo
            def _recorta_origem(traj):
                if not traj:
                    return traj
                idx = _geo.indice_saida_origem([(p['lat'], p['lng']) for p in traj],
                                               float(_olat), float(_olng))
                return traj[idx:]
            traj_cavalo = _recorta_origem(traj_cavalo)
            traj_c1 = _recorta_origem(traj_c1)
            traj_c2 = _recorta_origem(traj_c2)

        # Placa de rastreio principal (carreta1 → cavalo → carreta2)
        placa_track = rastreamento_worker._placa_tracking(
            carga['cavalo_placa'], carga.get('carreta1_placa'), carga.get('carreta2_placa'), cur)
        # Comparação NORMALIZADA: `placa_track` vem na grafia da 3S (pode ser a
        # antiga) e a placa da carga pode estar em Mercosul. Comparar cru fazia
        # cair no `else` e exibir o trajeto do CAVALO no lugar do da carreta.
        _pt = placas.mercosul(placa_track or '')
        if placa_track and _pt == placas.mercosul(carga.get('carreta1_placa') or ''):
            rastreado_via = {'placa': placa_track, 'tipo': 'carreta1'}
            traj_principal = traj_c1
        elif placa_track and _pt == placas.mercosul(carga.get('carreta2_placa') or ''):
            rastreado_via = {'placa': placa_track, 'tipo': 'carreta2'}
            traj_principal = traj_c2
        elif placa_track and _pt == placas.mercosul(carga['cavalo_placa'] or ''):
            rastreado_via = {'placa': placa_track, 'tipo': 'cavalo'}
            traj_principal = traj_cavalo
        else:
            rastreado_via = None
            traj_principal = traj_cavalo

        # ── FALLBACK DE EXIBIÇÃO (item 2): se a placa rastreada é a CARRETA e ela está MUDA
        # (sem ponto recente no trajeto), mostra o CAVALO — só exibição, não muda o fechamento.
        # Recorta o trajeto do cavalo na CHEGADA (no_local_desde) p/ contemplar só origem→destino
        # (a próxima viagem do cavalo é cortada, igual já se faz com carga entregue).
        fallback_cavalo = False
        if rastreado_via and rastreado_via['tipo'] in ('carreta1', 'carreta2') \
                and carga.get('status') != 'Desengatada':
            def _idade_h(iso):
                try:
                    return (_dt.utcnow() - _dt.fromisoformat(str(iso).replace('Z', ''))).total_seconds() / 3600.0
                except Exception:
                    return None
            ult_carreta = _idade_h(traj_principal[-1]['data']) if traj_principal else None
            carreta_muda = (not traj_principal) or (ult_carreta is not None and ult_carreta > rastreamento_worker.FRESCOR_H)
            if carreta_muda and traj_cavalo:
                _nld = carga.get('no_local_desde')
                if _nld is not None:
                    traj_cav_carga = [p for p in traj_cavalo
                                      if _dt.fromisoformat(p['data'].replace('Z', '')) <= _nld]
                else:
                    traj_cav_carga = traj_cavalo
                if traj_cav_carga:
                    traj_cavalo = traj_cav_carga          # linha desenhada (trajeto.cavalo)
                    traj_principal = traj_cav_carga
                    rastreado_via = {'placa': carga['cavalo_placa'], 'tipo': 'cavalo',
                                     'fallback_carreta_muda': True}
                    fallback_cavalo = True

        # ── Entregue: corta o desenho no ponto de CHEGADA no destino, decidido por POSIÇÃO
        # (onde ele parou ~1h no destino; senão, maior aproximação dentro do raio) — NÃO pelo
        # horário no_local_desde, que o 3S pode ter cravado a 100+ km por mentir o nome da
        # cidade. Assim o trajeto NUNCA "volta" ao finalizar. Chegada correta fica igual.
        if carga.get('status') == 'Entregue' and destinos:
            _dfin = destinos[-1]
            _dla, _dln = _dfin.get('latitude'), _dfin.get('longitude')
            if _dla is not None and _dln is not None:
                def _corta_chegada(traj):
                    idx = _indice_chegada_destino(traj, _dla, _dln)
                    return traj[:idx + 1] if idx is not None else traj
                traj_cavalo = _corta_chegada(traj_cavalo)
                traj_c1 = _corta_chegada(traj_c1)
                traj_c2 = _corta_chegada(traj_c2)
                traj_principal = _corta_chegada(traj_principal)

        # KPIs já consolidados?
        cur.execute("""
            SELECT distancia_metros, velocidade_max, velocidade_media,
                   tempo_movimento_seg, tempo_parado_seg, consolidado_final
            FROM embarques_cargas_rastreio_kpi WHERE carga_id=%s
        """, (carga_id,))
        rk = cur.fetchone()
        if rk and rk[5] and not fallback_cavalo:
            # KPI final consolidado (carga entregue) — usa o valor persistido
            kpi = {
                'distancia_km': round((rk[0] or 0) / 1000, 1),
                'velocidade_max': rk[1],
                'velocidade_media': float(rk[2]) if rk[2] is not None else None,
                'tempo_movimento_seg': rk[3],
                'tempo_parado_seg': rk[4],
                'consolidado_final': True,
            }
        else:
            # Em viagem: calcula ao vivo a partir do trajeto da placa rastreada (só da viagem)
            kpi = _kpi_ao_vivo(traj_principal)

        cur.close(); conn.close()

        # Última posição = último ponto do trajeto da placa rastreada (carreta primeiro)
        ultima = traj_principal[-1] if traj_principal else None

        # Rota é sempre origem→destino (completa). O que falta é derivado da posição atual
        # projetada nessa rota — assim a linha do mapa fica completa e o "km faltando" certo.
        from datetime import datetime as _dt3, timedelta as _td3
        pos_la = ultima['lat'] if ultima else None
        pos_ln = ultima['lng'] if ultima else None
        km_total = float(carga['distancia_planejada_km']) if carga.get('distancia_planejada_km') is not None else None
        dur_total = carga.get('duracao_estimada_min')
        km_restante = _km_restante(carga.get('rota_planejada_polyline'), pos_la, pos_ln)
        if km_restante is None:
            km_restante = km_total
        dur_restante = None
        if dur_total and km_total:
            dur_restante = max(0, round(dur_total * (km_restante / km_total)))
        agora = _dt3.utcnow()
        eta_chegada_iso = ((agora + _td3(minutes=dur_restante)).isoformat() + 'Z') if dur_restante is not None else None
        _eta = eta_realista(km_restante, agora, dur_restante)
        eta_iso = (_eta.isoformat() + 'Z') if _eta else None

        # Format origem/destinos pra JSON
        origem = {
            'cidade': carga['origem_cidade'], 'uf': carga['origem_uf'],
            'latitude': float(carga['origem_latitude']) if carga['origem_latitude'] is not None else None,
            'longitude': float(carga['origem_longitude']) if carga['origem_longitude'] is not None else None,
        }

        resp = {
            'ok': True,
            'carga': {
                'id': carga['id'],
                'numero': carga['numero'],
                'status': carga['status'],
                'cliente_nome': carga['cliente_nome'],
                'motorista_nome': carga['motorista_nome'],
                'cavalo_placa': carga['cavalo_placa'],
                'carreta1_placa': carga.get('carreta1_placa'),
                'carreta2_placa': carga.get('carreta2_placa'),
                'data_carregamento': carga['data_carregamento'].isoformat() if carga['data_carregamento'] else None,
                'data_saida_real': (carga['data_saida_real'].isoformat() + 'Z') if carga['data_saida_real'] else None,
                'data_conclusao': (carga['data_conclusao'].isoformat() + 'Z') if carga['data_conclusao'] else None,
                'no_local_desde': (carga['no_local_desde'].isoformat() + 'Z') if carga['no_local_desde'] else None,
                'saida_auto': carga['saida_auto'],
                'entregue_auto': carga['entregue_auto'],
                'desengatada_em': (carga['desengatada_em'].isoformat() + 'Z') if carga.get('desengatada_em') else None,
                'descarga_motorista_nome': carga.get('descarga_motorista_nome'),
                'descarga_cavalo_placa': carga.get('descarga_cavalo_placa'),
            },
            'origem': origem,
            'destinos': destinos,
            'trajeto': {
                'cavalo': traj_cavalo,
                'carreta1': traj_c1,
                'carreta2': traj_c2,
            },
            'rota_planejada': {
                'polyline': carga.get('rota_planejada_polyline'),
                'distancia_km': km_total,
                'duracao_min': dur_total,
                'distancia_restante_km': km_restante,
                'duracao_restante_min': dur_restante,
                'recalculada_em': (carga['rota_recalculada_em'].isoformat() + 'Z') if carga.get('rota_recalculada_em') else None,
                'eta_chegada_iso': eta_chegada_iso,
                'eta_realista_iso': eta_iso,
            },
            'ultima_posicao': ultima,
            'rastreado_via': rastreado_via,
            'kpi': kpi,
        }
        return jsonify(resp)
    except Exception as e:
        try: conn.close()
        except Exception: pass
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/rastreamento/cargas/<int:carga_id>/confirmar-entrega', methods=['POST'])
@login_required
def api_rastreamento_confirmar_entrega(carga_id):
    """Confirma manualmente a entrega. status='Entregue', entregue_auto=false.
    Aceita, no corpo (opcional), o cavalo/motorista que efetivou a descarga — usado
    ao finalizar uma carga 'Desengatada' (registro de quem fez o trampo). Se nada for
    enviado, preserva o que já estava (preenchido no desengate)."""
    b = request.get_json(silent=True) or {}
    desc_cav = (b.get('descarga_cavalo_placa') or '').strip().upper() or None
    desc_mot = (b.get('descarga_motorista_nome') or '').strip() or None
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT status FROM embarques_cargas WHERE id=%s", (carga_id,))
        r = cur.fetchone()
        if not r:
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': 'Carga não encontrada'}), 404
        if r[0] in ('Entregue', 'Cancelada'):
            cur.close(); conn.close()
            return jsonify({'ok': False, 'error': f'Carga já está {r[0]}'}), 400
        # COALESCE: só sobrescreve descarga_* quando enviado; senão mantém o do desengate.
        cur.execute("""
            UPDATE embarques_cargas
            SET status='Entregue', entregue_auto=FALSE, data_conclusao=NOW(), atualizado_em=NOW(),
                descarga_cavalo_placa = COALESCE(%s, descarga_cavalo_placa),
                descarga_motorista_nome = COALESCE(%s, descarga_motorista_nome)
            WHERE id=%s
        """, (desc_cav, desc_mot, carga_id))
        # Log no embarques_cargas_log
        cur.execute("""
            INSERT INTO embarques_cargas_log (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
            VALUES (%s, %s, %s, 'status', %s, 'Entregue')
        """, (carga_id, session.get('user_id'), session.get('nome'), r[0]))
        if desc_cav or desc_mot:
            quem = '; '.join(p for p in [
                ('cavalo ' + desc_cav) if desc_cav else '',
                ('motorista ' + desc_mot) if desc_mot else '',
            ] if p)
            cur.execute("""
                INSERT INTO embarques_cargas_log (carga_id, usuario_id, usuario_nome, campo, valor_anterior, valor_novo)
                VALUES (%s, %s, %s, 'descarga', NULL, %s)
            """, (carga_id, session.get('user_id'), session.get('nome'), quem))
        # Consolida KPI
        rastreamento_worker._consolidar_kpi(cur, carga_id, final=True)
        conn.commit()
        cur.close(); conn.close()
        return jsonify({'ok': True})
    except Exception as e:
        try: conn.rollback(); conn.close()
        except Exception: pass
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/rastreamento/sync-veiculos', methods=['POST'])
@admin_required
def api_rastreamento_sync_veiculos():
    """Força sync com /ListaVeiculos da 3S. UPSERT em embarques_veiculos_rastreio."""
    try:
        veiculos = tres_s_client.lista_veiculos()
        conn = get_db()
        cur = conn.cursor()
        novos = 0
        atualizados = 0
        for v in veiculos:
            placa = (v.get('placa') or '').strip().upper()
            id_veiculo = v.get('idVeiculo')
            if not placa or not id_veiculo:
                continue
            # Identidade do veículo é o id_veiculo_3s (a placa pode mudar: antiga -> Mercosul).
            cur.execute("SELECT 1 FROM embarques_veiculos_rastreio WHERE id_veiculo_3s=%s", (id_veiculo,))
            existe = cur.fetchone() is not None
            # Se essa placa estiver presa em OUTRA linha (placa realocada/órfã), libera antes.
            cur.execute(
                "DELETE FROM embarques_veiculos_rastreio WHERE placa=%s AND id_veiculo_3s<>%s",
                (placa, id_veiculo)
            )
            # UPSERT pela identidade do veículo: atualiza a placa no lugar (resolve troca de placa,
            # sem deixar linha órfã com a placa antiga).
            cur.execute("""
                INSERT INTO embarques_veiculos_rastreio
                    (placa, id_veiculo_3s, id_equipamento, frota, modelo, tipo, sincronizado_em)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ON CONFLICT (id_veiculo_3s) DO UPDATE SET
                    placa = EXCLUDED.placa,
                    id_equipamento = EXCLUDED.id_equipamento,
                    frota = EXCLUDED.frota,
                    modelo = EXCLUDED.modelo,
                    tipo = EXCLUDED.tipo,
                    sincronizado_em = NOW()
            """, (placa, id_veiculo, v.get('idEquipamento'), v.get('frota'),
                  v.get('modelo'), v.get('tipo')))
            # Limpa posição órfã da placa antiga (mesma identidade) — senão o veículo
            # aparece 2× no mapa após a troca p/ Mercosul.
            cur.execute(
                "DELETE FROM embarques_posicoes_atuais WHERE id_veiculo_3s=%s AND placa<>%s",
                (id_veiculo, placa))
            if existe:
                atualizados += 1
            else:
                novos += 1
        conn.commit()
        cur.close(); conn.close()
        return jsonify({'ok': True, 'total': len(veiculos), 'novos': novos, 'atualizados': atualizados})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/rastreamento/health')
@admin_required
def api_rastreamento_health():
    """Status do worker, última sync, token, contadores."""
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT MAX(atualizado_em) FROM embarques_posicoes_atuais")
        ultima_sync = cur.fetchone()[0]
        cur.execute("SELECT expiration FROM embarques_3s_token WHERE id=1")
        r = cur.fetchone()
        token_valido_ate = (r[0].isoformat() + 'Z') if r else None
        cur.execute("SELECT COUNT(*) FROM embarques_3s_log WHERE chamado_em > NOW() - INTERVAL '60 seconds'")
        chamadas_60s = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM embarques_3s_log WHERE provider='ORS' AND chamado_em > NOW() - INTERVAL '24 hours'")
        ors_24h = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM embarques_3s_log WHERE erro_codigo IS NOT NULL AND chamado_em > NOW() - INTERVAL '24 hours'")
        erros_24h = cur.fetchone()[0]
        cur.close(); conn.close()
        return jsonify({
            'ok': True,
            'worker_running': rastreamento_worker.is_running(),
            'modo_simulado': tres_s_client.is_modo_simulado(),
            'ultima_sync': (ultima_sync.isoformat() + 'Z') if ultima_sync else None,
            'token_valido_ate': token_valido_ate,
            'chamadas_60s': chamadas_60s,
            'ors_chamadas_24h': ors_24h,
            'erros_24h': erros_24h,
        })
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/api/rastreamento/log')
@admin_required
def api_rastreamento_log():
    """Últimas 200 linhas do log (3S + ORS + SIM)."""
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("""
            SELECT chamado_em, provider, endpoint, duracao_ms, status_http, erro_codigo, erro_msg
            FROM embarques_3s_log
            ORDER BY chamado_em DESC LIMIT 200
        """)
        cols = ['chamado_em', 'provider', 'endpoint', 'duracao_ms', 'status_http', 'erro_codigo', 'erro_msg']
        data = []
        for r in cur.fetchall():
            d = dict(zip(cols, r))
            d['chamado_em'] = d['chamado_em'].isoformat() + 'Z'
            data.append(d)
        cur.close(); conn.close()
        return jsonify({'ok': True, 'data': data, 'count': len(data)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


if __name__ == '__main__':
    print("\n⚡ Auditoria Receita — Backend")
    print("=" * 40)

    missing = [k for k, v in CONFIG.items() if not v]
    if missing:
        print(f"\n⚠️  Variáveis Power BI faltando no .env: {', '.join(missing)}")
    else:
        print("✅ Configuração Power BI OK")

    # Cache do cadastro de veículos para o PGR. Fica aqui (e não no worker)
    # porque é este lado que fala Power BI. A cada 12h, então nunca passa disso
    # de defasagem e não precisa combinar relógio com o job de apuração.
    if os.getenv('PGR_SYNC_CADASTRO', 'true').lower() == 'true':
        import threading as _th

        def _loop_cadastro_pgr():
            while True:
                try:
                    n = sincronizar_cadastro_pgr()
                    m = sincronizar_manifestos_pgr()
                    print(f"✅ PGR: cadastro sincronizado ({n} veículos, {m} manifestos)")
                except Exception as e:
                    # Falha aqui é macia: a apuração segue e o rótulo de tipo
                    # sai em branco. Não pode derrubar o processo.
                    print(f"⚠️  PGR: falha ao sincronizar cadastro: {e}")
                time.sleep(12 * 3600)

        _th.Thread(target=_loop_cadastro_pgr, daemon=True, name='PgrCadastro').start()

    # Lançamento automático de embarques a partir do manifesto do SSW.
    # Fica aqui (e não no worker) pelo mesmo motivo do cache do PGR: é este lado
    # que fala Power BI. Roda 1x/dia depois da carga full-refresh do SSW (~05:10).
    # A chave EMBARQUES_AUTO desliga o motor sem deploy — se o operacional voltar
    # a lançar à mão, basta virar a variável no Portainer e reiniciar.
    if embarques_auto.ligado():
        import threading as _th_ea

        def _loop_embarques_auto():
            # datetime é importado localmente em todo o arquivo (não há import global)
            from datetime import datetime, timedelta
            hora = os.getenv('EMBARQUES_AUTO_HORA_BRT', '07:00')
            janela = int(os.getenv('EMBARQUES_AUTO_JANELA_DISPARO_MIN', '180'))
            try:
                _hh, _mm = [int(x) for x in hora.split(':')]
            except Exception:
                _hh, _mm = 7, 0
            ultimo_dia = None
            while True:
                try:
                    # BRT = UTC-3. Sem tz database: o container roda em UTC e a
                    # única coisa que importa é disparar DEPOIS da carga do SSW.
                    agora = datetime.utcnow() - timedelta(hours=3)
                    minutos = agora.hour * 60 + agora.minute
                    alvo_min = _hh * 60 + _mm
                    na_janela = alvo_min <= minutos <= alvo_min + janela
                    if na_janela and ultimo_dia != agora.date():
                        r = embarques_auto.executar()
                        ultimo_dia = agora.date()
                        if r.get('ok'):
                            print(f"✅ Embarques auto: {r['criadas']} criada(s), "
                                  f"{sum(r['fechadas'].values())} encerrada(s), "
                                  f"{r['reconciliadas']} reconciliada(s) "
                                  f"[janela {r['janela'][0]}..{r['janela'][1]}]")
                except Exception as e:
                    # Falha aqui não pode derrubar o processo: o pior caso é o dia
                    # ficar sem lançamento automático, e o índice único garante que
                    # a rodada seguinte não duplica o que já entrou.
                    print(f"⚠️  Embarques auto: falha na execução: {e}")
                time.sleep(600)

        _th_ea.Thread(target=_loop_embarques_auto, daemon=True, name='EmbarquesAuto').start()
        print("✅ Lançamento automático de embarques LIGADO "
              f"(diário às {os.getenv('EMBARQUES_AUTO_HORA_BRT', '07:00')} BRT)")
    else:
        print("ℹ️  Lançamento automático de embarques desligado (EMBARQUES_AUTO)")

    # Boot do worker de rastreamento
    if os.getenv('START_WORKER', '').lower() == 'true':
        try:
            rastreamento_worker.start()
            modo = 'SIMULADO' if tres_s_client.is_modo_simulado() else 'REAL'
            print(f"✅ Worker de rastreamento iniciado (modo {modo})")
        except Exception as e:
            print(f"⚠️  Worker não iniciou: {e}")
    else:
        print("ℹ️  Worker de rastreamento desligado (START_WORKER != true)")

    print(f"\n🌐 Acesse: http://localhost:5000\n")
    app.run(host='0.0.0.0', port=5000, debug=True, use_reloader=False)
